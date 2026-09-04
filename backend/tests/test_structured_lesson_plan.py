import sys
import unittest
import zipfile
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from unittest.mock import patch
from uuid import uuid4
import shutil

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from lesson_plan_docx import render_structured_lesson_plan_docx
from structured_documents import ValidationError, normalise_create_resources_result, structured_lesson_plan_validation_summary, validate_structured_lesson_plan_document


def lesson_plan_payload():
    return {
        "document": {
            "title": "Photosynthesis: energy transfer",
            "subject": "Biology",
            "level": "Leaving Certificate",
            "class_context": "5th Year",
            "duration": "55 minutes",
            "primary_outcome": "Students explain how light energy is converted into chemical energy during photosynthesis.",
            "learning_intentions": ["Describe the reactants and products of photosynthesis."],
            "success_criteria": ["I can use the word equation accurately."],
            "definitions": [{"term": "Chlorophyll", "definition": "A pigment that absorbs light energy."}],
            "resources": ["Equation handout", "Mini whiteboards"],
            "prior_knowledge": ["Plant cells contain chloroplasts."],
            "lesson_flow": [
                {"minutes": "0-8 min", "phase": "Starter", "teacher_action": "Elicit the plant-cell structures needed for photosynthesis.", "student_action": "Annotate a chloroplast diagram.", "check_for_understanding": "Students hold up the correct labelled diagram."},
                {"minutes": "8-45 min", "phase": "Development", "teacher_action": "Model the word equation and question likely misconceptions.", "student_action": "Complete and explain the equation in pairs.", "check_for_understanding": "Pairs justify where the energy enters the process."},
            ],
            "differentiation": ["Provide a labelled word bank where needed."],
            "misconceptions": ["Plants do not get their food from soil."],
            "assessment": ["Ask each student to complete the word equation independently."],
            "teacher_note": "Use the source vocabulary consistently before introducing the balanced chemical equation.",
            "homework": "Explain why a destarched leaf is used in the starch test.",
            "stopping_point": "Begin the chemical equation next lesson.",
        }
    }


class StructuredLessonPlanTests(unittest.TestCase):
    class _WorkspaceTempDirectory:
        def __init__(self):
            self.path = Path.cwd() / f".structured-pdf-test-{uuid4().hex}"

        def __enter__(self):
            self.path.mkdir()
            return str(self.path)

        def __exit__(self, exc_type, exc_value, traceback):
            shutil.rmtree(self.path)

    def test_valid_structured_lesson_plan_builds_document_and_export_text(self):
        result = normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback")

        self.assertEqual(result["title"], "Photosynthesis: energy transfer")
        self.assertEqual(result["document"]["schema_version"], 1)
        self.assertEqual(result["document"]["resource_type"], "lesson_plan")
        self.assertIn("## Lesson Flow", result["content"])
        self.assertIn("Photosynthesis: energy transfer", result["content"])

    def test_missing_required_core_field_is_rejected(self):
        payload = lesson_plan_payload()
        del payload["document"]["primary_outcome"]

        with self.assertRaises(ValidationError):
            normalise_create_resources_result("lesson_plan", payload, "Fallback")

    def test_ai_null_optional_lists_are_normalised_before_validation(self):
        payload = lesson_plan_payload()
        for field in ("definitions", "resources", "prior_knowledge", "differentiation", "misconceptions", "assessment"):
            payload["document"][field] = None

        result = normalise_create_resources_result("lesson_plan", payload, "Fallback")

        self.assertEqual(result["title"], "Photosynthesis: energy transfer")
        self.assertIn("Lesson flow", [block.get("title") for block in result["document"]["blocks"]])

    def test_reasonable_ai_lesson_plan_aliases_normalise_to_the_canonical_schema(self):
        payload = lesson_plan_payload()["document"]
        payload["learning_objectives"] = payload.pop("learning_intentions")
        payload["key_definitions"] = [{"name": "Chlorophyll", "meaning": "A pigment that absorbs light energy."}]
        del payload["definitions"]
        payload["lesson_timeline"] = payload.pop("lesson_flow")
        for item in payload["lesson_timeline"]:
            item["time"] = item.pop("minutes")
            item["teacher_activity"] = item.pop("teacher_action")
            item["student_activity"] = item.pop("student_action")
            item["check"] = item.pop("check_for_understanding")

        result = normalise_create_resources_result("lesson_plan", payload, "Fallback")

        timeline = next(block for block in result["document"]["blocks"] if block["type"] == "timeline")
        self.assertEqual(timeline["items"][0]["minutes"], "0-8 min")
        self.assertEqual(timeline["items"][0]["teacher_action"], "Elicit the plant-cell structures needed for photosynthesis.")
        definitions = next(block for block in result["document"]["blocks"] if block.get("label") == "Key definitions")
        self.assertEqual(definitions["definitions"][0]["term"], "Chlorophyll")

    def test_junior_cycle_outcome_language_normalises_to_the_canonical_structured_document(self):
        payload = lesson_plan_payload()["document"]
        payload["level"] = "Junior Cycle"
        payload["learning_outcomes"] = payload.pop("learning_intentions")
        payload["success_indicators"] = payload.pop("success_criteria")
        payload["assessment_for_learning"] = payload.pop("assessment")
        payload["differentiation_support"] = payload.pop("differentiation")

        result = normalise_create_resources_result("lesson_plan", payload, "Junior Cycle fallback")

        self.assertEqual(result["document"]["resource_type"], "lesson_plan")
        self.assertEqual(result["document"]["level"], "Junior Cycle")
        self.assertEqual(
            [block.get("title") or block.get("label") for block in result["document"]["blocks"]],
            [
                "Primary learning outcome",
                "Learning intentions",
                "Success criteria",
                "Key definitions",
                "Prior knowledge",
                "Lesson flow",
                "Resources",
                "Differentiation and support",
                "Misconceptions to address",
                "Assessment and checks for understanding",
                "Teacher reference",
                "Suggested homework",
                "Stopping point / next lesson",
            ],
        )

    def test_direct_lesson_plan_document_is_accepted_but_invalid_content_still_fails(self):
        direct_document = lesson_plan_payload()["document"]
        result = normalise_create_resources_result("lesson_plan", direct_document, "Fallback")
        self.assertEqual(result["title"], "Photosynthesis: energy transfer")

        direct_document["lesson_flow"] = [{"minutes": "0-5", "phase": "Starter", "teacher_action": "Ask", "student_action": "Answer"}]
        with self.assertRaises(ValidationError):
            normalise_create_resources_result("lesson_plan", direct_document, "Fallback")

    def test_ai_whole_minute_variants_normalise_to_canonical_minute_strings(self):
        payload = lesson_plan_payload()
        payload["document"]["lesson_flow"] = [
            {"minutes": 5, "phase": "Starter", "teacher_action": "Set the retrieval task.", "student_action": "Complete the retrieval task."},
            {"minutes": "10", "phase": "Model", "teacher_action": "Model the worked example.", "student_action": "Annotate the worked example."},
            {"minutes": "15 minutes", "phase": "Practice", "teacher_action": "Guide practice.", "student_action": "Complete guided questions."},
            {"minutes": "5 minute", "phase": "Check", "teacher_action": "Ask checking questions.", "student_action": "Explain answers."},
            {"minutes": "3 min", "phase": "Plenary", "teacher_action": "Set an exit prompt.", "student_action": "Answer the exit prompt."},
            {"minutes": "2 mins", "phase": "Close", "teacher_action": "Summarise learning.", "student_action": "Record the next step."},
        ]

        result = normalise_create_resources_result("lesson_plan", payload, "Fallback")
        timeline = next(block for block in result["document"]["blocks"] if block["type"] == "timeline")
        self.assertEqual([item["minutes"] for item in timeline["items"]], ["5", "10", "15", "5", "3", "2"])

    def test_existing_range_minute_strings_remain_compatible(self):
        payload = lesson_plan_payload()
        result = normalise_create_resources_result("lesson_plan", payload, "Fallback")

        timeline = next(block for block in result["document"]["blocks"] if block["type"] == "timeline")
        self.assertEqual([item["minutes"] for item in timeline["items"]], ["0-8 min", "8-45 min"])

    def test_requested_lesson_duration_accepts_matching_total_and_canonicalises_duration(self):
        payload = lesson_plan_payload()
        payload["document"]["duration"] = "58 min"
        payload["document"]["lesson_flow"] = [
            {"minutes": 8, "phase": "Starter", "teacher_action": "Set retrieval work.", "student_action": "Complete retrieval work."},
            {"minutes": "35 minutes", "phase": "Development", "teacher_action": "Model the new content.", "student_action": "Complete guided practice."},
            {"minutes": "15", "phase": "Plenary", "teacher_action": "Lead the final check.", "student_action": "Complete an exit response."},
        ]

        result = normalise_create_resources_result("lesson_plan", payload, "Fallback", expected_duration_minutes=58)
        timeline = next(block for block in result["document"]["blocks"] if block["type"] == "timeline")
        self.assertEqual(result["document"]["duration"], "58 minutes")
        self.assertEqual([item["minutes"] for item in timeline["items"]], ["8", "35", "15"])

    def test_requested_lesson_duration_rejects_non_matching_total(self):
        payload = lesson_plan_payload()
        payload["document"]["duration"] = "58 minutes"
        payload["document"]["lesson_flow"] = [
            {"minutes": 8, "phase": "Starter", "teacher_action": "Set retrieval work.", "student_action": "Complete retrieval work."},
            {"minutes": 35, "phase": "Development", "teacher_action": "Model the new content.", "student_action": "Complete guided practice."},
        ]

        with self.assertRaisesRegex(ValueError, "lesson flow duration total mismatch"):
            normalise_create_resources_result("lesson_plan", payload, "Fallback", expected_duration_minutes=58)

    def test_requested_lesson_duration_rejects_range_flow_values_without_affecting_legacy_documents(self):
        with self.assertRaisesRegex(ValueError, "invalid lesson flow minutes for requested duration"):
            normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback", expected_duration_minutes=55)

    def test_ai_invalid_non_text_lesson_flow_minutes_are_rejected(self):
        for invalid_minutes in (5.5, 0, -5, True):
            with self.subTest(invalid_minutes=invalid_minutes):
                payload = lesson_plan_payload()
                payload["document"]["lesson_flow"][0]["minutes"] = invalid_minutes
                with self.assertRaisesRegex(ValueError, "invalid lesson flow minutes"):
                    normalise_create_resources_result("lesson_plan", payload, "Fallback")

    def test_validation_summary_contains_only_field_names(self):
        payload = lesson_plan_payload()
        payload["document"]["learning_intentions"] = []
        try:
            normalise_create_resources_result("lesson_plan", payload, "Fallback")
        except ValidationError as exc:
            summary = structured_lesson_plan_validation_summary(exc)
        else:
            self.fail("Expected structured validation to fail")

        self.assertIn("learning_intentions", summary)
        self.assertNotIn("Photosynthesis", summary)

    def test_excessive_lesson_flow_is_rejected(self):
        payload = lesson_plan_payload()
        payload["document"]["lesson_flow"] *= 5

        with self.assertRaises(ValidationError):
            normalise_create_resources_result("lesson_plan", payload, "Fallback")

    def test_malformed_lesson_plan_document_is_rejected(self):
        with self.assertRaises(ValueError):
            normalise_create_resources_result("lesson_plan", {"title": "Not structured", "content": "Legacy text"}, "Fallback")

    def test_lesson_plan_document_rejects_unknown_fields(self):
        payload = lesson_plan_payload()
        payload["document"]["colour_scheme"] = "emerald"

        with self.assertRaises(ValidationError):
            normalise_create_resources_result("lesson_plan", payload, "Fallback")

    def test_non_lesson_plan_modes_keep_legacy_response_shape(self):
        result = normalise_create_resources_result("worksheet", {"title": "Cells worksheet", "content": "Questions"}, "Fallback")

        self.assertEqual(result, {"title": "Cells worksheet", "content": "Questions"})

    def test_structured_lesson_plan_docx_contains_document_sections_and_tables(self):
        from docx import Document

        result = normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback")
        document = validate_structured_lesson_plan_document(result["document"])
        data = render_structured_lesson_plan_docx(document, teacher="Ms Example", meta={"schoolName": "Example School"})

        self.assertTrue(zipfile.is_zipfile(BytesIO(data)))
        reopened = Document(BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in reopened.paragraphs)
        table_text = "\n".join(cell.text for table in reopened.tables for row in table.rows for cell in row.cells)
        self.assertIn("Photosynthesis: energy transfer", text)
        self.assertIn("PRIMARY LEARNING OUTCOME", table_text)
        self.assertIn("Time", table_text)
        self.assertIn("Teacher", table_text)
        self.assertIn("Chlorophyll", table_text)
        self.assertIn("Suggested homework", table_text)
        self.assertIn("Misconceptions to address", table_text)
        self.assertIn("Assessment and checks for understanding", table_text)

    def test_structured_lesson_plan_docx_embeds_selected_elume_logo_in_footer(self):
        result = normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback")
        data = render_structured_lesson_plan_docx(
            validate_structured_lesson_plan_document(result["document"]),
            teacher="Ms Example",
            meta={"brandingChoice": "elume", "schoolName": "Example School"},
        )

        with zipfile.ZipFile(BytesIO(data)) as archive:
            media_files = [name for name in archive.namelist() if name.startswith("word/media/")]
            header_files = [name for name in archive.namelist() if name.startswith("word/header") and name.endswith(".xml")]
            footer_files = [name for name in archive.namelist() if name.startswith("word/footer") and name.endswith(".xml")]
            self.assertTrue(media_files)
            self.assertTrue(header_files)
            self.assertTrue(footer_files)
            self.assertNotIn(b"a:blip", archive.read(header_files[0]))
            self.assertIn(b"a:blip", archive.read(footer_files[0]))

        from docx import Document

        reopened = Document(BytesIO(data))
        self.assertEqual(reopened.sections[0].header.paragraphs[0].text, "")
        self.assertEqual(reopened.sections[0].footer.paragraphs[0].text, "Ms Example  |  Example School\t")

    def test_structured_lesson_plan_docx_embeds_valid_school_logo_data_url(self):
        school_logo = "data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVQIHWP4z8DwHwAFgAI/ScL+XQAAAABJRU5ErkJggg=="
        result = normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback")
        data = render_structured_lesson_plan_docx(
            validate_structured_lesson_plan_document(result["document"]),
            meta={"brandingChoice": "school", "schoolLogoDataUrl": school_logo},
        )

        with zipfile.ZipFile(BytesIO(data)) as archive:
            self.assertTrue([name for name in archive.namelist() if name.startswith("word/media/")])
            footer_files = [name for name in archive.namelist() if name.startswith("word/footer") and name.endswith(".xml")]
            self.assertIn(b"a:blip", archive.read(footer_files[0]))

    def test_structured_lesson_plan_docx_ignores_bad_or_unselected_logo_data(self):
        result = normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback")
        document = validate_structured_lesson_plan_document(result["document"])
        for meta in (
            {"brandingChoice": "school", "schoolLogoDataUrl": "data:image/png;base64,not-a-logo"},
            {"brandingChoice": "none", "schoolLogoDataUrl": "data:image/png;base64,not-a-logo"},
        ):
            with self.subTest(meta=meta):
                data = render_structured_lesson_plan_docx(document, meta=meta)
                with zipfile.ZipFile(BytesIO(data)) as archive:
                    self.assertFalse([name for name in archive.namelist() if name.startswith("word/media/")])

    def test_structured_lesson_plan_docx_allows_absent_optional_sections(self):
        payload = lesson_plan_payload()
        for field in ("definitions", "resources", "prior_knowledge", "differentiation", "misconceptions", "assessment"):
            payload["document"][field] = []
        payload["document"]["teacher_note"] = None
        payload["document"]["homework"] = None
        payload["document"]["stopping_point"] = None

        result = normalise_create_resources_result("lesson_plan", payload, "Fallback")
        data = render_structured_lesson_plan_docx(validate_structured_lesson_plan_document(result["document"]))
        self.assertTrue(zipfile.is_zipfile(BytesIO(data)))

    def test_export_document_rejects_malformed_renderer_blocks(self):
        result = normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback")
        result["document"]["blocks"][0]["unexpected"] = "not allowed"

        with self.assertRaises(ValueError):
            validate_structured_lesson_plan_document(result["document"])

    def test_legacy_docx_exporter_remains_available(self):
        from main import _docx_from_markdownish

        data = _docx_from_markdownish("Legacy worksheet", "# Questions\n\n- Explain the process.")
        self.assertTrue(zipfile.is_zipfile(BytesIO(data)))

    def test_structured_lesson_plan_pdf_export_uses_canonical_docx_then_conversion(self):
        from main import ExportDocxRequest, export_pdf

        result = normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback")
        with patch("main.render_structured_lesson_plan_docx", return_value=b"canonical-docx") as render_docx, patch(
            "main._convert_structured_lesson_plan_docx_to_pdf", return_value=b"%PDF-1.7\nstructured"
        ) as convert_pdf:
            response = export_pdf(
                ExportDocxRequest(
                    title=result["title"],
                    content=result["content"],
                    document=result["document"],
                    teacher="Ms Example",
                )
            )

        render_docx.assert_called_once()
        convert_pdf.assert_called_once_with(b"canonical-docx")
        self.assertEqual(response.media_type, "application/pdf")
        self.assertIn('.pdf"', response.headers["content-disposition"])

    def test_structured_pdf_conversion_uses_isolated_temp_docx_and_cleans_up(self):
        from main import _convert_structured_lesson_plan_docx_to_pdf

        observed = {}

        def fake_run(args, **kwargs):
            source_path = Path(args[-1])
            converted_dir = Path(args[args.index("--outdir") + 1])
            observed["source_path"] = source_path
            observed["converted_dir"] = converted_dir
            observed["args"] = args
            self.assertEqual(source_path.read_bytes(), b"canonical-docx")
            (converted_dir / "lesson-plan.pdf").write_bytes(b"%PDF-1.7\nconverted")
            return SimpleNamespace(returncode=0)

        temporary_directory = self._WorkspaceTempDirectory()
        with patch("main.tempfile.TemporaryDirectory", return_value=temporary_directory) as temp_factory, patch(
            "main._libreoffice_command", return_value="libreoffice"
        ), patch("main.subprocess.run", side_effect=fake_run):
            result = _convert_structured_lesson_plan_docx_to_pdf(b"canonical-docx")

        self.assertTrue(result.startswith(b"%PDF-"))
        temp_factory.assert_called_once_with(prefix="elume-structured-pdf-")
        self.assertEqual(observed["args"][0], "libreoffice")
        self.assertIn("--headless", observed["args"])
        self.assertFalse(observed["source_path"].parent.exists())
        self.assertFalse(observed["converted_dir"].exists())

    def test_structured_pdf_conversion_failure_is_controlled_and_cleans_up(self):
        from fastapi import HTTPException
        from main import _convert_structured_lesson_plan_docx_to_pdf

        observed = {}

        def fake_run(args, **kwargs):
            observed["source_path"] = Path(args[-1])
            return SimpleNamespace(returncode=1)

        temporary_directory = self._WorkspaceTempDirectory()
        with patch("main.tempfile.TemporaryDirectory", return_value=temporary_directory), patch(
            "main._libreoffice_command", return_value="libreoffice"
        ), patch("main.subprocess.run", side_effect=fake_run):
            with self.assertRaises(HTTPException) as raised:
                _convert_structured_lesson_plan_docx_to_pdf(b"canonical-docx")

        self.assertEqual(raised.exception.status_code, 422)
        self.assertNotIn("elume-structured-pdf-", raised.exception.detail)
        self.assertFalse(observed["source_path"].parent.exists())

    def test_structured_pdf_export_rejects_invalid_document_and_legacy_pdf_still_works(self):
        from fastapi import HTTPException
        from main import ExportDocxRequest, _pdf_from_markdownish, export_pdf

        malformed = normalise_create_resources_result("lesson_plan", lesson_plan_payload(), "Fallback")
        malformed["document"]["blocks"][0]["unexpected"] = "invalid"
        with self.assertRaises(HTTPException) as raised:
            export_pdf(ExportDocxRequest(title="Bad", content="", document=malformed["document"]))
        self.assertEqual(raised.exception.status_code, 422)

        legacy_pdf = _pdf_from_markdownish("Legacy worksheet", "# Questions\n\n- Explain the process.")
        self.assertTrue(legacy_pdf.startswith(b"%PDF-"))


if __name__ == "__main__":
    unittest.main()
