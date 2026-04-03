from __future__ import annotations

import base64
import json, os, uuid
import hashlib
import re
import shutil
import random
import string
import struct
import textwrap
import zipfile
import zlib
from datetime import datetime, timedelta
from pathlib import Path
from typing import Any, List, Optional
from xml.etree import ElementTree as ET
import stripe

from copy import deepcopy

import json
from collections import Counter, defaultdict
from fastapi import WebSocket, WebSocketDisconnect

from dotenv import load_dotenv

from pathlib import Path



BASE_DIR = Path(__file__).resolve().parent
UPLOADS_DIR = BASE_DIR.parent / "uploads"
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)

load_dotenv(BASE_DIR / ".env")

from fastapi import Depends, FastAPI, File, Form, HTTPException, Request, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from sqlalchemy.orm import Session
from sqlalchemy import and_, Column, Integer, String, Text, Boolean, func
from sqlalchemy import text
from fastapi import Header
from passlib.context import CryptContext
from jose import jwt, JWTError
import secrets

import logging

logger = logging.getLogger(__name__)

from io import BytesIO
from fastapi.responses import FileResponse, StreamingResponse

from uuid import uuid4
from datetime import datetime, timedelta
from uuid import uuid4

from datetime import datetime, timedelta

from models import CollabSessionModel, CollabParticipantModel
from schemas import (
    CollabCreatePayload,
    CollabCreateResponse,
    CollabJoinPayload,
    CollabJoinResponse,
    CollabAssignmentsPayload,
    CollabStatusResponse,
)


import models  # IMPORTANT: needed because we reference models.Topic, models.Note, etc.
import schemas
from db import Base, SessionLocal, engine

from models import (
    ClassModel,
    PostModel,
    TestCategory,
    TestItem,
    StudentModel,
    ClassAssessmentModel,
    AssessmentResultModel,
    Cat4BaselineSetModel,
    Cat4StudentBaselineModel,
    Cat4TermResultSetModel,
    Cat4StudentTermResultModel,
    Cat4WorkbookVersionModel,
    LiveQuizSessionModel,
    LiveQuizParticipantModel,
    LiveQuizAnswerModel,
    LiveQuizAttemptModel,

)

def _dev_autologin_enabled() -> bool:
    return os.getenv("DEV_AUTO_LOGIN", "0").strip().lower() in {"1", "true", "yes", "on"}

def _is_local_request(host: Optional[str]) -> bool:
    host = (host or "").split(":")[0].strip().lower()
    return host in {"localhost", "127.0.0.1"}

def _app_env() -> str:
    return (
        os.getenv("APP_ENV")
        or os.getenv("ENV")
        or ""
    ).strip().lower()

def _is_explicit_development() -> bool:
    return _app_env() in {"development", "dev", "local"}

def _jwt_secret_is_weak(secret: str) -> bool:
    return not secret or secret == "change-me"


# =========================================================
# APP
# =========================================================
app = FastAPI()
print("✅ LOADED main.py FROM:", __file__)

PWD_CONTEXT = CryptContext(schemes=["bcrypt"], deprecated="auto")
JWT_SECRET = (os.getenv("JWT_SECRET") or "").strip()
APP_BASE_URL = (os.getenv("APP_BASE_URL") or "http://localhost:5173").strip()
JWT_ALG = "HS256"
JWT_EXPIRE_DAYS = 30

STRIPE_SECRET_KEY = (os.getenv("STRIPE_SECRET_KEY") or "").strip()
STRIPE_PRICE_MONTHLY_EUR = (os.getenv("STRIPE_PRICE_MONTHLY_EUR") or "").strip()
STRIPE_PRICE_ANNUAL_EUR = (os.getenv("STRIPE_PRICE_ANNUAL_EUR") or "").strip()
STRIPE_WEBHOOK_SECRET = (os.getenv("STRIPE_WEBHOOK_SECRET") or "").strip()

stripe.api_key = STRIPE_SECRET_KEY


if _jwt_secret_is_weak(JWT_SECRET) and not _is_explicit_development():
    raise RuntimeError(
        "JWT_SECRET is missing or insecure. Set a strong JWT_SECRET or run only in explicit development mode."
    )

class DevAutoLoginResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"

class AuthRegister(BaseModel):
    first_name: str
    last_name: str
    school_name: str
    email: str
    password: str

class AuthLogin(BaseModel):
    email: str
    password: str

class AuthToken(BaseModel):
    access_token: str
    token_type: str = "bearer"


class AuthRegisterResponse(BaseModel):
    success: bool
    message: str


class VerifyEmailPayload(BaseModel):
    token: str


class VerifyEmailResponse(BaseModel):
    success: bool
    message: str
    access_token: str
    token_type: str = "bearer"
    next_path: str = "/onboarding/billing"

class AdminCreateUser(BaseModel):
    email: str
    password: str

class AdminResetPassword(BaseModel):
    email: str
    new_password: str

class AdminRenameUser(BaseModel):
    old_email: str
    new_email: str


class AdminDeleteUser(BaseModel):
    email: str
    hard_delete: bool = False


class AdminTransferClassPayload(BaseModel):
    class_id: int
    target_email: str


class Cat4BaselineSetCreatePayload(BaseModel):
    title: str
    test_date: Optional[str] = None


class Cat4BaselineRowInput(BaseModel):
    raw_name: str
    verbal_sas: Optional[int] = None
    quantitative_sas: Optional[int] = None
    non_verbal_sas: Optional[int] = None
    spatial_sas: Optional[int] = None
    overall_sas: Optional[int] = None
    profile_label: Optional[str] = None
    confidence_note: Optional[str] = None


class Cat4BaselineRowsPayload(BaseModel):
    rows: List[Cat4BaselineRowInput]


class Cat4TermResultSetCreatePayload(BaseModel):
    title: str
    academic_year: Optional[str] = None
    term_key: Optional[str] = None


class Cat4TermResultRowInput(BaseModel):
    raw_name: str
    average_percent: Optional[int] = None
    subject_count: Optional[int] = None
    raw_subjects_json: Optional[Any] = None


class Cat4TermResultRowsPayload(BaseModel):
    rows: List[Cat4TermResultRowInput]

class CreateCheckoutSessionRequest(BaseModel):
    plan: str  # "monthly" | "annual"


class CreateCheckoutSessionResponse(BaseModel):
    checkout_url: str


class CreatePortalSessionResponse(BaseModel):
    portal_url: str


class StartTrialResponse(BaseModel):
    success: bool
    message: str


class ConfirmCheckoutSessionResponse(BaseModel):
    success: bool
    billing_status: dict[str, Any]


def _stripe_obj_get(obj: Any, key: str, default: Any = None) -> Any:
    if obj is None:
        return default
    if isinstance(obj, dict):
        return obj.get(key, default)
    getter = getattr(obj, "get", None)
    if callable(getter):
        return getter(key, default)
    return getattr(obj, key, default)


def _stripe_epoch_to_datetime(value: Any) -> Optional[datetime]:
    if value in (None, ""):
        return None
    try:
        return datetime.utcfromtimestamp(int(value))
    except Exception:
        return None


def _password_policy_error(password: str) -> Optional[str]:
    if len(password) < 8:
        return "Password must be at least 8 characters"
    if not any(ch.isupper() for ch in password):
        return "Password must include at least one uppercase letter"
    if not any(ch.islower() for ch in password):
        return "Password must include at least one lowercase letter"
    if not any(ch.isdigit() for ch in password):
        return "Password must include at least one number"
    return None


def _resolve_billing_interval(obj: Any, fallback: Optional[str] = None) -> Optional[str]:
    metadata = _stripe_obj_get(obj, "metadata") or {}
    plan = (metadata.get("plan") or "").strip().lower()
    if plan in {"monthly", "annual"}:
        return plan

    items = _stripe_obj_get(obj, "items") or {}
    data = _stripe_obj_get(items, "data") or []
    if data:
        first_item = data[0]
        price = _stripe_obj_get(first_item, "price") or {}
        recurring = _stripe_obj_get(price, "recurring") or {}
        interval = (_stripe_obj_get(recurring, "interval") or "").strip().lower()
        if interval == "month":
            interval_count = int(_stripe_obj_get(recurring, "interval_count", 1) or 1)
            return "annual" if interval_count >= 12 else "monthly"
        if interval == "year":
            return "annual"

    return fallback


def _utcnow() -> datetime:
    return datetime.utcnow()


def _is_paid_subscription_active(user: models.UserModel) -> bool:
    return (user.subscription_status or "").strip().lower() == "active"


def _is_subscription_trialing(user: models.UserModel) -> bool:
    return (user.subscription_status or "").strip().lower() == "trialing"


def _has_billing_access(user: models.UserModel) -> bool:
    return (user.subscription_status or "").strip().lower() in {"active", "trialing"}


def _is_trial_active(user: models.UserModel) -> bool:
    if _is_subscription_trialing(user):
        return bool(user.trial_ends_at and user.trial_ends_at > _utcnow())
    if _is_paid_subscription_active(user):
        return False
    return bool(user.trial_ends_at and user.trial_ends_at > _utcnow())


def _refresh_ai_daily_limit(user: models.UserModel) -> int:
    # TEMP: AI limits disabled for launch week
    target = 9999
    if user.ai_daily_limit != target:
        user.ai_daily_limit = target
    return target


def _reset_ai_prompt_counter_if_needed(user: models.UserModel) -> None:
    today = _utcnow().date()
    existing = user.ai_prompt_count_date.date() if user.ai_prompt_count_date else None
    if existing != today:
        user.ai_prompt_count = 0
        user.ai_prompt_count_date = _utcnow()


def _billing_status_payload(user: models.UserModel) -> dict[str, Any]:
    _refresh_ai_daily_limit(user)
    _reset_ai_prompt_counter_if_needed(user)
    return {
        "subscription_status": user.subscription_status or "inactive",
        "billing_interval": user.billing_interval,
        "current_period_end": user.current_period_end,
        "has_stripe_customer": bool(user.stripe_customer_id),
        "billing_onboarding_required": bool(user.billing_onboarding_required),
        "trial_started_at": user.trial_started_at,
        "trial_ends_at": user.trial_ends_at,
        "trial_active": _is_trial_active(user),
        "prompt_usage_today": int(user.ai_prompt_count or 0),
        "prompt_limit_today": int(user.ai_daily_limit or 0),
    }


def _is_missing_stripe_customer_error(exc: Exception) -> bool:
    return isinstance(exc, stripe.error.InvalidRequestError) and (
        getattr(exc, "code", None) == "resource_missing" or "No such customer" in str(exc)
    )


def _validate_stored_stripe_customer(
    db: Session,
    user: models.UserModel,
    *,
    commit_on_clear: bool = False,
) -> Optional[str]:
    customer_id = (user.stripe_customer_id or "").strip()
    if not customer_id:
        return None

    try:
        stripe.Customer.retrieve(customer_id)
        return customer_id
    except Exception as exc:
        if _is_missing_stripe_customer_error(exc):
            user.stripe_customer_id = None
            if commit_on_clear:
                db.commit()
            return None
        raise


def _enforce_ai_prompt_limit(db: Session, user: models.UserModel) -> None:
    _refresh_ai_daily_limit(user)
    _reset_ai_prompt_counter_if_needed(user)
    limit = int(user.ai_daily_limit or 0)
    used = int(user.ai_prompt_count or 0)
    if limit <= 0:
        raise HTTPException(status_code=403, detail="AI access requires an active plan or trial.")
    if used >= limit:
        raise HTTPException(status_code=403, detail=f"Today's AI prompt limit reached ({limit}/{limit}).")
    db.commit()


def _record_ai_prompt_usage(db: Session, user: models.UserModel) -> None:
    _refresh_ai_daily_limit(user)
    _reset_ai_prompt_counter_if_needed(user)
    user.ai_prompt_count = int(user.ai_prompt_count or 0) + 1
    user.ai_prompt_count_date = _utcnow()
    db.commit()


def _find_billing_user(db: Session, stripe_obj: Any) -> Optional[models.UserModel]:
    metadata = _stripe_obj_get(stripe_obj, "metadata") or {}
    user_id = _stripe_obj_get(metadata, "user_id")
    customer_id = _stripe_obj_get(stripe_obj, "customer")
    subscription_id = _stripe_obj_get(stripe_obj, "subscription")
    checkout_session_id = _stripe_obj_get(stripe_obj, "id") if _stripe_obj_get(stripe_obj, "object") == "checkout.session" else None
    email = (
        _stripe_obj_get(metadata, "email")
        or _stripe_obj_get(stripe_obj, "customer_email")
        or _stripe_obj_get(stripe_obj, "email")
    )

    if user_id:
        try:
            user = db.query(models.UserModel).filter(models.UserModel.id == int(user_id)).first()
            if user:
                return user
        except Exception:
            pass

    if subscription_id:
        user = db.query(models.UserModel).filter(models.UserModel.stripe_subscription_id == str(subscription_id)).first()
        if user:
            return user

    if customer_id:
        user = db.query(models.UserModel).filter(models.UserModel.stripe_customer_id == str(customer_id)).first()
        if user:
            return user

    if checkout_session_id:
        user = db.query(models.UserModel).filter(models.UserModel.stripe_checkout_session_id == str(checkout_session_id)).first()
        if user:
            return user

    if email:
        user = db.query(models.UserModel).filter(models.UserModel.email == str(email).strip().lower()).first()
        if user:
            return user

    return None


def _apply_subscription_update(user: models.UserModel, subscription_obj: Any) -> None:
    customer_id = _stripe_obj_get(subscription_obj, "customer")
    subscription_id = _stripe_obj_get(subscription_obj, "id")
    status = (_stripe_obj_get(subscription_obj, "status") or user.subscription_status or "inactive").strip().lower()
    created_at = _stripe_epoch_to_datetime(_stripe_obj_get(subscription_obj, "created"))
    period_end = _stripe_epoch_to_datetime(_stripe_obj_get(subscription_obj, "current_period_end"))
    trial_end = _stripe_epoch_to_datetime(_stripe_obj_get(subscription_obj, "trial_end"))
    billing_interval = _resolve_billing_interval(subscription_obj, user.billing_interval)
    metadata = _stripe_obj_get(subscription_obj, "metadata") or {}

    if customer_id:
        user.stripe_customer_id = str(customer_id)
    if subscription_id:
        user.stripe_subscription_id = str(subscription_id)

    user.subscription_status = status or "inactive"
    user.billing_interval = billing_interval
    user.current_period_end = period_end
    user.trial_ends_at = trial_end if status == "trialing" else None
    if status == "trialing" and not user.trial_started_at:
        user.trial_started_at = created_at or datetime.utcnow()

    if status in {"active", "trialing"} and not user.subscription_started_at:
        user.subscription_started_at = created_at or datetime.utcnow()
    if status in {"active", "trialing"}:
        user.billing_onboarding_required = False

    if (metadata.get("launch_offer_candidate") or "").strip().lower() == "true" and status in {"active", "trialing"}:
        user.launch_offer_applied = True

def _make_token(user_id: int, email: str) -> str:
    exp = datetime.utcnow() + timedelta(days=JWT_EXPIRE_DAYS)
    return jwt.encode(
        {"sub": str(user_id), "email": email, "exp": exp},
        JWT_SECRET,
        algorithm=JWT_ALG,
    )

def _annual_launch_offer_applies(now: datetime | None = None) -> bool:
    now = now or datetime.utcnow()

    # Adjust this date if you want a different launch window start
    offer_start = datetime(2026, 4, 1)

    return now >= offer_start

# -------------------------
# DOCX Export (editable Word doc)
# -------------------------

class ExportDocxRequest(BaseModel):
    title: str
    content: str
    teacher: str | None = None
    meta: dict | None = None  # optional: scope/level/template/date etc.


def _pdf_escape(text: str) -> str:
    safe = (text or "").replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    return safe.encode("latin-1", errors="replace").decode("latin-1")


def _pdf_wrap(text: str, width: int) -> list[str]:
    clean = " ".join((text or "").split())
    if not clean:
        return [""]
    return textwrap.wrap(clean, width=width) or [clean]


LESSON_PLAN_SECTION_ORDER = [
    "Learning Overview",
    "Learning Intentions",
    "Success Criteria",
    "Lesson Flow",
    "Resources",
    "Differentiation",
    "Assessment",
    "Suggested Homework",
    "Reflection",
]

LESSON_PLAN_SECTION_DISPLAY = {
    "Learning Overview": "Learning Overview",
    "Learning Intentions": "Learning Intentions",
    "Success Criteria": "Success Criteria",
    "Lesson Flow": "Lesson Flow",
    "Resources": "Resources",
    "Differentiation": "Differentiation",
    "Assessment": "Assessment",
    "Suggested Homework": "Suggested Homework",
    "Reflection": "Reflection",
}

LESSON_FLOW_SUBHEADINGS = [
    "Starter (5 Minutes)",
    "Teaching and Development (35 Minutes)",
    "Activity and Application (20 Minutes)",
    "Plenary and Closure (5 Minutes)",
]

LESSON_PLAN_SUBHEADING_DISPLAY = {
    "starter": ("Lesson Flow", "Starter (5 Minutes)"),
    "teaching / development": ("Lesson Flow", "Teaching and Development (35 Minutes)"),
    "teaching and development": ("Lesson Flow", "Teaching and Development (35 Minutes)"),
    "development": ("Lesson Flow", "Teaching and Development (35 Minutes)"),
    "activity / application": ("Lesson Flow", "Activity and Application (20 Minutes)"),
    "activity and application": ("Lesson Flow", "Activity and Application (20 Minutes)"),
    "application": ("Lesson Flow", "Activity and Application (20 Minutes)"),
    "plenary / closure": ("Lesson Flow", "Plenary and Closure (5 Minutes)"),
    "plenary and closure": ("Lesson Flow", "Plenary and Closure (5 Minutes)"),
    "closure": ("Lesson Flow", "Plenary and Closure (5 Minutes)"),
}

WORKSHEET_SECTION_DISPLAY = {
    "Instructions": "Instructions",
    "Tasks": "Tasks",
    "Extension Challenge": "Extension Challenge",
    "Answer Key": "Answer Key",
}


def _strip_export_preamble(content: str) -> str:
    text = (content or "").replace("\r\n", "\n").strip()
    if "\n---\n" in text:
        return text.split("\n---\n", 1)[1].strip()
    lines = text.split("\n")
    filtered: list[str] = []
    skipping = True
    for raw in lines:
        line = raw.strip()
        if skipping and (
            not line
            or line.startswith("ELume")
            or line.startswith("Type:")
            or line.startswith("Scope:")
            or line.startswith("Level:")
            or line.startswith("Detail:")
            or line.startswith("Save to:")
            or line.startswith("Created:")
        ):
            continue
        skipping = False
        filtered.append(raw)
    return "\n".join(filtered).strip()


def _clean_heading_text(line: str) -> str:
    text = (line or "").strip()
    text = re.sub(r"^#{1,6}\s*", "", text)
    text = re.sub(r"^\d+\.\s*", "", text)
    return text.strip(" :")


def _clean_lesson_plan_text(value: str) -> str:
    text = (value or "").replace("\r", " ").replace("\t", " ").strip()
    text = text.replace("â€¢", " ").replace("•", " ")
    text = re.sub(r"^\s*[-*?]+\s*", "", text)
    text = re.sub(r"^\s*\d+[\.\)]\s*", "", text)
    text = re.sub(r"^\s*[a-zA-Z][\.\)]\s*", "", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip(" :-")


def _clean_lesson_plan_title(value: str) -> str:
    text = _clean_lesson_plan_text(_clean_heading_text(value or ""))
    text = re.sub(r"^(lesson\s+plan)\s*[:\-|]\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^(lesson\s+plan)\s+[·•]\s*", "", text, flags=re.IGNORECASE)
    return text or "Lesson Plan"


def _normalise_lesson_plan_key(value: str) -> str:
    text = _clean_lesson_plan_text(_clean_heading_text(value or "")).lower()
    text = text.replace("&", "and")
    text = text.replace(" / ", "/")
    text = re.sub(r"\s*/\s*", " / ", text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _is_lesson_plan_bullet(raw: str) -> bool:
    stripped = (raw or "").lstrip()
    return bool(
        stripped.startswith(("- ", "* ", "• ", "? "))
        or re.match(r"^\d+[\.\)]\s+", stripped)
        or re.match(r"^[a-zA-Z][\.\)]\s+", stripped)
    )


def _lesson_plan_scope_parts(scope_label: str) -> tuple[str, str]:
    parts = [part.strip() for part in re.split(r"\s*[•|]\s*", scope_label or "") if part.strip()]
    if len(parts) >= 2:
        return parts[0], parts[1]
    label = (scope_label or "").strip()
    return label, ""


def _lesson_plan_duration_from_text(value: str) -> str:
    text = _clean_lesson_plan_text(value)
    if not text:
        return ""
    if re.search(r"\b(min|mins|minute|minutes|hour|hours|class|classes|lesson|lessons)\b", text, re.IGNORECASE):
        return text
    return ""


def _normalise_section_name(name: str) -> str:
    key = _normalise_lesson_plan_key(name)
    mapping = {
        "lesson title, class, duration, topic": "__meta__",
        "lesson title class duration topic": "__meta__",
        "lesson title": "__meta__",
        "learning overview": "Learning Overview",
        "overview": "Learning Overview",
        "topic": "Learning Overview",
        "prior knowledge": "Learning Overview",
        "learning intentions": "Learning Intentions",
        "learning intention": "Learning Intentions",
        "learning outcomes": "Learning Intentions",
        "learning outcome": "Learning Intentions",
        "success criteria": "Success Criteria",
        "success criterion": "Success Criteria",
        "lesson flow": "Lesson Flow",
        "activity and application": "Lesson Flow",
        "activity / application": "Lesson Flow",
        "plenary and closure": "Lesson Flow",
        "plenary / closure": "Lesson Flow",
        "resources": "Resources",
        "differentiation": "Differentiation",
        "assessment": "Assessment",
        "homework": "Suggested Homework",
        "suggested homework": "Suggested Homework",
        "reflection": "Reflection",
        "footer metadata": "__ignore__",
        "manual notes supplied": "__ignore__",
    }
    return mapping.get(key, "")


def _normalise_lesson_plan_subheading(name: str) -> tuple[str, str] | None:
    return LESSON_PLAN_SUBHEADING_DISPLAY.get(_normalise_lesson_plan_key(name))


def _append_lesson_plan_item(
    sections: dict[str, list[tuple[str, str]]],
    section_name: str,
    item_kind: str,
    item_text: str,
) -> None:
    clean_text = _clean_lesson_plan_text(item_text)
    if not clean_text or section_name not in sections:
        return
    target = sections[section_name]
    if target and target[-1] == (item_kind, clean_text):
        return
    if item_kind == "subheading" and any(kind == "subheading" and text == clean_text for kind, text in target):
        return
    target.append((item_kind, clean_text))


def _lesson_flow_items_for_subheading(items: list[tuple[str, str]], subheading: str) -> list[tuple[str, str]]:
    active = False
    collected: list[tuple[str, str]] = []
    for kind, text in items:
        if kind == "subheading":
            if text == subheading:
                active = True
                continue
            if active:
                break
            continue
        if active:
            collected.append((kind, text))
    return collected


def _lesson_flow_has_subheading(items: list[tuple[str, str]], subheading: str) -> bool:
    return any(kind == "subheading" and text == subheading for kind, text in items)


def _lesson_flow_has_content(items: list[tuple[str, str]], subheading: str) -> bool:
    return bool(_lesson_flow_items_for_subheading(items, subheading))


def _ensure_lesson_flow_subheading(
    sections: dict[str, list[tuple[str, str]]],
    subheading: str,
    fallback_lines: list[str],
) -> None:
    flow_items = sections["Lesson Flow"]
    if _lesson_flow_has_content(flow_items, subheading):
        return
    if not _lesson_flow_has_subheading(flow_items, subheading):
        _append_lesson_plan_item(sections, "Lesson Flow", "subheading", subheading)
    for line in fallback_lines:
        _append_lesson_plan_item(sections, "Lesson Flow", "bullet", line)


def _extract_homework_text(text: str) -> str | None:
    clean = _clean_lesson_plan_text(text)
    match = re.match(r"^(suggested homework|homework)\s*:\s*(.+)$", clean, flags=re.IGNORECASE)
    if match:
        return _clean_lesson_plan_text(match.group(2))
    return None


def _lesson_plan_meta_line(meta: dict[str, str]) -> str:
    parts = []
    if meta.get("Subject"):
        parts.append(meta["Subject"])
    if meta.get("Level"):
        parts.append(meta["Level"])
    if meta.get("Duration"):
        parts.append(meta["Duration"])
    return " | ".join(parts)


def _lesson_plan_footer_label(scope_label: str) -> str:
    class_label, subject_label = _lesson_plan_scope_parts(scope_label)
    parts = [part for part in [class_label, subject_label] if part]
    return " | ".join(parts)


def _decode_data_url_bytes(data_url: str) -> bytes | None:
    raw = (data_url or "").strip()
    if not raw.startswith("data:") or "," not in raw:
        return None
    header, payload = raw.split(",", 1)
    try:
        if ";base64" in header:
            return base64.b64decode(payload)
    except Exception:
        return None
    return None


def _paeth_predictor(a: int, b: int, c: int) -> int:
    p = a + b - c
    pa = abs(p - a)
    pb = abs(p - b)
    pc = abs(p - c)
    if pa <= pb and pa <= pc:
        return a
    if pb <= pc:
        return b
    return c


def _load_footer_logo_bytes(branding_choice: str, meta: dict | None = None) -> bytes | None:
    meta = meta or {}
    if branding_choice == "school":
        return _decode_data_url_bytes(str(meta.get("schoolLogoDataUrl") or ""))
    if branding_choice != "elume":
        return None

    current_logo_path = BASE_DIR.parent / "frontend" / "src" / "assets" / "ELogo2.png"
    try:
        if current_logo_path.exists():
            return current_logo_path.read_bytes()
    except Exception:
        return None
    return None


def _prepare_pdf_png_image(data: bytes | None) -> dict[str, Any] | None:
    if not data:
        return None
    signature = b"\x89PNG\r\n\x1a\n"
    if not data.startswith(signature):
        return None

    offset = len(signature)
    width = height = bit_depth = color_type = None
    idat_parts: list[bytes] = []

    while offset + 8 <= len(data):
        length = struct.unpack(">I", data[offset:offset + 4])[0]
        chunk_type = data[offset + 4:offset + 8]
        chunk_start = offset + 8
        chunk_end = chunk_start + length
        chunk_data = data[chunk_start:chunk_end]
        offset = chunk_end + 4

        if chunk_type == b"IHDR" and len(chunk_data) >= 13:
            width, height, bit_depth, color_type = struct.unpack(">IIBB", chunk_data[:10])
        elif chunk_type == b"IDAT":
            idat_parts.append(chunk_data)
        elif chunk_type == b"IEND":
            break

    if not width or not height or bit_depth != 8 or color_type not in {2, 6} or not idat_parts:
        return None

    try:
        decompressed = zlib.decompress(b"".join(idat_parts))
    except Exception:
        return None

    bytes_per_pixel = 3 if color_type == 2 else 4
    stride = width * bytes_per_pixel
    expected = height * (stride + 1)
    if len(decompressed) < expected:
        return None

    rgb = bytearray()
    alpha = bytearray() if color_type == 6 else None
    prev_row = bytearray(stride)
    pos = 0

    for _ in range(height):
        filter_type = decompressed[pos]
        pos += 1
        row = bytearray(decompressed[pos:pos + stride])
        pos += stride
        recon = bytearray(stride)

        for i in range(stride):
            left = recon[i - bytes_per_pixel] if i >= bytes_per_pixel else 0
            up = prev_row[i]
            up_left = prev_row[i - bytes_per_pixel] if i >= bytes_per_pixel else 0
            value = row[i]
            if filter_type == 0:
                recon[i] = value
            elif filter_type == 1:
                recon[i] = (value + left) & 255
            elif filter_type == 2:
                recon[i] = (value + up) & 255
            elif filter_type == 3:
                recon[i] = (value + ((left + up) // 2)) & 255
            elif filter_type == 4:
                recon[i] = (value + _paeth_predictor(left, up, up_left)) & 255
            else:
                return None

        prev_row = recon
        if color_type == 2:
            rgb.extend(recon)
        else:
            for i in range(0, stride, 4):
                rgb.extend(recon[i:i + 3])
                alpha.append(recon[i + 3])

    return {
        "width": width,
        "height": height,
        "rgb": zlib.compress(bytes(rgb)),
        "alpha": zlib.compress(bytes(alpha)) if alpha is not None else None,
    }


def _is_lesson_plan_pdf(title: str, meta: dict | None = None) -> bool:
    meta = meta or {}
    output_kind = str(meta.get("outputKind") or meta.get("kind") or "").strip().lower()
    if output_kind == "lesson_plan":
        return True
    return "lesson plan" in (title or "").strip().lower()


def _is_worksheet_pdf(title: str, meta: dict | None = None) -> bool:
    meta = meta or {}
    output_kind = str(meta.get("outputKind") or meta.get("kind") or "").strip().lower()
    if output_kind == "worksheet":
        return True
    return "worksheet" in (title or "").strip().lower()


def _normalise_worksheet_heading(name: str) -> str:
    key = _clean_lesson_plan_text(_clean_heading_text(name or "")).lower()
    mapping = {
        "instructions": "Instructions",
        "short instructions": "Instructions",
        "task 1": "Tasks",
        "task 2": "Tasks",
        "task 3": "Tasks",
        "task 4": "Tasks",
        "task 5": "Tasks",
        "tasks": "Tasks",
        "questions": "Tasks",
        "question 1": "Tasks",
        "question 2": "Tasks",
        "question 3": "Tasks",
        "question 4": "Tasks",
        "question 5": "Tasks",
        "extension challenge": "Extension Challenge",
        "extension": "Extension Challenge",
        "challenge": "Extension Challenge",
        "reflection": "Extension Challenge",
        "answer key": "Answer Key",
        "answers": "Answer Key",
    }
    return mapping.get(key, "")


def _clean_worksheet_line(value: str) -> str:
    return _clean_lesson_plan_text(_clean_heading_text(value or ""))


def _worksheet_display_line(value: str) -> str:
    text = _clean_worksheet_line(value)
    lowered = text.lower()
    if lowered.startswith("instructions:"):
        text = text.split(":", 1)[1].strip()
    if lowered.startswith("short instructions:"):
        text = text.split(":", 1)[1].strip()
    return text


def _parse_worksheet_pdf_content(title: str, content: str, meta: dict | None = None) -> dict[str, Any]:
    body = _strip_export_preamble(content)
    meta = meta or {}
    worksheet_title = (title or "").strip() or "Worksheet"
    instructions: list[str] = []
    tasks: list[dict[str, Any]] = []
    extension: list[str] = []
    answers: list[str] = []
    class_line = ""
    found_structure = False

    current_section = ""
    current_task: dict[str, Any] | None = None

    def flush_task() -> None:
        nonlocal current_task
        if current_task and (current_task.get("title") or current_task.get("lines")):
            tasks.append(current_task)
        current_task = None

    for raw in body.split("\n"):
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped or stripped == "---":
            continue

        if stripped.lower().startswith("worksheet:"):
            candidate = _clean_lesson_plan_title(stripped)
            if candidate:
                worksheet_title = candidate
                found_structure = True
            continue

        if stripped.lower().startswith("worksheet title:"):
            candidate = _clean_lesson_plan_text(stripped.split(":", 1)[1])
            if candidate:
                worksheet_title = candidate
                found_structure = True
            continue

        if stripped.lower().startswith("class:"):
            class_line = _clean_lesson_plan_text(stripped.split(":", 1)[1])
            if class_line:
                found_structure = True
            continue

        if stripped.lower().startswith("student name:") or stripped.lower().startswith("date:"):
            found_structure = True
            continue

        heading = _normalise_worksheet_heading(stripped)
        if heading:
            if heading != "Tasks":
                flush_task()
            current_section = heading
            found_structure = True
            continue

        clean = _worksheet_display_line(stripped)
        if not clean:
            continue

        if current_section == "Instructions":
            instructions.append(clean)
            found_structure = True
            continue

        if current_section == "Tasks":
            task_heading = _clean_heading_text(stripped)
            if re.match(r"^(task|question)\s*\d+(\s*[:\-]\s*.+)?$", task_heading, flags=re.IGNORECASE):
                flush_task()
                title_match = re.match(r"^((?:task|question)\s*\d+)\s*[:\-]?\s*(.*)$", task_heading, flags=re.IGNORECASE)
                if title_match:
                    current_task = {"title": title_match.group(1).title(), "lines": []}
                    trailing = _clean_worksheet_line(title_match.group(2))
                    if trailing:
                        current_task["lines"].append(trailing)
                else:
                    current_task = {"title": task_heading, "lines": []}
                found_structure = True
                continue
            numbered = re.match(r"^(\d+)[\.\)]\s*(.+)$", stripped)
            if numbered:
                flush_task()
                current_task = {"title": f"Task {numbered.group(1)}", "lines": [_clean_lesson_plan_text(numbered.group(2))]}
                found_structure = True
                continue
            if stripped.startswith("#"):
                continue
            if current_task is None:
                current_task = {"title": f"Task {len(tasks) + 1}", "lines": []}
            current_task["lines"].append(clean)
            found_structure = True
            continue

        if current_section == "Extension Challenge":
            extension.append(clean)
            found_structure = True
            continue

        if current_section == "Answer Key":
            answers.append(clean)
            found_structure = True
            continue

        task_heading = _clean_heading_text(stripped)
        if re.match(r"^(task|question)\s*\d+(\s*[:\-]\s*.+)?$", task_heading, flags=re.IGNORECASE):
            flush_task()
            title_match = re.match(r"^((?:task|question)\s*\d+)\s*[:\-]?\s*(.*)$", task_heading, flags=re.IGNORECASE)
            if title_match:
                current_task = {"title": title_match.group(1).title(), "lines": []}
                trailing = _clean_worksheet_line(title_match.group(2))
                if trailing:
                    current_task["lines"].append(trailing)
            else:
                current_task = {"title": task_heading, "lines": []}
            current_section = "Tasks"
            found_structure = True
            continue

        numbered = re.match(r"^(\d+)[\.\)]\s*(.+)$", stripped)
        if numbered:
            flush_task()
            current_task = {"title": f"Task {numbered.group(1)}", "lines": [_clean_worksheet_line(numbered.group(2))]}
            current_section = "Tasks"
            found_structure = True
            continue

        if not current_section and clean:
            if not instructions:
                instructions.append(clean)
                found_structure = True
                continue
            if current_task is None:
                current_task = {"title": f"Task {len(tasks) + 1}", "lines": []}
            current_section = "Tasks"
            current_task["lines"].append(clean)
            found_structure = True
            continue

    flush_task()

    if not instructions:
        instructions = [
            "Read each task carefully and answer in clear subject-specific language.",
            "Show your working or explanation where appropriate.",
        ]

    if not tasks and not found_structure:
        tasks = [
            {"title": "Task 1", "lines": ["Complete the first question using the topic and key vocabulary provided."]},
            {"title": "Task 2", "lines": ["Apply what you know in a short written response or worked example."]},
        ]

    for idx, task in enumerate(tasks, start=1):
        task["title"] = task.get("title") or f"Task {idx}"
        task["lines"] = [line for line in task.get("lines", []) if line]

    include_answers = bool(meta.get("worksheetIncludeAnswers"))
    if not include_answers:
        answers = []

    return {
        "title": worksheet_title,
        "classLine": class_line or str(meta.get("scopeLabel") or "").strip(),
        "instructions": instructions,
        "tasks": tasks,
        "extension": extension,
        "answers": answers,
    }


def _parse_lesson_plan_pdf_content(title: str, content: str, meta: dict | None = None) -> dict:
    body = _strip_export_preamble(content)
    meta = meta or {}
    sections: dict[str, list[tuple[str, str]]] = {name: [] for name in LESSON_PLAN_SECTION_ORDER}
    top_meta = {
        "Lesson Title": _clean_lesson_plan_title(title or ""),
        "Class": "",
        "Subject": "",
        "Level": "",
        "Duration": "",
        "Topic": "",
    }

    current_section = "Learning Overview"
    section_zero = False
    seen_structured_section = False

    for raw in body.split("\n"):
        line = raw.rstrip()
        stripped = line.strip()
        if not stripped or stripped == "---":
            continue

        if stripped.lower().startswith("lesson plan:"):
            candidate_title = _clean_lesson_plan_title(stripped)
            if candidate_title:
                top_meta["Lesson Title"] = candidate_title
            continue

        heading = _clean_heading_text(stripped)
        normalised = _normalise_section_name(heading)
        if normalised == "__meta__":
            section_zero = True
            current_section = ""
            continue
        if normalised == "__ignore__":
            break
        if normalised:
            current_section = normalised
            section_zero = False
            seen_structured_section = True
            if heading in {"Activity and Application", "Plenary and Closure"}:
                mapped = _normalise_lesson_plan_subheading(heading)
                if mapped:
                    current_section = mapped[0]
                    _append_lesson_plan_item(sections, current_section, "subheading", mapped[1])
            continue

        subheading = _normalise_lesson_plan_subheading(heading)
        if subheading and (stripped.startswith("### ") or current_section == "Lesson Flow"):
            current_section = subheading[0]
            section_zero = False
            seen_structured_section = True
            _append_lesson_plan_item(sections, current_section, "subheading", subheading[1])
            continue

        if not seen_structured_section and "|" in stripped and not stripped.startswith("#"):
            meta_parts = [part.strip() for part in stripped.split("|") if part.strip()]
            if len(meta_parts) >= 3:
                if not top_meta["Subject"]:
                    top_meta["Subject"] = _clean_lesson_plan_text(meta_parts[0])
                if not top_meta["Level"]:
                    top_meta["Level"] = _clean_lesson_plan_text(meta_parts[1])
                if not top_meta["Duration"]:
                    top_meta["Duration"] = _clean_lesson_plan_text(meta_parts[2])
                continue

        if ":" in stripped:
            key, value = stripped.lstrip("- ").split(":", 1)
            key_clean = _normalise_lesson_plan_key(key)
            value_clean = _clean_lesson_plan_text(value)
            if key_clean in {"homework", "suggested homework"}:
                current_section = "Suggested Homework"
                section_zero = False
                if value_clean:
                    _append_lesson_plan_item(sections, current_section, "bullet", value_clean)
                continue
            if key_clean in {"support", "extension"}:
                current_section = "Differentiation"
                section_zero = False
                if value_clean:
                    _append_lesson_plan_item(sections, current_section, "bullet", f"{key.strip().title()}: {value_clean}")
                continue
            if key_clean == "reflection":
                current_section = "Reflection"
                section_zero = False
                if value_clean:
                    _append_lesson_plan_item(sections, current_section, "bullet", value_clean)
                continue
            if key_clean in {"lesson title", "class", "subject", "duration", "topic", "level"} and value_clean:
                target_key = {
                    "lesson title": "Lesson Title",
                    "class": "Class",
                    "subject": "Subject",
                    "duration": "Duration",
                    "topic": "Topic",
                    "level": "Level",
                }[key_clean]
                top_meta[target_key] = value_clean
                continue

        if section_zero:
            continue

        clean_text = _clean_lesson_plan_text(stripped)
        if not clean_text:
            continue

        if clean_text.lower() in {
            top_meta["Lesson Title"].lower(),
            f"lesson plan: {top_meta['Lesson Title'].lower()}",
            f"topic: {top_meta['Topic'].lower()}" if top_meta["Topic"] else "",
        }:
            continue

        target_section = current_section or "Learning Overview"
        item_kind = "bullet" if _is_lesson_plan_bullet(stripped) else "p"
        _append_lesson_plan_item(sections, target_section, item_kind, clean_text)

    class_from_scope, subject_from_scope = _lesson_plan_scope_parts(str(meta.get("scopeLabel") or ""))
    if not top_meta["Class"] and class_from_scope:
        top_meta["Class"] = class_from_scope
    if not top_meta["Subject"] and subject_from_scope:
        top_meta["Subject"] = subject_from_scope
    if not top_meta["Level"]:
        top_meta["Level"] = _clean_lesson_plan_text(str(meta.get("level") or ""))
    if not top_meta["Duration"]:
        top_meta["Duration"] = _lesson_plan_duration_from_text(str(meta.get("detail") or "")) or "60 Minutes"
    if not top_meta["Topic"]:
        top_meta["Topic"] = top_meta["Lesson Title"]

    overview_items: list[tuple[str, str]] = []
    if top_meta["Topic"]:
        overview_items.append(("bullet", f"Topic: {top_meta['Topic']}"))
    for kind, text in sections["Learning Overview"]:
        if text.lower() != f"topic: {top_meta['Topic']}".lower():
            overview_items.append((kind, text))
    sections["Learning Overview"] = overview_items

    moved_homework: list[tuple[str, str]] = []
    for section_name in list(sections.keys()):
        if section_name == "Suggested Homework":
            continue
        cleaned_items: list[tuple[str, str]] = []
        for item_kind, item_text in sections[section_name]:
            extracted_homework = _extract_homework_text(item_text)
            if extracted_homework:
                moved_homework.append(("bullet", extracted_homework))
                continue
            cleaned_items.append((item_kind, item_text))
        sections[section_name] = cleaned_items
    for item_kind, item_text in moved_homework:
        _append_lesson_plan_item(sections, "Suggested Homework", item_kind, item_text)

    _ensure_lesson_flow_subheading(
        sections,
        "Starter (5 Minutes)",
        ["Begin with a short retrieval or settling task linked to recent learning."],
    )
    _ensure_lesson_flow_subheading(
        sections,
        "Teaching and Development (35 Minutes)",
        ["Explain and model the new learning using clear examples, questioning, and guided teacher input."],
    )
    _ensure_lesson_flow_subheading(
        sections,
        "Activity and Application (20 Minutes)",
        ["Students complete a focused task that applies the learning with teacher circulation and feedback."],
    )
    _ensure_lesson_flow_subheading(
        sections,
        "Plenary and Closure (5 Minutes)",
        ["Close with a brief review, exit prompt, or check for understanding."],
    )

    if not sections["Reflection"]:
        sections["Reflection"] = []

    ordered_sections = [
        (name, sections[name])
        for name in LESSON_PLAN_SECTION_ORDER
        if sections[name] or name == "Reflection"
    ]
    return {"meta": top_meta, "sections": ordered_sections}


def _docx_from_markdownish(title: str, content: str, teacher: str | None = None, meta: dict | None = None) -> bytes:
    """
    Very simple parser:
    - Lines starting with ### / ## / # become headings
    - Lines starting with "- " become bullet points
    - Blank lines separate paragraphs
    """
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"python-docx not installed/available: {e}")

    doc = Document()

    # Title
    doc.add_heading(title.strip() or "ELume Resource", level=0)

    # Subheader
    sub = []
    if teacher:
        sub.append(f"Teacher: {teacher}")
    if meta:
        # keep it neat: only show a few common fields if present
        for k in ["template", "level", "tone", "scopeLabel", "createdAt"]:
            if meta.get(k):
                label = {
                    "scopeLabel": "Scope",
                    "createdAt": "Created",
                }.get(k, k.capitalize())
                sub.append(f"{label}: {meta.get(k)}")
    if sub:
        p = doc.add_paragraph(" • ".join(sub))
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_paragraph("")  # spacer

    lines = (content or "").replace("\r\n", "\n").split("\n")
    in_bullets = False

    for raw in lines:
        line = raw.rstrip()

        if not line.strip():
            doc.add_paragraph("")
            in_bullets = False
            continue

        # headings
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            in_bullets = False
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            in_bullets = False
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
            in_bullets = False
            continue

        # bullets
        if line.lstrip().startswith("- "):
            text = line.lstrip()[2:].strip()
            doc.add_paragraph(text, style="List Bullet")
            in_bullets = True
            continue

        # normal paragraph
        doc.add_paragraph(line)
        in_bullets = False

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _pdf_from_markdownish(title: str, content: str, teacher: str | None = None, meta: dict | None = None) -> bytes:
    page_width = 595
    page_height = 842
    left = 56
    right = 56
    top = 54
    bottom = 58

    title = (title or "").strip() or "ELume Resource"
    content = (content or "").replace("\r\n", "\n")
    meta = meta or {}

    teacher_short = str(meta.get("teacherDisplayNameShort") or "").strip() or teacher or ""
    school_name = str(meta.get("schoolName") or "").strip()
    branding_choice = str(meta.get("brandingChoice") or "").strip().lower()
    scope_label = str(meta.get("scopeLabel") or "").strip()
    level = str(meta.get("level") or "").strip()
    detail = str(meta.get("detail") or "").strip()
    created_at = str(meta.get("createdAt") or "").strip()
    output_kind = str(meta.get("outputKind") or meta.get("kind") or "").strip()
    school_logo_available = bool(meta.get("schoolLogoAvailable"))
    lesson_plan_mode = _is_lesson_plan_pdf(title, meta)
    worksheet_mode = _is_worksheet_pdf(title, meta)
    footer_logo = _prepare_pdf_png_image(_load_footer_logo_bytes(branding_choice, meta))
    footer_reserved_height = 74 if lesson_plan_mode or worksheet_mode else bottom
    usable_width = page_width - left - right

    pages: list[str] = []
    ops: list[str] = []
    page_number = 0
    y = page_height - top

    def add_text_line(
        text_value: str,
        x: float,
        y_value: float,
        font_size: int,
        font_name: str = "F1",
        color: tuple[float, float, float] = (0.12, 0.16, 0.22),
    ) -> None:
        safe = _pdf_escape(text_value)
        r, g, b = color
        ops.append("BT")
        ops.append(f"{r:.3f} {g:.3f} {b:.3f} rg")
        ops.append(f"/{font_name} {font_size} Tf")
        ops.append(f"1 0 0 1 {x:.2f} {y_value:.2f} Tm")
        ops.append(f"({safe}) Tj")
        ops.append("ET")

    def add_rule(y_value: float, color: tuple[float, float, float] = (0.84, 0.88, 0.92), width: float = 1.0) -> None:
        r, g, b = color
        ops.append(f"{r:.3f} {g:.3f} {b:.3f} RG")
        ops.append(f"{width:.2f} w")
        ops.append(f"{left} {y_value:.2f} m {page_width - right} {y_value:.2f} l S")

    def add_rule_segment(x1: float, x2: float, y_value: float, color: tuple[float, float, float] = (0.84, 0.88, 0.92), width: float = 1.0) -> None:
        r, g, b = color
        ops.append(f"{r:.3f} {g:.3f} {b:.3f} RG")
        ops.append(f"{width:.2f} w")
        ops.append(f"{x1:.2f} {y_value:.2f} m {x2:.2f} {y_value:.2f} l S")

    def add_light_box(
        x: float,
        y_value: float,
        width: float,
        height: float,
        fill: tuple[float, float, float] = (0.97, 0.98, 0.99),
        stroke: tuple[float, float, float] = (0.86, 0.90, 0.94),
    ) -> None:
        fr, fg, fb = fill
        sr, sg, sb = stroke
        ops.append(f"{fr:.3f} {fg:.3f} {fb:.3f} rg")
        ops.append(f"{sr:.3f} {sg:.3f} {sb:.3f} RG")
        ops.append(f"{x:.2f} {y_value:.2f} {width:.2f} {height:.2f} re B")

    def fit_logo_to_box(
        image: dict[str, Any] | None,
        max_width: float,
        max_height: float,
    ) -> tuple[float, float]:
        if not image:
            return 0.0, 0.0
        source_width = float(image["width"])
        source_height = float(image["height"])
        if source_width <= 0 or source_height <= 0:
            return 0.0, 0.0
        scale = min(max_width / source_width, max_height / source_height, 1.0)
        return source_width * scale, source_height * scale

    def trim_footer_text(text_value: str, max_chars: int) -> str:
        clean = " ".join((text_value or "").split())
        if len(clean) <= max_chars:
            return clean
        return clean[: max(0, max_chars - 1)].rstrip() + "…"

    def draw_footer_logo(y_value: float, logo_right_x: float, max_width: float, max_height: float) -> None:
        if not footer_logo:
            return
        logo_width, logo_height = fit_logo_to_box(footer_logo, max_width=max_width, max_height=max_height)
        if logo_width <= 0 or logo_height <= 0:
            return
        logo_x = logo_right_x - logo_width
        logo_y = y_value - (logo_height / 2.0)
        ops.append("q")
        ops.append(f"{logo_width:.2f} 0 0 {logo_height:.2f} {logo_x:.2f} {logo_y:.2f} cm")
        ops.append("/Im1 Do")
        ops.append("Q")

    def add_bullet_line(
        text_value: str,
        indent: float = 12,
        width_chars: int = 68,
        font_size: int = 11,
        line_height: float = 14,
        color: tuple[float, float, float] = (0.14, 0.18, 0.24),
        extra_after: float = 2.0,
    ) -> None:
        nonlocal y
        wrapped = _pdf_wrap(text_value, width_chars)
        ensure_room(line_height * len(wrapped) + extra_after)
        first = True
        for line in wrapped:
            prefix = f"{chr(183)} " if first else "  "
            add_text_line(f"{prefix}{line}", left + indent, y, font_size, color=color)
            y -= line_height
            first = False
        y -= extra_after

    def add_writing_lines(count: int = 3, gap: float = 18.0) -> None:
        nonlocal y
        ensure_room((count * gap) + 4)
        for _ in range(count):
            add_rule(y - 8, color=(0.88, 0.91, 0.95), width=0.8)
            y -= gap

    def add_header(current_page: int) -> None:
        nonlocal y
        if lesson_plan_mode:
            y = page_height - top
            return
        if worksheet_mode:
            y = page_height - top
            return
        add_text_line(title, left, page_height - 34, 10, font_name="F2", color=(0.20, 0.25, 0.33))
        if school_name:
            add_text_line(school_name, page_width - right - 140, page_height - 34, 9, color=(0.38, 0.45, 0.55))
        add_rule(page_height - 42, color=(0.84, 0.89, 0.94), width=0.8)
        y = page_height - top

    def finish_page() -> None:
        nonlocal ops, y
        footer_y = 34
        rule_y = footer_y + (18 if lesson_plan_mode else 12)
        add_rule(rule_y, color=(0.88, 0.91, 0.95), width=0.8)
        if lesson_plan_mode:
            footer_y = 28
            footer_left = " | ".join([part for part in [teacher_short, school_name or _lesson_plan_footer_label(scope_label)] if part])
            footer_left = trim_footer_text(footer_left, 42)
            logo_area_width = 96.0 if footer_logo else 0.0
            page_x = (page_width / 2) - 10
            if footer_left:
                add_text_line(footer_left, left, footer_y, 9, color=(0.42, 0.48, 0.56))
            add_text_line(f"Page {page_number}", page_x, footer_y, 9, color=(0.42, 0.48, 0.56))
            if footer_logo:
                draw_footer_logo(footer_y + 2, page_width - right, max_width=logo_area_width, max_height=20.0)
        elif worksheet_mode:
            footer_y = 28
            footer_left = " | ".join([part for part in [teacher_short, school_name] if part])
            footer_left = trim_footer_text(footer_left, 42)
            logo_area_width = 88.0 if footer_logo else 0.0
            if footer_left:
                add_text_line(footer_left, left, footer_y, 9, color=(0.42, 0.48, 0.56))
            add_text_line(f"Page {page_number}", (page_width / 2) - 10, footer_y, 9, color=(0.42, 0.48, 0.56))
            if footer_logo:
                draw_footer_logo(footer_y + 2, page_width - right, max_width=logo_area_width, max_height=18.0)
        else:
            footer_left = " | ".join([part for part in [teacher_short, school_name] if part])
            footer_left = trim_footer_text(footer_left, 48)
            if footer_left:
                add_text_line(footer_left, left, footer_y, 9, color=(0.42, 0.48, 0.56))
            add_text_line(f"Page {page_number}", (page_width / 2) - 14, footer_y, 9, color=(0.42, 0.48, 0.56))
            if footer_logo:
                draw_footer_logo(footer_y, page_width - right, max_width=58.0, max_height=15.0)
        pages.append("\n".join(ops))
        ops = []
        y = page_height - top

    def start_page() -> None:
        nonlocal page_number
        page_number += 1
        add_header(page_number)

    def ensure_room(height_needed: float) -> None:
        nonlocal y
        if y - height_needed < footer_reserved_height + 22:
            finish_page()
            start_page()

    def draw_wrapped(
        text_value: str,
        x: float,
        font_size: int,
        width_chars: int,
        line_height: float,
        font_name: str = "F1",
        color: tuple[float, float, float] = (0.12, 0.16, 0.22),
        extra_after: float = 0.0,
    ) -> None:
        nonlocal y
        wrapped = _pdf_wrap(text_value, width_chars)
        ensure_room(line_height * len(wrapped) + extra_after)
        for line in wrapped:
            add_text_line(line, x, y, font_size, font_name=font_name, color=color)
            y -= line_height
        y -= extra_after

    start_page()

    if lesson_plan_mode:
        parsed = _parse_lesson_plan_pdf_content(title, content, meta=meta)
        lesson_meta = parsed["meta"]
        sections = parsed["sections"]
        display_title = f"Lesson Plan: {lesson_meta.get('Lesson Title') or _clean_lesson_plan_title(title)}"
        meta_line = _lesson_plan_meta_line(lesson_meta)

        draw_wrapped(display_title, left, 22, 44, 24, font_name="F2", color=(0.08, 0.14, 0.22), extra_after=3)
        if meta_line:
            draw_wrapped(meta_line, left, 10, 94, 13, color=(0.38, 0.45, 0.55), extra_after=5)
        add_rule(y, color=(0.83, 0.89, 0.95), width=1.0)
        y -= 16

        for section_name, items in sections:
            ensure_room(34)
            draw_wrapped(
                LESSON_PLAN_SECTION_DISPLAY.get(section_name, section_name),
                left,
                15,
                66,
                18,
                font_name="F2",
                color=(0.10, 0.30, 0.37),
                extra_after=3,
            )
            if section_name == "Reflection":
                for item_kind, item_text in items:
                    if item_kind == "bullet":
                        add_bullet_line(item_text, indent=10, width_chars=90, font_size=11, line_height=14, extra_after=2)
                    elif item_kind == "p":
                        draw_wrapped(item_text, left, 11, 96, 14, color=(0.14, 0.18, 0.24), extra_after=3)
                add_writing_lines(4, gap=16.0)
                y -= 8
                continue
            for item_kind, item_text in items:
                if item_kind == "subheading":
                    draw_wrapped(item_text, left, 12, 82, 15, font_name="F2", color=(0.17, 0.23, 0.31), extra_after=1)
                    continue
                if item_kind == "bullet":
                    add_bullet_line(item_text, indent=10, width_chars=90, font_size=11, line_height=14, extra_after=2)
                    continue
                draw_wrapped(item_text, left, 11, 96, 14, color=(0.14, 0.18, 0.24), extra_after=3)
            y -= 10
    elif worksheet_mode:
        parsed = _parse_worksheet_pdf_content(title, content, meta=meta)
        display_title = parsed["title"]
        meta_parts = [part for part in [level] if part]
        meta_line = " | ".join(meta_parts)

        def draw_worksheet_lines(
            lines: list[str],
            *,
            x: float,
            text_width: int = 72,
            font_size: int = 11,
            line_height: float = 14,
            extra_after: float = 2,
        ) -> None:
            for item in lines:
                clean_item = _worksheet_display_line(item)
                if not clean_item:
                    continue
                numbered_match = re.match(r"^(\d+[\.\)])\s+(.+)$", clean_item)
                bullet_match = clean_item.startswith("- ") or clean_item.startswith("* ")
                if numbered_match:
                    draw_wrapped(f"{numbered_match.group(1)} {numbered_match.group(2)}", x, font_size, text_width, line_height, color=(0.14, 0.18, 0.24), extra_after=extra_after)
                    continue
                if bullet_match:
                    add_bullet_line(clean_item[2:].strip(), indent=max(int(x - left), 10), width_chars=text_width - 2, font_size=font_size, line_height=line_height, extra_after=extra_after)
                    continue
                draw_wrapped(clean_item, x, font_size, text_width, line_height, color=(0.14, 0.18, 0.24), extra_after=extra_after)

        draw_wrapped(display_title, left, 22, 34, 24, font_name="F2", color=(0.08, 0.14, 0.22), extra_after=4)
        if meta_line:
            draw_wrapped(meta_line, left, 10, 78, 13, color=(0.38, 0.45, 0.55), extra_after=4)
        add_rule(y, color=(0.83, 0.89, 0.95), width=1.0)
        y -= 18

        ensure_room(64)
        add_text_line("Student Name", left, y, 10, font_name="F2", color=(0.25, 0.31, 0.40))
        add_rule_segment(left + 74, left + 224, y - 4, color=(0.78, 0.83, 0.89), width=0.9)
        add_text_line("Class", left + 244, y, 10, font_name="F2", color=(0.25, 0.31, 0.40))
        add_rule_segment(left + 278, left + 392, y - 4, color=(0.78, 0.83, 0.89), width=0.9)
        add_text_line("Date", left + 408, y, 10, font_name="F2", color=(0.25, 0.31, 0.40))
        add_rule_segment(left + 440, page_width - right, y - 4, color=(0.78, 0.83, 0.89), width=0.9)
        y -= 22
        if parsed["classLine"]:
            draw_wrapped(f"Class group: {parsed['classLine']}", left, 10, 78, 13, color=(0.38, 0.45, 0.55), extra_after=8)

        draw_wrapped("Instructions", left, 15, 54, 18, font_name="F2", color=(0.10, 0.30, 0.37), extra_after=2)
        draw_worksheet_lines(parsed["instructions"], x=left, text_width=76, font_size=11, line_height=14, extra_after=3)
        y -= 8

        for idx, task in enumerate(parsed["tasks"], start=1):
            task_lines = [line for line in task["lines"] if _worksheet_display_line(line)]
            writing_line_count = 4 if len(task_lines) <= 2 else 3
            estimated_height = 92 + (len(task_lines) * 16) + (writing_line_count * 12)
            ensure_room(estimated_height)
            add_light_box(left, y - 12, usable_width, 32)
            draw_wrapped(task["title"] or f"Task {idx}", left + 12, 13, 60, 16, font_name="F2", color=(0.10, 0.24, 0.31), extra_after=1)
            y -= 18
            draw_worksheet_lines(task_lines, x=left + 12, text_width=72, font_size=11, line_height=15, extra_after=3)
            add_writing_lines(writing_line_count)
            y -= 8

        if parsed["extension"]:
            draw_wrapped("Extension Challenge", left, 15, 54, 18, font_name="F2", color=(0.10, 0.30, 0.37), extra_after=3)
            draw_worksheet_lines(parsed["extension"], x=left + 12, text_width=72, font_size=11, line_height=15, extra_after=3)
            add_writing_lines(2)
            y -= 8

        if parsed["answers"]:
            draw_wrapped("Teacher Answer Key", left, 15, 54, 18, font_name="F2", color=(0.10, 0.30, 0.37), extra_after=3)
            draw_worksheet_lines(parsed["answers"], x=left + 12, text_width=72, font_size=11, line_height=15, extra_after=3)
            y -= 6
    else:
        body = _strip_export_preamble(content)
        blocks: list[tuple[str, str]] = [("title", title)]
        sub = [part for part in [scope_label, level, detail, created_at] if part]
        if sub:
            blocks.append(("sub", " • ".join(sub)))
        blocks.append(("spacer", ""))

        for raw in body.split("\n"):
            line = raw.rstrip()
            if not line.strip():
                blocks.append(("spacer", ""))
                continue
            if line == "---":
                blocks.append(("rule", ""))
                continue
            if line.startswith("### "):
                blocks.append(("h3", _clean_heading_text(line)))
                continue
            if line.startswith("## "):
                blocks.append(("h2", _clean_heading_text(line)))
                continue
            if line.startswith("# "):
                blocks.append(("h1", _clean_heading_text(line)))
                continue
            if line.lstrip().startswith("- "):
                blocks.append(("bullet", line.lstrip()[2:].strip()))
                continue
            blocks.append(("p", line))

        for kind, value in blocks:
            if kind == "spacer":
                y -= 10
                continue
            if kind == "rule":
                ensure_room(16)
                add_rule(y, color=(0.86, 0.90, 0.94), width=0.9)
                y -= 12
                continue
            if kind == "title":
                draw_wrapped(value, left, 20, 40, 22, font_name="F2", color=(0.08, 0.14, 0.22), extra_after=5)
                continue
            if kind == "sub":
                draw_wrapped(value, left, 10, 76, 13, color=(0.38, 0.45, 0.55), extra_after=4)
                continue
            if kind == "h1":
                draw_wrapped(value, left, 17, 46, 20, font_name="F2", color=(0.10, 0.32, 0.38), extra_after=2)
                continue
            if kind == "h2":
                draw_wrapped(value, left, 15, 54, 18, font_name="F2", color=(0.10, 0.32, 0.38), extra_after=1)
                continue
            if kind == "h3":
                draw_wrapped(value, left, 12, 64, 15, font_name="F2", color=(0.17, 0.23, 0.31))
                continue
            if kind == "bullet":
                add_bullet_line(value, indent=10, width_chars=70, font_size=11, line_height=14, extra_after=2)
                continue
            draw_wrapped(value, left, 11, 80, 14, color=(0.14, 0.18, 0.24), extra_after=1)

    if ops or not pages:
        finish_page()

    objects: list[bytes] = []

    def add_obj(data: bytes) -> int:
        objects.append(data)
        return len(objects)

    font_regular = add_obj(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    font_bold = add_obj(b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica-Bold >>")
    footer_logo_ref: int | None = None

    if footer_logo:
        alpha_ref: int | None = None
        if footer_logo.get("alpha"):
            alpha_stream = footer_logo["alpha"]
            alpha_ref = add_obj(
                (
                    f"<< /Type /XObject /Subtype /Image /Width {footer_logo['width']} /Height {footer_logo['height']} "
                    f"/ColorSpace /DeviceGray /BitsPerComponent 8 /Filter /FlateDecode /Length {len(alpha_stream)} >>\nstream\n"
                ).encode("ascii")
                + alpha_stream
                + b"\nendstream"
            )

        rgb_stream = footer_logo["rgb"]
        image_dict = (
            f"<< /Type /XObject /Subtype /Image /Width {footer_logo['width']} /Height {footer_logo['height']} "
            f"/ColorSpace /DeviceRGB /BitsPerComponent 8 /Filter /FlateDecode /Length {len(rgb_stream)} "
        )
        if alpha_ref is not None:
            image_dict += f"/SMask {alpha_ref} 0 R "
        image_dict += ">>\nstream\n"
        footer_logo_ref = add_obj(image_dict.encode("ascii") + rgb_stream + b"\nendstream")

    page_ids: list[int] = []

    for page_ops in pages:
        stream = page_ops.encode("latin-1", errors="replace")
        content_obj = add_obj(b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream")
        xobject_part = f"/XObject << /Im1 {footer_logo_ref} 0 R >> " if footer_logo_ref is not None else ""
        page_obj = add_obj(
            (
                f"<< /Type /Page /Parent PAGES_REF 0 R /MediaBox [0 0 {page_width} {page_height}] "
                f"/Resources << /Font << /F1 {font_regular} 0 R /F2 {font_bold} 0 R >> {xobject_part}>> /Contents {content_obj} 0 R >>"
            ).encode("ascii")
        )
        page_ids.append(page_obj)

    kids = " ".join(f"{pid} 0 R" for pid in page_ids)
    pages_obj_num = add_obj(f"<< /Type /Pages /Kids [{kids}] /Count {len(page_ids)} >>".encode("ascii"))

    for pid in page_ids:
        objects[pid - 1] = objects[pid - 1].replace(b"PAGES_REF", str(pages_obj_num).encode("ascii"))

    catalog_obj = add_obj(f"<< /Type /Catalog /Pages {pages_obj_num} 0 R >>".encode("ascii"))

    buf = BytesIO()
    buf.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for idx, obj in enumerate(objects, start=1):
        offsets.append(buf.tell())
        buf.write(f"{idx} 0 obj\n".encode("ascii"))
        buf.write(obj)
        buf.write(b"\nendobj\n")
    xref_pos = buf.tell()
    buf.write(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    buf.write(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        buf.write(f"{offset:010d} 00000 n \n".encode("ascii"))
    buf.write(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root {catalog_obj} 0 R >>\n"
            f"startxref\n{xref_pos}\n%%EOF"
        ).encode("ascii")
    )
    return buf.getvalue()


@app.post("/exports/docx")
def export_docx(payload: ExportDocxRequest):
    title = (payload.title or "").strip() or "ELume Resource"
    content = payload.content or ""
    teacher = (payload.teacher or "").strip() or None

    data = _docx_from_markdownish(title, content, teacher=teacher, meta=payload.meta or {})

    filename_safe = re.sub(r"[^a-zA-Z0-9_\- ]+", "", title).strip().replace(" ", "_")
    if not filename_safe:
        filename_safe = "ELume_Resource"
    filename = f"{filename_safe}.docx"

    return StreamingResponse(
        BytesIO(data),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/exports/pdf")
def export_pdf(payload: ExportDocxRequest):
    title = (payload.title or "").strip() or "ELume Resource"
    content = payload.content or ""
    teacher = (payload.teacher or "").strip() or None

    data = _pdf_from_markdownish(title, content, teacher=teacher, meta=payload.meta or {})

    filename_safe = re.sub(r"[^a-zA-Z0-9_\- ]+", "", title).strip().replace(" ", "_")
    if not filename_safe:
        filename_safe = "ELume_Resource"
    filename = f"{filename_safe}.pdf"

    return StreamingResponse(
        BytesIO(data),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )

# =========================================================
# CORS
# =========================================================
from starlette.middleware.cors import CORSMiddleware

app.add_middleware(
    CORSMiddleware,
    allow_origin_regex=r"^https?://(localhost|127\.0\.0\.1|192\.168\.\d+\.\d+)(:\d+)?$",
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

@app.post("/auth/register", response_model=AuthRegisterResponse)
def auth_register(payload: AuthRegister, db: Session = Depends(get_db)):
    first_name = (payload.first_name or "").strip()
    last_name = (payload.last_name or "").strip()
    school_name = (payload.school_name or "").strip()
    email = (payload.email or "").strip().lower()
    password = (payload.password or "").strip()
    logger.info("auth_register started for %s", email)

    if not first_name or not last_name or not school_name:
        raise HTTPException(status_code=400, detail="First name, last name, and school name are required")
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Valid email required")
    password_error = _password_policy_error(password)
    if password_error:
        raise HTTPException(status_code=400, detail=password_error)

    existing = db.query(models.UserModel).filter(models.UserModel.email == email).first()
    if existing:
        raise HTTPException(status_code=400, detail="Email already registered")

    user = models.UserModel(
        first_name=first_name,
        last_name=last_name,
        school_name=school_name,
        email=email,
        password_hash=PWD_CONTEXT.hash(password),
        email_verified=False,
        billing_onboarding_required=True,
    )

    try:
        db.add(user)
        db.flush()

        raw_token = secrets.token_urlsafe(32)
        now = datetime.utcnow()

        db.add(
            models.EmailVerificationTokenModel(
                user_id=user.id,
                token_hash=_hash_email_verification_token(raw_token),
                expires_at=now + timedelta(hours=24),
            )
        )

        verify_link = f"{APP_BASE_URL.rstrip('/')}/#/verify-email?token={raw_token}"
        body = (
            f"Hello {first_name},\n\n"
            "Welcome to Elume.\n\n"
            "Please verify your email to activate your teacher account:\n"
            f"{verify_link}\n\n"
            "This link will expire in 24 hours.\n\n"
            "Elume"
        )

        _send_email(user.email, "Verify your Elume account", body)

        admin_notify_email = (os.getenv("ADMIN_SIGNUP_NOTIFY_EMAIL") or "admin@elume.ie").strip()
        if admin_notify_email:
            admin_body = (
                "New Elume teacher signup\n\n"
                f"First name: {first_name}\n"
                f"Last name: {last_name}\n"
                f"School name: {school_name}\n"
                f"Email: {email}\n"
                f"Signup timestamp (UTC): {now.isoformat()}Z\n"
            )
            try:
                _send_email(admin_notify_email, "New Elume teacher signup", admin_body)
            except Exception:
                logger.exception("Failed to send admin signup notification for %s", email)

        db.commit()
        db.refresh(user)

    except HTTPException:
        db.rollback()
        raise
    except Exception as e:
        db.rollback()
        logger.exception("Failed to register teacher account for %s; error=%r", email, e)
        raise HTTPException(status_code=500, detail="Failed to create account")

    return {
        "success": True,
        "message": "Account created. Please check your email to verify your account before signing in.",
    }

@app.post("/auth/dev-auto-login", response_model=DevAutoLoginResponse)
def auth_dev_auto_login(
    db: Session = Depends(get_db),
    host: Optional[str] = Header(default=None, alias="host"),
):
    if not _dev_autologin_enabled():
        raise HTTPException(status_code=404, detail="Not found")

    if not _is_explicit_development():
        raise HTTPException(status_code=404, detail="Not found")

    if not _is_local_request(host):
        raise HTTPException(status_code=403, detail="Local development only")

    if _jwt_secret_is_weak(JWT_SECRET):
        raise HTTPException(status_code=403, detail="JWT secret is not configured securely")

    email = "admin@elume.ie"

    user = db.query(models.UserModel).filter(models.UserModel.email == email).first()
    if not user:
        user = models.UserModel(
            email=email,
            password_hash=PWD_CONTEXT.hash("dev-only-placeholder-password"),
            email_verified=True,
        )
        db.add(user)
        db.commit()
        db.refresh(user)

    return {
        "access_token": _make_token(user.id, user.email),
        "token_type": "bearer",
    }

@app.post("/auth/login", response_model=AuthToken)
def auth_login(payload: AuthLogin, db: Session = Depends(get_db)):
    email = (payload.email or "").strip().lower()
    password = (payload.password or "").strip()

    user = db.query(models.UserModel).filter(models.UserModel.email == email).first()
    if not user or not PWD_CONTEXT.verify(password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid credentials")
    if not bool(user.email_verified):
        raise HTTPException(status_code=401, detail="Please verify your email before signing in.")

    return {"access_token": _make_token(user.id, user.email), "token_type": "bearer"}


@app.post("/auth/verify-email", response_model=VerifyEmailResponse)
def auth_verify_email(payload: VerifyEmailPayload, db: Session = Depends(get_db)):
    token = (payload.token or "").strip()
    if not token:
        raise HTTPException(status_code=400, detail="Verification token required")

    token_row = db.query(models.EmailVerificationTokenModel).filter(
        models.EmailVerificationTokenModel.token_hash == _hash_email_verification_token(token),
        models.EmailVerificationTokenModel.used_at.is_(None),
        models.EmailVerificationTokenModel.expires_at > datetime.utcnow(),
    ).first()
    if not token_row:
        raise HTTPException(status_code=400, detail="Invalid or expired verification token")

    user = db.query(models.UserModel).filter(models.UserModel.id == token_row.user_id).first()
    if not user:
        raise HTTPException(status_code=400, detail="Invalid or expired verification token")

    try:
        user.email_verified = True
        token_row.used_at = datetime.utcnow()
        _seed_demo_class(db, user)
        db.commit()
    except HTTPException:
        db.rollback()
        raise
    except Exception:
        db.rollback()
        logger.exception("Failed to verify teacher account for %s", user.email)
        raise HTTPException(status_code=500, detail="Failed to verify email")

    return {
        "success": True,
        "message": "Email verified. Let's finish your Elume setup.",
        "access_token": _make_token(user.id, user.email),
        "token_type": "bearer",
        "next_path": "/onboarding/billing",
    }


@app.post("/auth/forgot-password")
def auth_forgot_password(payload: schemas.ForgotPasswordRequest, db: Session = Depends(get_db)):
    email = (payload.email or "").strip().lower()
    generic_response = {
        "success": True,
        "message": "If that email is registered, a password reset link has been sent.",
    }

    if not email:
        return generic_response

    user = db.query(models.UserModel).filter(models.UserModel.email == email).first()
    if not user:
        return generic_response

    now = datetime.utcnow()
    db.query(models.PasswordResetTokenModel).filter(
        models.PasswordResetTokenModel.user_id == user.id,
        models.PasswordResetTokenModel.used_at.is_(None),
    ).update({"used_at": now}, synchronize_session=False)

    raw_token = secrets.token_urlsafe(32)
    db.add(
        models.PasswordResetTokenModel(
            user_id=user.id,
            token_hash=_hash_reset_token(raw_token),
            expires_at=now + timedelta(hours=1),
        )
    )
    db.commit()

    reset_link = f"{APP_BASE_URL.rstrip('/')}/#/reset-password?token={raw_token}"
    body = (
        "Hello,\n\n"
        "We received a request to reset your Elume password.\n\n"
        f"Reset your password here:\n{reset_link}\n\n"
        "If you did not request this, you can ignore this email.\n"
        "This link will expire in 1 hour.\n\n"
        "Elume"
    )

    try:
        _send_email(user.email, "Reset your Elume password", body)
    except Exception:
        logger.exception("Failed to send password reset email to %s", user.email)
        raise HTTPException(status_code=500, detail="Failed to send reset email")

    return generic_response

@app.post("/auth/reset-password")
def auth_reset_password(payload: schemas.ResetPasswordRequest, db: Session = Depends(get_db)):
    token = (payload.token or "").strip()
    new_password = (payload.new_password or "").strip()

    password_error = _password_policy_error(new_password)
    if password_error:
        raise HTTPException(status_code=400, detail=password_error)

    token_row = db.query(models.PasswordResetTokenModel).filter(
        models.PasswordResetTokenModel.token_hash == _hash_reset_token(token),
        models.PasswordResetTokenModel.used_at.is_(None),
        models.PasswordResetTokenModel.expires_at > datetime.utcnow(),
    ).first()
    if not token_row:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")

    user = db.query(models.UserModel).filter(models.UserModel.id == token_row.user_id).first()
    if not user:
        raise HTTPException(status_code=400, detail="Invalid or expired reset token")

    user.password_hash = PWD_CONTEXT.hash(new_password)
    token_row.used_at = datetime.utcnow()
    db.commit()

    return {"success": True, "message": "Password reset successful."}

def get_current_user(
    authorization: Optional[str] = Header(default=None),
    db: Session = Depends(get_db),
) -> models.UserModel:
    if not authorization or not authorization.lower().startswith("bearer "):
        raise HTTPException(status_code=401, detail="Missing bearer token")

    token = authorization.split(" ", 1)[1].strip()

    try:
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALG])
        user_id = int(payload.get("sub"))
    except (JWTError, ValueError):
        raise HTTPException(status_code=401, detail="Invalid token")

    user = db.query(models.UserModel).filter(models.UserModel.id == user_id).first()
    if not user:
        raise HTTPException(status_code=401, detail="User not found")

    return user

@app.post("/billing/create-checkout-session", response_model=CreateCheckoutSessionResponse)
def create_checkout_session(
    payload: CreateCheckoutSessionRequest,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    if not STRIPE_SECRET_KEY:
        raise HTTPException(status_code=500, detail="Stripe is not configured")
    if not STRIPE_PRICE_MONTHLY_EUR or not STRIPE_PRICE_ANNUAL_EUR:
        raise HTTPException(status_code=500, detail="Stripe prices are not configured")
    if not bool(user.email_verified):
        raise HTTPException(status_code=403, detail="Please verify your email before starting billing.")
    if (user.subscription_status or "").strip().lower() in {"active", "trialing"}:
        raise HTTPException(
            status_code=400,
            detail="You already have an active subscription. Use the billing portal to manage it.",
        )

    plan = (payload.plan or "").strip().lower()
    if plan not in {"monthly", "annual"}:
        raise HTTPException(status_code=400, detail="plan must be 'monthly' or 'annual'")

    price_id = STRIPE_PRICE_MONTHLY_EUR if plan == "monthly" else STRIPE_PRICE_ANNUAL_EUR
    launch_offer_candidate = (plan == "annual" and _annual_launch_offer_applies())

    try:
        valid_customer_id = _validate_stored_stripe_customer(db, user)
        session_kwargs = {
            "mode": "subscription",
            "payment_method_types": ["card"],
            "payment_method_collection": "always",
            "client_reference_id": str(user.id),
            "line_items": [
                {
                    "price": price_id,
                    "quantity": 1,
                }
            ],
            "success_url": f"{APP_BASE_URL.rstrip('/')}/#/billing/success?session_id={{CHECKOUT_SESSION_ID}}",
            "cancel_url": f"{APP_BASE_URL.rstrip('/')}/#/billing/cancel",
            "metadata": {
                "user_id": str(user.id),
                "email": user.email,
                "plan": plan,
                "launch_offer_candidate": "true" if launch_offer_candidate else "false",
            },
            "subscription_data": {
                "trial_period_days": 14,
                "metadata": {
                    "user_id": str(user.id),
                    "email": user.email,
                    "plan": plan,
                    "launch_offer_candidate": "true" if launch_offer_candidate else "false",
                }
            },
        }

        if valid_customer_id:
            session_kwargs["customer"] = valid_customer_id
        else:
            session_kwargs["customer_email"] = user.email

        session = stripe.checkout.Session.create(
            **session_kwargs,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create Stripe checkout session: {e}")

    user.subscription_status = "pending"
    user.billing_interval = plan
    user.stripe_checkout_session_id = session.id
    db.commit()

    return {"checkout_url": session.url}


@app.post("/billing/create-portal-session", response_model=CreatePortalSessionResponse)
def create_portal_session(
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    if not STRIPE_SECRET_KEY:
        raise HTTPException(status_code=500, detail="Stripe is not configured")
    if not (user.stripe_customer_id or "").strip():
        raise HTTPException(status_code=400, detail="No Stripe customer found for this account")

    try:
        valid_customer_id = _validate_stored_stripe_customer(db, user, commit_on_clear=True)
        if not valid_customer_id:
            raise HTTPException(
                status_code=400,
                detail="This account's saved billing customer is no longer valid in the current Stripe environment. Please reconnect billing.",
            )
        session = stripe.billing_portal.Session.create(
            customer=valid_customer_id,
            return_url=f"{APP_BASE_URL.rstrip('/')}/#/",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create billing portal session: {e}")

    return {"portal_url": session.url}


@app.post("/billing/confirm-checkout-session", response_model=ConfirmCheckoutSessionResponse)
def confirm_checkout_session(
    session_id: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    if not STRIPE_SECRET_KEY:
        raise HTTPException(status_code=500, detail="Stripe is not configured")

    checkout_session_id = (session_id or "").strip()
    if not checkout_session_id:
        raise HTTPException(status_code=400, detail="Checkout session id required")

    try:
        session = stripe.checkout.Session.retrieve(checkout_session_id)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to confirm Stripe checkout session: {e}")

    matched_user = _find_billing_user(db, session)
    if not matched_user or matched_user.id != user.id:
        raise HTTPException(status_code=404, detail="Checkout session not found for this account")

    session_customer = _stripe_obj_get(session, "customer")
    session_subscription = _stripe_obj_get(session, "subscription")
    session_status = (_stripe_obj_get(session, "payment_status") or "").strip().lower()
    billing_interval = _resolve_billing_interval(session, user.billing_interval)
    metadata = _stripe_obj_get(session, "metadata") or {}

    if session_customer:
        user.stripe_customer_id = str(session_customer)
    if session_subscription:
        user.stripe_subscription_id = str(session_subscription)

    user.stripe_checkout_session_id = str(_stripe_obj_get(session, "id") or user.stripe_checkout_session_id or "")
    user.billing_interval = billing_interval

    if session_subscription:
        try:
            subscription = stripe.Subscription.retrieve(str(session_subscription))
            _apply_subscription_update(user, subscription)
        except Exception:
            logger.warning("Failed to refresh subscription %s during checkout confirmation", session_subscription)
            user.subscription_status = "active" if session_status == "paid" else (user.subscription_status or "pending")
    else:
        user.subscription_status = "active" if session_status == "paid" else (user.subscription_status or "pending")

    if user.subscription_status == "active" and not user.subscription_started_at:
        user.subscription_started_at = datetime.utcnow()
    if user.subscription_status in {"active", "trialing"}:
        user.billing_onboarding_required = False
    if (metadata.get("launch_offer_candidate") or "").strip().lower() == "true" and user.subscription_status in {"active", "trialing"}:
        user.launch_offer_applied = True

    db.commit()
    return {
        "success": True,
        "billing_status": _billing_status_payload(user),
    }


@app.post("/billing/start-trial", response_model=StartTrialResponse)
def start_billing_trial(
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    if not bool(user.email_verified):
        raise HTTPException(status_code=403, detail="Please verify your email before starting your trial.")
    if _is_paid_subscription_active(user):
        raise HTTPException(status_code=400, detail="You already have an active plan.")
    if user.trial_started_at or user.subscription_started_at or user.stripe_customer_id:
        raise HTTPException(status_code=400, detail="The free trial has already been used for this account.")

    now = _utcnow()
    user.trial_started_at = now
    user.trial_ends_at = now + timedelta(days=14)
    user.ai_daily_limit = 5
    user.ai_prompt_count = 0
    user.ai_prompt_count_date = now
    db.commit()

    return {
        "success": True,
        "message": "Your 14-day Elume trial is now active.",
    }


@app.post("/billing/webhook")
async def stripe_billing_webhook(request: Request, db: Session = Depends(get_db)):
    if not STRIPE_SECRET_KEY or not STRIPE_WEBHOOK_SECRET:
        raise HTTPException(status_code=500, detail="Stripe webhook is not configured")

    payload = await request.body()
    sig_header = request.headers.get("stripe-signature")
    if not sig_header:
        raise HTTPException(status_code=400, detail="Missing Stripe signature")

    try:
        event = stripe.Webhook.construct_event(payload, sig_header, STRIPE_WEBHOOK_SECRET)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid webhook payload")
    except stripe.error.SignatureVerificationError:
        raise HTTPException(status_code=400, detail="Invalid webhook signature")

    event_type = (event.get("type") or "").strip()
    stripe_obj = event.get("data", {}).get("object", {})

    if event_type not in {
        "checkout.session.completed",
        "customer.subscription.created",
        "customer.subscription.updated",
        "customer.subscription.deleted",
        "invoice.paid",
        "invoice.payment_failed",
    }:
        return {"received": True}

    user = _find_billing_user(db, stripe_obj)
    if not user:
        logger.warning("Stripe webhook could not match user for event %s", event_type)
        return {"received": True}

    if event_type == "checkout.session.completed":
        session_customer = _stripe_obj_get(stripe_obj, "customer")
        session_subscription = _stripe_obj_get(stripe_obj, "subscription")
        session_status = (_stripe_obj_get(stripe_obj, "payment_status") or "").strip().lower()
        billing_interval = _resolve_billing_interval(stripe_obj, user.billing_interval)
        metadata = _stripe_obj_get(stripe_obj, "metadata") or {}

        if session_customer:
            user.stripe_customer_id = str(session_customer)
        if session_subscription:
            user.stripe_subscription_id = str(session_subscription)

        user.stripe_checkout_session_id = str(_stripe_obj_get(stripe_obj, "id") or user.stripe_checkout_session_id or "")
        user.billing_interval = billing_interval
        user.subscription_status = "active" if session_status == "paid" else "pending"

        if user.subscription_status == "active" and not user.subscription_started_at:
            user.subscription_started_at = datetime.utcnow()
        if user.subscription_status == "active":
            user.trial_started_at = None
            user.trial_ends_at = None

        if (metadata.get("launch_offer_candidate") or "").strip().lower() == "true" and user.subscription_status == "active":
            user.launch_offer_applied = True

    elif event_type in {"customer.subscription.created", "customer.subscription.updated"}:
        _apply_subscription_update(user, stripe_obj)
    if _is_paid_subscription_active(user):
        user.trial_started_at = None
        user.trial_ends_at = None
        user.billing_onboarding_required = False

    elif event_type == "customer.subscription.deleted":
        _apply_subscription_update(user, stripe_obj)
        user.subscription_status = "canceled"

    elif event_type == "invoice.payment_failed":
        customer_id = _stripe_obj_get(stripe_obj, "customer")
        subscription_id = _stripe_obj_get(stripe_obj, "subscription")
        if customer_id:
            user.stripe_customer_id = str(customer_id)
        if subscription_id:
            user.stripe_subscription_id = str(subscription_id)
        user.subscription_status = "past_due"

    elif event_type == "invoice.paid":
        customer_id = _stripe_obj_get(stripe_obj, "customer")
        subscription_id = _stripe_obj_get(stripe_obj, "subscription")
        if customer_id:
            user.stripe_customer_id = str(customer_id)
        if subscription_id:
            user.stripe_subscription_id = str(subscription_id)

        refreshed = False
        if subscription_id:
            try:
                subscription = stripe.Subscription.retrieve(str(subscription_id))
                _apply_subscription_update(user, subscription)
                refreshed = True
            except Exception:
                logger.warning("Failed to refresh subscription %s after invoice.paid", subscription_id)

        if not refreshed:
            user.subscription_status = "active"
        if _is_paid_subscription_active(user):
            user.trial_started_at = None
            user.trial_ends_at = None
            user.billing_onboarding_required = False

    _refresh_ai_daily_limit(user)
    db.commit()
    return {"received": True}


@app.get("/billing/me", response_model=schemas.BillingStatusOut)
def billing_me(
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    payload = _billing_status_payload(user)
    db.commit()
    return payload

def get_owned_class_or_404(
    class_id: int,
    db: Session,
    user: models.UserModel,
) -> ClassModel:
    cls = db.query(ClassModel).filter(
        ClassModel.id == class_id,
        ClassModel.owner_user_id == user.id,
    ).first()
    if not cls:
        raise HTTPException(status_code=404, detail="Class not found")
    return cls


def _get_owned_class_for_cat4_or_403(
    class_id: int,
    db: Session,
    user: models.UserModel,
) -> ClassModel:
    require_cat4_access(user)
    return get_owned_class_or_404(class_id, db, user)


def _cat4_baseline_sets_for_class(class_id: int, db: Session) -> list[Cat4BaselineSetModel]:
    return (
        db.query(Cat4BaselineSetModel)
        .filter(Cat4BaselineSetModel.class_id == class_id)
        .order_by(Cat4BaselineSetModel.test_date.desc(), Cat4BaselineSetModel.created_at.desc(), Cat4BaselineSetModel.id.desc())
        .all()
    )


def _cat4_term_sets_for_class(class_id: int, db: Session) -> list[Cat4TermResultSetModel]:
    return (
        db.query(Cat4TermResultSetModel)
        .filter(Cat4TermResultSetModel.class_id == class_id)
        .order_by(Cat4TermResultSetModel.created_at.desc(), Cat4TermResultSetModel.id.desc())
        .all()
    )


def _build_cat4_meta_payload(class_id: int, db: Session) -> dict[str, Any]:
    baseline_sets = _cat4_baseline_sets_for_class(class_id, db)
    term_sets = _cat4_term_sets_for_class(class_id, db)
    workbook_versions = (
        db.query(Cat4WorkbookVersionModel)
        .filter(Cat4WorkbookVersionModel.class_id == class_id)
        .order_by(Cat4WorkbookVersionModel.version_number.desc(), Cat4WorkbookVersionModel.uploaded_at.desc())
        .all()
    )

    baseline_ids = [row.id for row in baseline_sets]
    term_ids = [row.id for row in term_sets]

    baseline_rows = (
        db.query(Cat4StudentBaselineModel)
        .filter(Cat4StudentBaselineModel.baseline_set_id.in_(baseline_ids))
        .all()
        if baseline_ids
        else []
    )
    term_rows = (
        db.query(Cat4StudentTermResultModel)
        .filter(Cat4StudentTermResultModel.result_set_id.in_(term_ids))
        .all()
        if term_ids
        else []
    )

    baseline_rows_by_set: dict[int, list[Cat4StudentBaselineModel]] = defaultdict(list)
    for row in baseline_rows:
        baseline_rows_by_set[row.baseline_set_id].append(row)

    term_rows_by_set: dict[int, list[Cat4StudentTermResultModel]] = defaultdict(list)
    for row in term_rows:
        term_rows_by_set[row.result_set_id].append(row)

    return {
        "feature_enabled": True,
        "active_workbook": next(
            (
                {
                    "id": item.id,
                    "version_number": item.version_number,
                    "workbook_name": item.workbook_name,
                    "uploaded_by_email": item.uploaded_by_email,
                    "uploaded_at": item.uploaded_at.isoformat() if item.uploaded_at else None,
                    "validation_summary": json.loads(item.validation_summary_json or "{}"),
                }
                for item in workbook_versions
                if item.is_active
            ),
            None,
        ),
        "workbook_versions": [
            {
                "id": item.id,
                "version_number": item.version_number,
                "workbook_name": item.workbook_name,
                "uploaded_by_email": item.uploaded_by_email,
                "uploaded_at": item.uploaded_at.isoformat() if item.uploaded_at else None,
                "is_active": bool(item.is_active),
                "validation_summary": json.loads(item.validation_summary_json or "{}"),
            }
            for item in workbook_versions
        ],
        "baseline_sets": [
            {
                "id": item.id,
                "title": item.title,
                "test_date": item.test_date.date().isoformat() if item.test_date else None,
                "is_locked": bool(item.is_locked),
                "locked_at": item.locked_at.isoformat() if item.locked_at else None,
                "created_at": item.created_at.isoformat() if item.created_at else None,
                **_cat4_set_summary(baseline_rows_by_set.get(item.id, [])),
            }
            for item in baseline_sets
        ],
        "term_sets": [
            {
                "id": item.id,
                "title": item.title,
                "academic_year": item.academic_year,
                "term_key": item.term_key,
                "created_at": item.created_at.isoformat() if item.created_at else None,
                **_cat4_set_summary(term_rows_by_set.get(item.id, [])),
            }
            for item in term_sets
        ],
        "matched_counts": {
            "baseline_rows": sum(1 for row in baseline_rows if row.student_id),
            "baseline_unmatched": sum(1 for row in baseline_rows if not row.student_id),
            "term_rows": sum(1 for row in term_rows if row.student_id),
            "term_unmatched": sum(1 for row in term_rows if not row.student_id),
        },
    }


def _build_cat4_report_payload(
    class_id: int,
    db: Session,
    baseline_id: Optional[int] = None,
    term_set_id: Optional[int] = None,
) -> dict[str, Any]:
    baseline_sets = _cat4_baseline_sets_for_class(class_id, db)
    term_sets = _cat4_term_sets_for_class(class_id, db)

    if not baseline_sets or not term_sets:
        return {
            "feature_enabled": True,
            "baseline_set": None,
            "latest_term_set": None,
            "previous_term_set": None,
            "summary_cards": [],
            "at_risk": [],
            "excelling": [],
            "within_expected_range": [],
            "all_matched_students": [],
            "unmatched_cat4_rows": [],
            "unmatched_term_rows": [],
            "profile_distribution": [],
            "domain_commentary": [],
        }

    selected_baseline = next((item for item in baseline_sets if item.id == baseline_id), baseline_sets[0]) if baseline_id else baseline_sets[0]
    selected_term = next((item for item in term_sets if item.id == term_set_id), term_sets[0]) if term_set_id else term_sets[0]
    previous_term = next((item for item in term_sets if item.id != selected_term.id), None)

    baseline_rows = (
        db.query(Cat4StudentBaselineModel)
        .filter(Cat4StudentBaselineModel.baseline_set_id == selected_baseline.id)
        .all()
    )
    latest_rows = (
        db.query(Cat4StudentTermResultModel)
        .filter(Cat4StudentTermResultModel.result_set_id == selected_term.id)
        .all()
    )
    previous_rows = (
        db.query(Cat4StudentTermResultModel)
        .filter(Cat4StudentTermResultModel.result_set_id == previous_term.id)
        .all()
        if previous_term
        else []
    )

    students = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id)
        .order_by(StudentModel.first_name.asc())
        .all()
    )
    student_name_by_id = {student.id: student.first_name for student in students}

    baseline_by_student = {row.student_id: row for row in baseline_rows if row.student_id}
    latest_by_student = {
        row.student_id: row
        for row in latest_rows
        if row.student_id and row.average_percent is not None
    }
    previous_by_student = {
        row.student_id: row
        for row in previous_rows
        if row.student_id and row.average_percent is not None
    }

    latest_percentiles = _build_percentile_map(
        [(student_id, float(row.average_percent)) for student_id, row in latest_by_student.items()]
    )
    previous_percentiles = _build_percentile_map(
        [(student_id, float(row.average_percent)) for student_id, row in previous_by_student.items()]
    )

    matched_student_ids = sorted(set(baseline_by_student.keys()) & set(latest_by_student.keys()))
    matched_students: list[dict[str, Any]] = []

    domain_baseline_maps = {
        "verbal_domain_score": _build_percentile_map([(student_id, float(row.verbal_sas)) for student_id, row in baseline_by_student.items() if row.verbal_sas is not None]),
        "quantitative_domain_score": _build_percentile_map([(student_id, float(row.quantitative_sas)) for student_id, row in baseline_by_student.items() if row.quantitative_sas is not None]),
        "non_verbal_domain_score": _build_percentile_map([(student_id, float(row.non_verbal_sas)) for student_id, row in baseline_by_student.items() if row.non_verbal_sas is not None]),
        "spatial_domain_score": _build_percentile_map([(student_id, float(row.spatial_sas)) for student_id, row in baseline_by_student.items() if row.spatial_sas is not None]),
    }
    domain_term_maps = {
        "verbal_domain_score": _build_percentile_map([(student_id, float(row.verbal_domain_score)) for student_id, row in latest_by_student.items() if row.verbal_domain_score is not None]),
        "quantitative_domain_score": _build_percentile_map([(student_id, float(row.quantitative_domain_score)) for student_id, row in latest_by_student.items() if row.quantitative_domain_score is not None]),
        "non_verbal_domain_score": _build_percentile_map([(student_id, float(row.non_verbal_domain_score)) for student_id, row in latest_by_student.items() if row.non_verbal_domain_score is not None]),
        "spatial_domain_score": _build_percentile_map([(student_id, float(row.spatial_domain_score)) for student_id, row in latest_by_student.items() if row.spatial_domain_score is not None]),
    }
    domain_labels = {
        "verbal_domain_score": "Verbal",
        "quantitative_domain_score": "Quantitative",
        "non_verbal_domain_score": "Non-Verbal",
        "spatial_domain_score": "Spatial",
    }

    for student_id in matched_student_ids:
        baseline_row = baseline_by_student[student_id]
        latest_row = latest_by_student[student_id]
        previous_row = previous_by_student.get(student_id)

        latest_term_percentile = latest_percentiles.get(student_id)
        previous_term_percentile = previous_percentiles.get(student_id)
        latest_average = float(latest_row.average_percent) if latest_row.average_percent is not None else None
        previous_average = float(previous_row.average_percent) if previous_row and previous_row.average_percent is not None else None
        trend_delta = round(latest_average - previous_average, 1) if latest_average is not None and previous_average is not None else None
        domain_movements: dict[str, Optional[float]] = {}
        domain_components: list[float] = []
        for domain_key, label in domain_labels.items():
            baseline_pct = domain_baseline_maps[domain_key].get(student_id)
            latest_pct = domain_term_maps[domain_key].get(student_id)
            if baseline_pct is None or latest_pct is None:
                domain_movements[label] = None
                continue
            delta = round(latest_pct - baseline_pct, 1)
            domain_movements[label] = delta
            domain_components.append(delta)

        movement_score = round(sum(domain_components) / len(domain_components), 1) if domain_components else None
        value_added_delta = movement_score

        matched_students.append(
            {
                "student_id": student_id,
                "student_name": latest_row.matched_name or baseline_row.matched_name or student_name_by_id.get(student_id) or latest_row.raw_name,
                "profile_label": baseline_row.profile_label,
                "baseline_percentile": round(sum(domain_baseline_maps[key].get(student_id, 0.0) for key in domain_labels if domain_baseline_maps[key].get(student_id) is not None) / max(1, sum(1 for key in domain_labels if domain_baseline_maps[key].get(student_id) is not None)), 1) if any(domain_baseline_maps[key].get(student_id) is not None for key in domain_labels) else None,
                "latest_term_percentile": latest_term_percentile,
                "previous_term_percentile": previous_term_percentile,
                "value_added_delta": value_added_delta,
                "trend_delta": trend_delta,
                "latest_average_percent": latest_average,
                "previous_average_percent": previous_average,
                "movement_score": movement_score,
                "domain_movements": domain_movements,
                "status": "within_expected_range",
                "reasons": [],
            }
        )

    ranked = sorted(
        [row for row in matched_students if row.get("movement_score") is not None],
        key=lambda row: (float(row.get("movement_score") or 0), str(row.get("student_name") or "").lower()),
    )
    count = len(ranked)
    if count > 1:
        band_size = max(1, int(round(count * 0.1)))
        bottom_ids = {row["student_id"] for row in ranked[:band_size]}
        top_ids = {row["student_id"] for row in ranked[-band_size:]}
    else:
        bottom_ids = set()
        top_ids = set()

    for row in matched_students:
        if row["student_id"] in bottom_ids:
            row["status"] = "at_risk"
            row["reasons"] = ["Movement score is in the bottom cohort band"]
        elif row["student_id"] in top_ids:
            row["status"] = "excelling"
            row["reasons"] = ["Movement score is in the top cohort band"]
        else:
            row["status"] = "within_expected_range"
            row["reasons"] = ["Performance is within the expected cohort range"]

    matched_students.sort(key=lambda row: ({"at_risk": 0, "excelling": 1, "within_expected_range": 2}.get(str(row.get("status")), 3), (row.get("student_name") or "").lower()))

    profile_distribution = Counter(
        row.profile_label.strip()
        for row in baseline_rows
        if isinstance(row.profile_label, str) and row.profile_label.strip()
    )

    domain_commentary = []
    for domain_key, label in domain_labels.items():
        values = [row["domain_movements"].get(label) for row in matched_students if isinstance(row.get("domain_movements"), dict)]
        values = [float(value) for value in values if value is not None]
        if not values:
            continue
        avg_delta = round(sum(values) / len(values), 1)
        if avg_delta >= 8:
            comment = f"{label} performance is trending clearly above CAT4 baseline expectations."
        elif avg_delta <= -8:
            comment = f"{label} performance is trending below CAT4 baseline expectations."
        else:
            comment = f"{label} performance is broadly in line with CAT4 baseline expectations."
        domain_commentary.append({"domain": label, "average_movement": avg_delta, "commentary": comment})

    return {
        "feature_enabled": True,
        "baseline_set": {
            "id": selected_baseline.id,
            "title": selected_baseline.title,
            "test_date": selected_baseline.test_date.date().isoformat() if selected_baseline.test_date else None,
            "is_locked": bool(selected_baseline.is_locked),
            "locked_at": selected_baseline.locked_at.isoformat() if selected_baseline.locked_at else None,
        },
        "latest_term_set": {
            "id": selected_term.id,
            "title": selected_term.title,
            "academic_year": selected_term.academic_year,
            "term_key": selected_term.term_key,
        },
        "previous_term_set": (
            {
                "id": previous_term.id,
                "title": previous_term.title,
                "academic_year": previous_term.academic_year,
                "term_key": previous_term.term_key,
            }
            if previous_term
            else None
        ),
        "summary_cards": [
            {"key": "matched", "label": "Matched Students", "value": len(matched_students)},
            {"key": "at_risk", "label": "At Risk", "value": sum(1 for row in matched_students if row["status"] == "at_risk")},
            {"key": "excelling", "label": "Excelling", "value": sum(1 for row in matched_students if row["status"] == "excelling")},
            {"key": "within_expected_range", "label": "Within Expected Range", "value": sum(1 for row in matched_students if row["status"] == "within_expected_range")},
        ],
        "at_risk": [row for row in matched_students if row["status"] == "at_risk"],
        "excelling": [row for row in matched_students if row["status"] == "excelling"],
        "within_expected_range": [row for row in matched_students if row["status"] == "within_expected_range"],
        "all_matched_students": matched_students,
        "unmatched_cat4_rows": [
            {
                "id": row.id,
                "raw_name": row.raw_name,
                "matched_name": row.matched_name,
                "confidence_note": row.confidence_note,
                "overall_sas": row.overall_sas,
                "profile_label": row.profile_label,
            }
            for row in baseline_rows
            if not row.student_id
        ],
        "unmatched_term_rows": [
            {
                "id": row.id,
                "raw_name": row.raw_name,
                "matched_name": row.matched_name,
                "confidence_note": None,
                "average_percent": row.average_percent,
                "subject_count": row.subject_count,
            }
            for row in latest_rows
            if not row.student_id
        ],
        "profile_distribution": [
            {"label": label, "count": count}
            for label, count in sorted(profile_distribution.items(), key=lambda item: (-item[1], item[0].lower()))
        ],
        "domain_commentary": domain_commentary,
    }

def get_owned_livequiz_session_or_404(
    code: str,
    db: Session,
    user: models.UserModel,
) -> LiveQuizSessionModel:
    session = db.query(LiveQuizSessionModel).filter(
        LiveQuizSessionModel.session_code == code
    ).first()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    get_owned_class_or_404(session.class_id, db, user)
    return session

def get_owned_collab_session_or_404(
    code: str,
    db: Session,
    user: models.UserModel,
) -> CollabSessionModel:
    session = db.query(CollabSessionModel).filter(
        CollabSessionModel.session_code == code
    ).first()
    if not session:
        raise HTTPException(status_code=404, detail="Session not found")
    get_owned_class_or_404(session.class_id, db, user)
    return session

def get_active_student_access_link_or_404(
    token: str,
    db: Session,
) -> models.StudentAccessLink:
    link = db.query(models.StudentAccessLink).filter(
        models.StudentAccessLink.token == token,
        models.StudentAccessLink.is_active == True,
    ).first()
    if not link:
        raise HTTPException(status_code=404, detail="Invalid link")
    return link


def _get_or_create_active_student_access_link(
    class_id: int,
    db: Session,
) -> models.StudentAccessLink:
    link = db.query(models.StudentAccessLink).filter(
        models.StudentAccessLink.class_id == class_id,
        models.StudentAccessLink.is_active == True,
    ).order_by(models.StudentAccessLink.created_at.desc()).first()

    if link:
        return link

    link = models.StudentAccessLink(
        class_id=class_id,
        token=secrets.token_urlsafe(24),
    )
    db.add(link)
    db.commit()
    db.refresh(link)
    return link


def _class_access_out(cls: ClassModel) -> schemas.ClassAccessOut:
    return schemas.ClassAccessOut(
        class_id=cls.id,
        class_code=(cls.class_code or "").strip(),
        class_pin=(cls.class_pin or "").strip(),
    )

def _internal_post_upload_relpath_or_none(link: str) -> Optional[str]:
    value = (link or "").strip()
    for prefix in ("/uploads/posts/", "/uploads/whiteboards/"):
        if value.startswith(prefix):
            return value[len("/uploads/"):]
    return None

def _resolve_upload_relpath_or_none(rel_path: str | None) -> Optional[Path]:
    if not rel_path:
        return None

    normalized = str(rel_path).replace("\\", "/").lstrip("/")
    direct_path = (UPLOADS_DIR / Path(normalized)).resolve()
    try:
        direct_path.relative_to(UPLOADS_DIR.resolve())
    except ValueError:
        direct_path = None
    if direct_path and direct_path.exists() and direct_path.is_file():
        return direct_path

    uploads_prefixed = f"uploads/{normalized}"
    return _resolve_stored_upload_path_or_none(uploads_prefixed)

def _post_attachment_path_or_404(post: PostModel, attachment_index: int) -> tuple[Path, str]:
    links = _links_to_list(getattr(post, "links", None))
    if attachment_index < 0 or attachment_index >= len(links):
        raise HTTPException(status_code=404, detail="Attachment not found")

    rel = _internal_post_upload_relpath_or_none(links[attachment_index])
    if not rel:
        raise HTTPException(status_code=404, detail="Attachment not found")

    path = _resolve_upload_relpath_or_none(rel)
    if not path:
        raise HTTPException(status_code=404, detail="File not found")

    return path, path.name

def _resolve_stored_upload_path_or_none(stored_path: str | None) -> Optional[Path]:
    if not stored_path:
        return None

    direct_path = Path(stored_path)
    if direct_path.exists() and direct_path.is_file():
        return direct_path

    normalized = str(stored_path).replace("\\", "/")
    marker = "uploads/"
    marker_index = normalized.lower().find(marker)
    if marker_index != -1:
        rel_suffix = normalized[marker_index + len(marker):].lstrip("/")
        fallback_path = (UPLOADS_DIR / Path(rel_suffix)).resolve()
        try:
            fallback_path.relative_to(UPLOADS_DIR.resolve())
        except ValueError:
            return None
        if fallback_path.exists() and fallback_path.is_file():
            return fallback_path

    return None


SAVED_WHITEBOARDS_TOPIC_NAME = "Saved Whiteboards"


def _get_or_create_saved_whiteboards_topic(class_id: int, db: Session) -> models.Topic:
    topic = (
        db.query(models.Topic)
        .filter(
            models.Topic.class_id == class_id,
            models.Topic.name == f"{NOTES_PREFIX}{SAVED_WHITEBOARDS_TOPIC_NAME}",
        )
        .first()
    )
    if topic:
        return topic

    topic = models.Topic(class_id=class_id, name=f"{NOTES_PREFIX}{SAVED_WHITEBOARDS_TOPIC_NAME}")
    db.add(topic)
    db.flush()
    return topic


def _safe_whiteboard_note_filename(title: str, suffix: str) -> str:
    base = (title or "Whiteboard").strip() or "Whiteboard"
    safe = re.sub(r'[<>:"/\\\\|?*]+', "", base).strip().rstrip(".")
    safe = safe or "Whiteboard"
    ext = suffix if suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"} else ".png"
    return f"{safe}{ext}"

def require_super_admin(user: models.UserModel):
    if (user.email or "").strip().lower() != "admin@elume.ie":
        raise HTTPException(status_code=403, detail="Not authorised")


def user_has_cat4_access(user: models.UserModel) -> bool:
    email = (user.email or "").strip().lower()
    if email == "admin@elume.ie":
        return True

    raw = (
        os.getenv("CAT4_ACCESS_EMAILS")
        or os.getenv("CAT4_ALLOWLIST_EMAILS")
        or ""
    )
    allowlist = {
        item.strip().lower()
        for item in raw.split(",")
        if item.strip()
    }
    return email in allowlist


def require_cat4_access(user: models.UserModel):
    if not user_has_cat4_access(user):
        raise HTTPException(status_code=403, detail="CAT4 Insights not enabled for this account")


def _normalise_student_name(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip()).lower()


def _cat4_name_index(students: list[StudentModel]) -> dict[str, list[StudentModel]]:
    index: dict[str, list[StudentModel]] = defaultdict(list)
    for student in students:
        key = _normalise_student_name(student.first_name)
        if key:
            index[key].append(student)
    return index


def _cat4_first_token(value: str) -> str:
    key = _normalise_student_name(value)
    if not key:
        return ""
    return key.split(" ", 1)[0].strip()


def _match_cat4_student_name(raw_name: str, name_index: dict[str, list[StudentModel]]) -> tuple[Optional[int], Optional[str], Optional[str]]:
    key = _normalise_student_name(raw_name)
    if not key:
        return None, None, "Missing student name"

    matches = name_index.get(key, [])
    if len(matches) > 1:
        return None, None, "Multiple class students matched"
    if len(matches) == 1:
        student = matches[0]
        return student.id, student.first_name, None

    first_token = _cat4_first_token(raw_name)
    if not first_token:
        return None, None, "No class student matched"

    first_token_matches = name_index.get(first_token, [])
    if len(first_token_matches) == 1:
        student = first_token_matches[0]
        return student.id, student.first_name, "Matched on first name only"
    if len(first_token_matches) > 1:
        return None, None, "Multiple class students matched on first name"

    return None, None, "No class student matched"


def _parse_optional_date(value: Optional[str]) -> Optional[datetime]:
    text_value = (value or "").strip()
    if not text_value:
        return None
    try:
        return datetime.fromisoformat(text_value)
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid date format; use YYYY-MM-DD")


def _json_text_or_none(value: Any) -> Optional[str]:
    if value in (None, ""):
        return None
    if isinstance(value, str):
        stripped = value.strip()
        return stripped or None
    try:
        return json.dumps(value, ensure_ascii=False)
    except Exception:
        raise HTTPException(status_code=400, detail="raw_subjects_json must be JSON-serialisable")


CAT4_DOMAIN_SUBJECTS: dict[str, list[str]] = {
    "verbal_domain_score": ["english", "irish", "french", "spanish", "history", "geography", "business_studies"],
    "quantitative_domain_score": ["mathematics", "science", "business_studies", "graphics"],
    "non_verbal_domain_score": ["science", "graphics", "geography", "visual_art"],
    "spatial_domain_score": ["graphics", "geography", "visual_art", "home_economics", "science"],
}

CAT4_SUBJECT_ALIASES: dict[str, str] = {
    "maths": "mathematics",
    "home ec": "home_economics",
    "home_ec": "home_economics",
    "business": "business_studies",
    "art": "visual_art",
    "name": "student_name",
}

CAT4_IGNORED_SCORE_VALUES = {"", "-", "n/a", "na"}


def _normalise_subject_key(value: str) -> str:
    key = _normalise_student_name(value).replace("/", " ").replace("-", " ").replace(".", " ")
    key = re.sub(r"\s+", "_", key.strip())
    return CAT4_SUBJECT_ALIASES.get(key, key)


def _parse_cat4_subject_scores(raw_subjects_json: Optional[str]) -> dict[str, int]:
    if not raw_subjects_json:
        return {}
    try:
        parsed = json.loads(raw_subjects_json)
    except Exception:
        return {}
    if not isinstance(parsed, dict):
        return {}

    scores: dict[str, int] = {}
    for raw_key, raw_value in parsed.items():
        key = _normalise_subject_key(str(raw_key))
        try:
            number = int(round(float(raw_value)))
        except Exception:
            continue
        scores[key] = number
    return scores


def _calculate_cat4_term_metrics(
    raw_subjects_json: Optional[str],
    fallback_average: Optional[int] = None,
    fallback_subject_count: Optional[int] = None,
) -> tuple[Optional[int], Optional[int], dict[str, Optional[int]], Optional[str]]:
    subject_scores = _parse_cat4_subject_scores(raw_subjects_json)
    values = list(subject_scores.values())
    average_percent = int(round(sum(values) / len(values))) if values else fallback_average
    subject_count = len(values) if values else fallback_subject_count

    domain_scores: dict[str, Optional[int]] = {}
    for domain_key, subjects in CAT4_DOMAIN_SUBJECTS.items():
        present = [subject_scores[subject] for subject in subjects if subject in subject_scores]
        domain_scores[domain_key] = int(round(sum(present) / len(present))) if present else None

    raw_json = _json_text_or_none(subject_scores) if subject_scores else _json_text_or_none(raw_subjects_json)
    return average_percent, subject_count, domain_scores, raw_json


def _xlsx_col_index(cell_ref: str) -> int:
    letters = "".join(ch for ch in (cell_ref or "") if ch.isalpha()).upper()
    index = 0
    for char in letters:
        index = index * 26 + (ord(char) - 64)
    return max(index - 1, 0)


def _xlsx_read_sheet(zf: zipfile.ZipFile, path: str, shared_strings: list[str]) -> list[list[str]]:
    ns = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    root = ET.fromstring(zf.read(path))
    rows: list[list[str]] = []
    for row in root.findall(".//x:sheetData/x:row", ns):
        cells: dict[int, str] = {}
        max_index = -1
        for cell in row.findall("x:c", ns):
            ref = cell.attrib.get("r", "")
            idx = _xlsx_col_index(ref)
            max_index = max(max_index, idx)
            cell_type = cell.attrib.get("t")
            value = ""
            if cell_type == "inlineStr":
                text_node = cell.find("x:is/x:t", ns)
                value = text_node.text if text_node is not None and text_node.text is not None else ""
            else:
                value_node = cell.find("x:v", ns)
                raw = value_node.text if value_node is not None and value_node.text is not None else ""
                if cell_type == "s":
                    try:
                        value = shared_strings[int(raw)]
                    except Exception:
                        value = raw
                else:
                    value = raw
            cells[idx] = str(value).strip()
        if max_index < 0:
            continue
        rows.append([cells.get(i, "").strip() for i in range(max_index + 1)])
    return rows


def _read_xlsx_workbook(file_bytes: bytes) -> dict[str, list[list[str]]]:
    ns_main = {"x": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
    ns_rel = {"r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships"}
    rel_ns = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}

    with zipfile.ZipFile(BytesIO(file_bytes)) as zf:
        shared_strings: list[str] = []
        if "xl/sharedStrings.xml" in zf.namelist():
            shared_root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
            for item in shared_root.findall(".//x:si", ns_main):
                text_parts = [node.text or "" for node in item.findall(".//x:t", ns_main)]
                shared_strings.append("".join(text_parts))

        workbook_root = ET.fromstring(zf.read("xl/workbook.xml"))
        rel_root = ET.fromstring(zf.read("xl/_rels/workbook.xml.rels"))
        rels = {
            rel.attrib.get("Id"): f"xl/{rel.attrib.get('Target', '').lstrip('/')}"
            for rel in rel_root.findall("rel:Relationship", rel_ns)
        }

        sheets: dict[str, list[list[str]]] = {}
        for sheet in workbook_root.findall("x:sheets/x:sheet", ns_main):
            name = sheet.attrib.get("name", "Sheet")
            rid = sheet.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id")
            target = rels.get(rid)
            if not target or target not in zf.namelist():
                continue
            sheets[name] = _xlsx_read_sheet(zf, target, shared_strings)
        return sheets


def _sheet_rows_to_dicts(rows: list[list[str]]) -> tuple[list[str], list[dict[str, str]]]:
    if not rows:
        return [], []
    headers = [_normalise_subject_key(cell) for cell in rows[0]]
    records: list[dict[str, str]] = []
    for row in rows[1:]:
        if not any((cell or "").strip() for cell in row):
            continue
        record: dict[str, str] = {}
        for idx, header in enumerate(headers):
            if not header:
                continue
            record[header] = (row[idx] if idx < len(row) else "").strip()
        records.append(record)
    return headers, records


def _validate_sheet_names(
    sheet_name: str,
    names: list[str],
    canonical_names: list[str],
    errors: list[str],
    warnings: list[str],
):
    clean_names = [_normalise_student_name(name) for name in names if _normalise_student_name(name)]
    if len(clean_names) != len(names):
        errors.append(f"{sheet_name}: blank student names found")
    if len(set(clean_names)) != len(clean_names):
        errors.append(f"{sheet_name}: duplicate student names found")
    if canonical_names:
        unknown = sorted(set(clean_names) - set(canonical_names))
        missing = sorted(set(canonical_names) - set(clean_names))
        if unknown:
            errors.append(f"{sheet_name}: unknown names {', '.join(unknown[:5])}")
        if missing:
            errors.append(f"{sheet_name}: missing names {', '.join(missing[:5])}")
        ordered = [name for name in clean_names if name in canonical_names]
        if ordered and ordered != canonical_names[: len(ordered)]:
            warnings.append(f"{sheet_name}: row order differs from cohort order")


def _build_cat4_workbook_preview(sheets: dict[str, list[list[str]]]) -> dict[str, Any]:
    errors: list[str] = []
    warnings: list[str] = []

    baseline_sheet_name = next((name for name in sheets if "cat4" in _normalise_student_name(name) or "baseline" in _normalise_student_name(name)), None)
    cohort_sheet_name = next((name for name in sheets if "cohort" in _normalise_student_name(name) or "name" in _normalise_student_name(name)), None)
    term_sheet_names = [name for name in sheets.keys() if name not in {baseline_sheet_name, cohort_sheet_name}]

    cohort_names: list[str] = []
    if cohort_sheet_name:
        _, cohort_records = _sheet_rows_to_dicts(sheets[cohort_sheet_name])
        cohort_names = [_normalise_student_name(record.get("student_name") or record.get("name") or "") for record in cohort_records if (record.get("student_name") or record.get("name") or "").strip()]

    baseline_rows: list[dict[str, Any]] = []
    if baseline_sheet_name:
        _, baseline_records = _sheet_rows_to_dicts(sheets[baseline_sheet_name])
        raw_names = [record.get("student_name") or "" for record in baseline_records]
        _validate_sheet_names(baseline_sheet_name, raw_names, cohort_names, errors, warnings)
        for record in baseline_records:
            row: dict[str, Any] = {
                "raw_name": (record.get("student_name") or "").strip(),
                "profile_label": (record.get("profile_label") or "").strip() or None,
                "confidence_note": (record.get("note") or "").strip() or None,
            }
            for field in ["verbal_sas", "quantitative_sas", "non_verbal_sas", "spatial_sas", "overall_sas"]:
                raw_value = (record.get(field) or "").strip()
                if not raw_value:
                    row[field] = None
                    continue
                try:
                    row[field] = int(round(float(raw_value)))
                except Exception:
                    row[field] = None
                    errors.append(f"{baseline_sheet_name}: non-numeric CAT4 value for {row['raw_name']} in {field}")
            baseline_rows.append(
                row
            )
        for row in baseline_rows:
            missing = [field for field in ["verbal_sas", "quantitative_sas", "non_verbal_sas", "spatial_sas", "overall_sas"] if row.get(field) is None]
            if missing:
                errors.append(f"{baseline_sheet_name}: {row['raw_name']} missing {', '.join(missing)}")

    term_sets: list[dict[str, Any]] = []
    canonical = cohort_names or [_normalise_student_name(row["raw_name"]) for row in baseline_rows if row.get("raw_name")]
    for sheet_name in term_sheet_names:
        headers, records = _sheet_rows_to_dicts(sheets[sheet_name])
        if not records:
            continue
        raw_names = [record.get("student_name") or "" for record in records]
        _validate_sheet_names(sheet_name, raw_names, canonical, errors, warnings)
        rows_payload: list[dict[str, Any]] = []
        for record in records:
            subject_scores: dict[str, int] = {}
            for header in headers:
                if header in {"student_name", "academic_year", "term_key", "average"}:
                    continue
                raw_value = (record.get(header) or "").strip()
                if raw_value.lower() in CAT4_IGNORED_SCORE_VALUES:
                    continue
                if not raw_value:
                    continue
                try:
                    subject_scores[header] = int(round(float(raw_value)))
                except Exception:
                    errors.append(f"{sheet_name}: non-numeric score for {(record.get('student_name') or '').strip()} in {header}")
            average_percent, subject_count, _, raw_json = _calculate_cat4_term_metrics(json.dumps(subject_scores), None, None)
            rows_payload.append(
                {
                    "raw_name": (record.get("student_name") or "").strip(),
                    "average_percent": average_percent,
                    "subject_count": subject_count,
                    "raw_subjects_json": raw_json,
                }
            )
        term_sets.append(
            {
                "title": sheet_name,
                "academic_year": records[0].get("academic_year") or None,
                "term_key": records[0].get("term_key") or None,
                "rows": rows_payload,
            }
        )

    return {
        "baseline_sheet_name": baseline_sheet_name,
        "cohort_sheet_name": cohort_sheet_name,
        "term_sheet_names": term_sheet_names,
        "baseline_rows": baseline_rows,
        "term_sets": term_sets,
        "errors": errors,
        "warnings": warnings,
    }


def _replace_cat4_data_from_workbook_payload(class_id: int, payload: dict[str, Any], db: Session) -> None:
    db.query(Cat4StudentBaselineModel).filter(Cat4StudentBaselineModel.class_id == class_id).delete()
    db.query(Cat4StudentTermResultModel).filter(Cat4StudentTermResultModel.class_id == class_id).delete()
    db.query(Cat4BaselineSetModel).filter(Cat4BaselineSetModel.class_id == class_id).delete()
    db.query(Cat4TermResultSetModel).filter(Cat4TermResultSetModel.class_id == class_id).delete()

    baseline_rows = payload.get("baseline_rows") or []
    if baseline_rows:
        baseline_set = Cat4BaselineSetModel(
            class_id=class_id,
            title=(payload.get("baseline_sheet_name") or "CAT4 Baseline").strip() or "CAT4 Baseline",
            is_locked=True,
            locked_at=datetime.utcnow(),
        )
        db.add(baseline_set)
        db.flush()

        students = (
            db.query(StudentModel)
            .filter(StudentModel.class_id == class_id)
            .filter(StudentModel.active == True)  # noqa: E712
            .all()
        )
        name_index = _cat4_name_index(students)

        for row in baseline_rows:
            student_id, matched_name, match_note = _match_cat4_student_name(str(row.get("raw_name") or ""), name_index)
            db.add(
                Cat4StudentBaselineModel(
                    baseline_set_id=baseline_set.id,
                    class_id=class_id,
                    student_id=student_id,
                    raw_name=str(row.get("raw_name") or "").strip(),
                    matched_name=matched_name,
                    verbal_sas=row.get("verbal_sas"),
                    quantitative_sas=row.get("quantitative_sas"),
                    non_verbal_sas=row.get("non_verbal_sas"),
                    spatial_sas=row.get("spatial_sas"),
                    overall_sas=row.get("overall_sas"),
                    profile_label=row.get("profile_label"),
                    confidence_note=row.get("confidence_note") or match_note,
                )
            )

    students = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id)
        .filter(StudentModel.active == True)  # noqa: E712
        .all()
    )
    name_index = _cat4_name_index(students)

    for term_set_payload in payload.get("term_sets") or []:
        term_set = Cat4TermResultSetModel(
            class_id=class_id,
            title=(term_set_payload.get("title") or "Term Results").strip() or "Term Results",
            academic_year=term_set_payload.get("academic_year"),
            term_key=term_set_payload.get("term_key"),
        )
        db.add(term_set)
        db.flush()

        for row in term_set_payload.get("rows") or []:
            raw_name = str(row.get("raw_name") or "").strip()
            if not raw_name:
                continue
            student_id, matched_name, _ = _match_cat4_student_name(raw_name, name_index)
            average_percent, subject_count, domain_scores, raw_json = _calculate_cat4_term_metrics(
                _json_text_or_none(row.get("raw_subjects_json")),
                row.get("average_percent"),
                row.get("subject_count"),
            )
            db.add(
                Cat4StudentTermResultModel(
                    result_set_id=term_set.id,
                    class_id=class_id,
                    student_id=student_id,
                    raw_name=raw_name,
                    matched_name=matched_name,
                    average_percent=average_percent,
                    subject_count=subject_count,
                    raw_subjects_json=raw_json,
                    verbal_domain_score=domain_scores.get("verbal_domain_score"),
                    quantitative_domain_score=domain_scores.get("quantitative_domain_score"),
                    non_verbal_domain_score=domain_scores.get("non_verbal_domain_score"),
                    spatial_domain_score=domain_scores.get("spatial_domain_score"),
                )
            )


def _build_percentile_map(pairs: list[tuple[int, float]]) -> dict[int, float]:
    if not pairs:
        return {}

    ordered = sorted(pairs, key=lambda item: (item[1], item[0]))
    result: dict[int, float] = {}
    n = len(ordered)
    i = 0
    while i < n:
        j = i
        value = ordered[i][1]
        while j + 1 < n and ordered[j + 1][1] == value:
            j += 1
        avg_rank = ((i + 1) + (j + 1)) / 2
        percentile = round((avg_rank / n) * 100, 1)
        for idx in range(i, j + 1):
            result[ordered[idx][0]] = percentile
        i = j + 1
    return result


def _cat4_status_for_student(
    latest_average: Optional[float],
    trend_delta: Optional[float],
    value_added_delta: Optional[float],
) -> tuple[str, list[str]]:
    reasons: list[str] = []

    if latest_average is not None and latest_average < 55:
        reasons.append("Latest average below 55%")
    if trend_delta is not None and trend_delta <= -7:
        reasons.append("Trend dropped by 7 points or more")
    if value_added_delta is not None and value_added_delta <= -20:
        reasons.append("Working well below CAT4 baseline percentile")

    if latest_average is not None and latest_average >= 80 and value_added_delta is not None and value_added_delta >= 15:
        return "excelling", ["Latest average is 80%+", "Performing well above CAT4 baseline percentile"]

    if reasons:
        return "at_risk", reasons

    improving_reasons: list[str] = []
    if trend_delta is not None and trend_delta >= 5:
        improving_reasons.append("Trend improved by 5 points or more")
    if value_added_delta is not None and value_added_delta >= 10:
        improving_reasons.append("Performing above CAT4 baseline percentile")
    if improving_reasons:
        return "improving", improving_reasons

    return "stable", ["Broadly on track"]


def _cat4_set_summary(rows: list[Any]) -> dict[str, int]:
    matched = sum(1 for row in rows if getattr(row, "student_id", None))
    return {
        "row_count": len(rows),
        "matched_count": matched,
        "unmatched_count": len(rows) - matched,
    }


# =========================================================
# WAITLIST (Elume early access)
# =========================================================

from pydantic import EmailStr
import smtplib
from email.mime.text import MIMEText


class WaitlistRequest(BaseModel):
    name: str
    email: EmailStr
    school: str | None = None


@app.post("/waitlist")
def join_waitlist(payload: WaitlistRequest):

    message = f"""
New Elume waitlist signup

Name: {payload.name}
Email: {payload.email}
School: {payload.school or "Not provided"}
"""

    msg = MIMEText(message)
    msg["Subject"] = "New Elume Waitlist Signup"
    msg["From"] = "admin@elume.ie"
    msg["To"] = "admin@elume.ie"

    try:
        with smtplib.SMTP("smtp.zoho.eu", 587) as server:
            server.starttls()
            server.login("admin@elume.ie", os.getenv("EMAIL_PASSWORD"))
            server.send_message(msg)

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Email failed: {str(e)}")

    return {"success": True}

# =========================================================
# TEACHER ADMIN STATE (Profile + Timetable) — synced across devices
# =========================================================
from sqlalchemy.exc import IntegrityError  # add near imports if you prefer

@app.get("/teacher-admin/state", response_model=schemas.TeacherAdminStateOut)
def get_teacher_admin_state(
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    row = (
        db.query(models.TeacherAdminStateModel)
        .filter(models.TeacherAdminStateModel.owner_user_id == user.id)
        .first()
    )
    if not row:
        return {"state": {}, "updated_at": None}

    try:
        parsed = json.loads(row.state_json or "{}")
    except Exception:
        parsed = {}

    return {"state": parsed, "updated_at": row.updated_at}


@app.put("/teacher-admin/state", response_model=schemas.TeacherAdminStateOut)
def save_teacher_admin_state(
    payload: schemas.TeacherAdminStateSave,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    # validate it is JSON-serialisable
    try:
        raw = json.dumps(payload.state, ensure_ascii=False)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"State must be JSON-serialisable: {e}")

    row = (
        db.query(models.TeacherAdminStateModel)
        .filter(models.TeacherAdminStateModel.owner_user_id == user.id)
        .first()
    )

    now = datetime.utcnow()

    if row:
        row.state_json = raw
        row.updated_at = now
        db.commit()
        db.refresh(row)
    else:
        row = models.TeacherAdminStateModel(
            owner_user_id=user.id,
            state_json=raw,
            updated_at=now,
        )
        db.add(row)
        try:
            db.commit()
        except IntegrityError:
            # rare race: if two devices create simultaneously
            db.rollback()
            row = (
                db.query(models.TeacherAdminStateModel)
                .filter(models.TeacherAdminStateModel.owner_user_id == user.id)
                .first()
            )
            if row:
                row.state_json = raw
                row.updated_at = now
                db.commit()
                db.refresh(row)
            else:
                raise
        else:
            db.refresh(row)

    return {"state": json.loads(row.state_json), "updated_at": row.updated_at}

@app.get("/admin/users")
def admin_list_users(
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    require_super_admin(user)

    users = (
        db.query(models.UserModel)
        .order_by(models.UserModel.id.asc())
        .all()
    )

    return [
        {
            "id": u.id,
            "email": u.email,
            "created_at": u.created_at.isoformat() if getattr(u, "created_at", None) else None,
            "subscription_status": (u.subscription_status or "inactive"),
            "billing_interval": u.billing_interval,
            "current_period_end": u.current_period_end.isoformat() if getattr(u, "current_period_end", None) else None,
        }
        for u in users
    ]


@app.post("/admin/users")
def admin_create_user(
    payload: AdminCreateUser,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    require_super_admin(user)

    email = (payload.email or "").strip().lower()
    password = (payload.password or "").strip()

    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Valid email required")
    password_error = _password_policy_error(password)
    if password_error:
        raise HTTPException(status_code=400, detail=password_error)

    existing = db.query(models.UserModel).filter(models.UserModel.email == email).first()
    if existing:
        raise HTTPException(status_code=400, detail="Email already exists")

    new_user = models.UserModel(
        email=email,
        password_hash=PWD_CONTEXT.hash(password),
        email_verified=True,
    )
    try:
        db.add(new_user)
        db.flush()
        _seed_demo_class(db, new_user)
        db.refresh(new_user)
    except HTTPException:
        db.rollback()
        raise
    except Exception:
        db.rollback()
        logger.exception("Failed to provision demo class during admin user creation for %s", email)
        raise HTTPException(status_code=500, detail="Failed to finish user setup")

    return {
        "message": "User created",
        "id": new_user.id,
        "email": new_user.email,
    }


@app.post("/admin/users/reset-password")
def admin_reset_password(
    payload: AdminResetPassword,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    require_super_admin(user)

    email = (payload.email or "").strip().lower()
    new_password = (payload.new_password or "").strip()

    if not email:
        raise HTTPException(status_code=400, detail="Email required")
    password_error = _password_policy_error(new_password)
    if password_error:
        raise HTTPException(status_code=400, detail=password_error)

    target = db.query(models.UserModel).filter(models.UserModel.email == email).first()
    if not target:
        raise HTTPException(status_code=404, detail="User not found")

    target.password_hash = PWD_CONTEXT.hash(new_password)
    db.commit()

    return {"message": f"Password reset for {target.email}"}


@app.post("/admin/users/rename")
def admin_rename_user(
    payload: AdminRenameUser,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    require_super_admin(user)

    old_email = (payload.old_email or "").strip().lower()
    new_email = (payload.new_email or "").strip().lower()

    if not old_email or not new_email or "@" not in new_email:
        raise HTTPException(status_code=400, detail="Valid old and new email required")

    target = db.query(models.UserModel).filter(models.UserModel.email == old_email).first()
    if not target:
        raise HTTPException(status_code=404, detail="User not found")

    clash = db.query(models.UserModel).filter(models.UserModel.email == new_email).first()
    if clash and clash.id != target.id:
        raise HTTPException(status_code=400, detail="New email already in use")

    target.email = new_email
    db.commit()
    db.refresh(target)

    return {
        "message": "User email updated",
        "id": target.id,
        "email": target.email,
    }


@app.delete("/admin/users")
def admin_delete_user(
    payload: AdminDeleteUser,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    require_super_admin(user)

    email = (payload.email or "").strip().lower()
    hard_delete = bool(payload.hard_delete)
    if not email or "@" not in email:
        raise HTTPException(status_code=400, detail="Valid email required")
    if email == "admin@elume.ie":
        raise HTTPException(status_code=400, detail="Cannot delete admin@elume.ie")
    if email == (user.email or "").strip().lower():
        raise HTTPException(status_code=400, detail="Cannot delete the currently logged in super admin account")

    target = db.query(models.UserModel).filter(models.UserModel.email == email).first()
    if not target:
        raise HTTPException(status_code=404, detail="User not found")

    owned_classes = (
        db.query(ClassModel)
        .filter(ClassModel.owner_user_id == target.id)
        .order_by(ClassModel.id.asc())
        .all()
    )
    if owned_classes and not hard_delete:
        raise HTTPException(
            status_code=409,
            detail={
                "detail": "User still owns classes.",
                "code": "USER_OWNS_CLASSES",
                "class_count": len(owned_classes),
                "classes": [
                    {"id": cls.id, "name": cls.name, "subject": cls.subject}
                    for cls in owned_classes
                ],
            },
        )

    class_ids = [cls.id for cls in owned_classes]

    try:
        if class_ids:
            _delete_class_dependencies(db, class_ids, owned_classes)

        db.query(models.TeacherAdminStateModel).filter(
            models.TeacherAdminStateModel.owner_user_id == target.id
        ).delete(synchronize_session=False)
        db.query(models.CalendarEvent).filter(
            models.CalendarEvent.owner_user_id == target.id
        ).delete(synchronize_session=False)
        db.query(models.WhiteboardStateModel).filter(
            models.WhiteboardStateModel.owner_user_id == target.id
        ).delete(synchronize_session=False)
        remaining_quizzes = db.query(models.SavedQuizModel).filter(
            models.SavedQuizModel.owner_user_id == target.id
        ).all()
        for quiz in remaining_quizzes:
            db.delete(quiz)
        db.query(models.PasswordResetTokenModel).filter(
            models.PasswordResetTokenModel.user_id == target.id
        ).delete(synchronize_session=False)
        db.query(models.EmailVerificationTokenModel).filter(
            models.EmailVerificationTokenModel.user_id == target.id
        ).delete(synchronize_session=False)

        target_email = target.email
        db.delete(target)
        db.commit()
    except HTTPException:
        db.rollback()
        raise
    except Exception:
        db.rollback()
        logger.exception("Failed to delete user %s", email)
        raise HTTPException(status_code=500, detail="Failed to delete user")

    return {
        "success": True,
        "message": f"Deleted user {target_email}",
        "hard_deleted": hard_delete,
    }


def _delete_class_dependencies(
    db: Session,
    class_ids: list[int],
    class_rows: Optional[list[ClassModel]] = None,
) -> None:
    if not class_ids:
        return

    assessment_ids = [
        row.id
        for row in db.query(ClassAssessmentModel.id)
        .filter(ClassAssessmentModel.class_id.in_(class_ids))
        .all()
    ]
    if assessment_ids:
        db.query(AssessmentResultModel).filter(
            AssessmentResultModel.assessment_id.in_(assessment_ids)
        ).delete(synchronize_session=False)
        db.query(ClassAssessmentModel).filter(
            ClassAssessmentModel.id.in_(assessment_ids)
        ).delete(synchronize_session=False)

    baseline_set_ids = [
        row.id
        for row in db.query(Cat4BaselineSetModel.id)
        .filter(Cat4BaselineSetModel.class_id.in_(class_ids))
        .all()
    ]
    if baseline_set_ids:
        db.query(Cat4StudentBaselineModel).filter(
            Cat4StudentBaselineModel.baseline_set_id.in_(baseline_set_ids)
        ).delete(synchronize_session=False)
        db.query(Cat4BaselineSetModel).filter(
            Cat4BaselineSetModel.id.in_(baseline_set_ids)
        ).delete(synchronize_session=False)

    term_set_ids = [
        row.id
        for row in db.query(Cat4TermResultSetModel.id)
        .filter(Cat4TermResultSetModel.class_id.in_(class_ids))
        .all()
    ]
    if term_set_ids:
        db.query(Cat4StudentTermResultModel).filter(
            Cat4StudentTermResultModel.result_set_id.in_(term_set_ids)
        ).delete(synchronize_session=False)
        db.query(Cat4TermResultSetModel).filter(
            Cat4TermResultSetModel.id.in_(term_set_ids)
        ).delete(synchronize_session=False)

    db.query(Cat4WorkbookVersionModel).filter(
        Cat4WorkbookVersionModel.class_id.in_(class_ids)
    ).delete(synchronize_session=False)

    db.query(models.Note).filter(
        models.Note.class_id.in_(class_ids)
    ).delete(synchronize_session=False)
    db.query(models.Topic).filter(
        models.Topic.class_id.in_(class_ids)
    ).delete(synchronize_session=False)

    db.query(models.TestItem).filter(
        models.TestItem.class_id.in_(class_ids)
    ).delete(synchronize_session=False)
    db.query(models.TestCategory).filter(
        models.TestCategory.class_id.in_(class_ids)
    ).delete(synchronize_session=False)

    db.query(models.StudentAccessLink).filter(
        models.StudentAccessLink.class_id.in_(class_ids)
    ).delete(synchronize_session=False)
    db.query(models.CalendarEvent).filter(
        models.CalendarEvent.class_id.in_(class_ids)
    ).delete(synchronize_session=False)
    db.query(models.WhiteboardStateModel).filter(
        models.WhiteboardStateModel.class_id.in_(class_ids)
    ).delete(synchronize_session=False)

    live_session_ids = [
        row.id
        for row in db.query(models.LiveQuizSessionModel.id)
        .filter(models.LiveQuizSessionModel.class_id.in_(class_ids))
        .all()
    ]
    if live_session_ids:
        db.query(models.LiveQuizAttemptModel).filter(
            models.LiveQuizAttemptModel.session_id.in_(live_session_ids)
        ).delete(synchronize_session=False)
        participant_ids = [
            row.id
            for row in db.query(models.LiveQuizParticipantModel.id)
            .filter(models.LiveQuizParticipantModel.session_id.in_(live_session_ids))
            .all()
        ]
        db.query(models.LiveQuizAnswerModel).filter(
            models.LiveQuizAnswerModel.session_id.in_(live_session_ids)
        ).delete(synchronize_session=False)
        if participant_ids:
            db.query(models.LiveQuizAnswerModel).filter(
                models.LiveQuizAnswerModel.participant_id.in_(participant_ids)
            ).delete(synchronize_session=False)
        db.query(models.LiveQuizParticipantModel).filter(
            models.LiveQuizParticipantModel.session_id.in_(live_session_ids)
        ).delete(synchronize_session=False)
        db.query(models.LiveQuizSessionModel).filter(
            models.LiveQuizSessionModel.id.in_(live_session_ids)
        ).delete(synchronize_session=False)

    db.query(models.LiveQuizAttemptModel).filter(
        models.LiveQuizAttemptModel.class_id.in_(class_ids)
    ).delete(synchronize_session=False)

    collab_session_ids = [
        row.id
        for row in db.query(models.CollabSessionModel.id)
        .filter(models.CollabSessionModel.class_id.in_(class_ids))
        .all()
    ]
    if collab_session_ids:
        db.query(models.CollabParticipantModel).filter(
            models.CollabParticipantModel.session_id.in_(collab_session_ids)
        ).delete(synchronize_session=False)
        db.query(models.CollabSessionModel).filter(
            models.CollabSessionModel.id.in_(collab_session_ids)
        ).delete(synchronize_session=False)

    class_quizzes = db.query(models.SavedQuizModel).filter(
        models.SavedQuizModel.class_id.in_(class_ids)
    ).all()
    for quiz in class_quizzes:
        db.delete(quiz)

    if class_rows:
        for cls in class_rows:
            db.delete(cls)


@app.post("/admin/classes/transfer")
def admin_transfer_class(
    payload: AdminTransferClassPayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    require_super_admin(user)

    target_email = (payload.target_email or "").strip().lower()
    if not target_email or "@" not in target_email:
        raise HTTPException(status_code=400, detail="Valid target email required")

    target_user = db.query(models.UserModel).filter(models.UserModel.email == target_email).first()
    if not target_user:
        raise HTTPException(status_code=404, detail="Target user not found")

    cls = db.query(ClassModel).filter(ClassModel.id == payload.class_id).first()
    if not cls:
        raise HTTPException(status_code=404, detail="Class not found")

    if cls.owner_user_id == target_user.id:
        raise HTTPException(status_code=400, detail="Class already belongs to target user")

    conflicting_class = (
        db.query(ClassModel)
        .filter(
            ClassModel.owner_user_id == target_user.id,
            ClassModel.id != cls.id,
            func.lower(ClassModel.name) == (cls.name or "").strip().lower(),
        )
        .first()
    )
    if conflicting_class:
        raise HTTPException(
            status_code=400,
            detail=f"Target user already owns a class named '{cls.name}'",
        )

    old_owner = (
        db.query(models.UserModel)
        .filter(models.UserModel.id == cls.owner_user_id)
        .first()
        if cls.owner_user_id
        else None
    )

    cls.owner_user_id = target_user.id
    db.commit()
    db.refresh(cls)

    return {
        "success": True,
        "class_id": cls.id,
        "class_name": cls.name,
        "old_owner_user_id": old_owner.id if old_owner else None,
        "old_owner_email": old_owner.email if old_owner else None,
        "new_owner_user_id": target_user.id,
        "new_owner_email": target_user.email,
    }

def _send_email(to_email: str, subject: str, body: str):
    msg = MIMEText(body)
    msg["Subject"] = subject
    smtp_host = (os.getenv("SMTP_HOST") or "smtp.zoho.eu").strip()
    smtp_port = int((os.getenv("SMTP_PORT") or "587").strip() or "587")
    smtp_user = (os.getenv("SMTP_USER") or "admin@elume.ie").strip()
    smtp_password = (os.getenv("SMTP_PASSWORD") or os.getenv("EMAIL_PASSWORD") or "").strip()
    smtp_from = (os.getenv("SMTP_FROM_EMAIL") or smtp_user or "admin@elume.ie").strip()
    smtp_use_tls = (os.getenv("SMTP_USE_TLS") or "true").strip().lower() in {"1", "true", "yes", "on"}

    if not smtp_host or not smtp_user or not smtp_password:
        raise RuntimeError("SMTP is not configured")

    msg["From"] = smtp_from
    msg["To"] = to_email

    with smtplib.SMTP(smtp_host, smtp_port) as server:
        if smtp_use_tls:
            server.starttls()
        server.login(smtp_user, smtp_password)
        server.send_message(msg)


def _hash_reset_token(raw_token: str) -> str:
    return hashlib.sha256((raw_token or "").encode("utf-8")).hexdigest()


def _hash_email_verification_token(raw_token: str) -> str:
    return hashlib.sha256((raw_token or "").encode("utf-8")).hexdigest()


# =========================================================
# ADMIN (Students + Assessments/Results)
# =========================================================

class StudentCreate(BaseModel):
    first_name: str
    notes: Optional[str] = None

class StudentBulkCreate(BaseModel):
    names: List[str]
    notes: Optional[str] = None

class StudentUpdate(BaseModel):
    first_name: Optional[str] = None
    notes: Optional[str] = None
    active: Optional[bool] = None

class AssessmentCreate(BaseModel):
    title: str
    # optional YYYY-MM-DD; if omitted use today
    assessment_date: Optional[str] = None

class AssessmentUpdate(BaseModel):
    title: str
    assessment_date: Optional[str] = None

class ResultUpsert(BaseModel):
    student_id: int
    score_percent: Optional[int] = None
    absent: bool = False

class BulkResultsUpdate(BaseModel):
    results: List[ResultUpsert]

# =========================================================
# LIVE QUIZ / POLL
# =========================================================


class LiveQuizCreateQuestion(BaseModel):
    id: str
    prompt: str
    choices: dict
    correct: Optional[str] = None  # "A"|"B"|"C"|"D" or None

class LiveQuizCreateRequest(BaseModel):
    class_id: int
    title: str
    anonymous: bool = True
    quiz_id: Optional[str] = None
    seconds_per_question: Optional[int] = 20
    shuffle_questions: bool = False
    auto_play: bool = False
    auto_end_when_all_answered: bool = True
    questions: List[LiveQuizCreateQuestion]

class LiveQuizCreateResponse(BaseModel):
    session_code: str
    join_url: Optional[str] = None

class LiveQuizJoinRequest(BaseModel):
    anon_id: Optional[str] = None  # allow client to send same anon_id again (optional)
    name: Optional[str] = None     # ✅ student-entered name

class LiveQuizJoinResponse(BaseModel):
    anon_id: str
    nickname: Optional[str] = None

class LiveQuizAnswerRequest(BaseModel):
    anon_id: str
    question_id: str
    choice: str  # A/B/C/D

class LiveQuizAttemptExcludeRequest(BaseModel):
    excluded: bool

class CollabUpdatePayload(BaseModel):
    room_count: Optional[int] = None
    timer_minutes: Optional[int] = None

class TeacherPlannerStateModel(Base):
    __tablename__ = "teacher_planner_state"

    id = Column(Integer, primary_key=True, index=True)
    teacher_id = Column(Integer, unique=True, index=True)
    state_json = Column(Text, default='{"notes":[],"tasks":[]}')
    updated_at = Column(Text, nullable=True)

@app.post("/collab/{code}/config")
def collab_update_config(
    code: str,
    payload: CollabUpdatePayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_collab_session_or_404(code, db, user)

    if s.state == "live":
        raise HTTPException(status_code=400, detail="Cannot change config after breakout has started")

    if payload.room_count is not None:
        s.room_count = max(1, min(12, int(payload.room_count)))

        participants = (
            db.query(CollabParticipantModel)
            .filter(CollabParticipantModel.session_id == s.id)
            .all()
        )
        for p in participants:
            if p.room_number is not None and p.room_number > s.room_count:
                p.room_number = s.room_count

    if payload.timer_minutes is not None:
        s.timer_minutes = max(1, min(60, int(payload.timer_minutes)))

    db.commit()
    db.refresh(s)

    return {
        "session_code": s.session_code,
        "room_count": s.room_count,
        "timer_minutes": s.timer_minutes,
        "state": s.state,
    }

class CollabRoomManager:
    def __init__(self):
        # key = (session_code, room_key)
        self.rooms: dict[tuple[str, str], list[WebSocket]] = defaultdict(list)

    async def connect(self, session_code: str, room_key: str, websocket: WebSocket):
        await websocket.accept()
        key = (session_code, room_key)
        self.rooms[key].append(websocket)

    def disconnect(self, session_code: str, room_key: str, websocket: WebSocket):
        key = (session_code, room_key)
        if key in self.rooms and websocket in self.rooms[key]:
            self.rooms[key].remove(websocket)
        if key in self.rooms and not self.rooms[key]:
            del self.rooms[key]

    async def broadcast(self, session_code: str, room_key: str, payload: dict):
        key = (session_code, room_key)
        dead = []
        for ws in self.rooms.get(key, []):
            try:
                await ws.send_json(payload)
            except Exception:
                dead.append(ws)

        for ws in dead:
            self.disconnect(session_code, room_key, ws)

collab_room_manager = CollabRoomManager()
CLASS_CODE_ALPHABET = "ABCDEFGHJKMNPQRSTUVWXYZ23456789"

def _collab_room_key(room_number: int | None) -> str:
    if room_number is None:
        return "teacher-main"
    return f"room-{int(room_number)}"

def _rand_code(n: int = 6) -> str:
    return "".join(random.choice(CLASS_CODE_ALPHABET) for _ in range(n))


def _normalise_class_code(value: str) -> str:
    cleaned = re.sub(r"[^A-Z0-9]", "", (value or "").upper())
    return cleaned[:6]


def _rand_class_code(db: Session) -> str:
    for _ in range(20):
        code = "".join(random.choice(CLASS_CODE_ALPHABET) for _ in range(6))
        exists = db.query(ClassModel).filter(ClassModel.class_code == code).first()
        if not exists:
            return code
    raise HTTPException(status_code=500, detail="Could not generate class code")


def _rand_class_pin() -> str:
    return "".join(random.choice(string.digits) for _ in range(4))

def _kid_name() -> str:
    adj = ["Sunny","Rocket","Brave","Clever","Happy","Chill","Mighty","Rapid","Cosmic","Bouncy","Quiet","Zippy"]
    animal = ["Panda","Koala","Tiger","Otter","Fox","Dolphin","Eagle","Turtle","Penguin","Lemur","Rabbit","Seal"]
    return f"{random.choice(adj)} {random.choice(animal)}"

def _load_questions(session: LiveQuizSessionModel) -> list:
    try:
        q = json.loads(session.questions_json or "[]")
        return q if isinstance(q, list) else []
    except Exception:
        return []

def _time_left_seconds(session: LiveQuizSessionModel) -> Optional[int]:
    if session.state != "live":
        return None
    if not session.seconds_per_question:
        return None
    if session.question_closed_at is not None:
        return 0
    if not session.question_started_at:
        return None
    elapsed = (datetime.utcnow() - session.question_started_at).total_seconds()
    left = int(session.seconds_per_question - elapsed)
    return max(0, left)

def _current_question(session: LiveQuizSessionModel) -> Optional[dict]:
    qs = _load_questions(session)

    # IMPORTANT: current_index can be 0, so do NOT use "or -1"
    if session.current_index is None:
        idx = -1
    else:
        try:
            idx = int(session.current_index)
        except Exception:
            idx = -1

    if idx < 0 or idx >= len(qs):
        return None
    return qs[idx]

def _livequiz_timer_expired(session: LiveQuizSessionModel) -> bool:
    if session.state != "live":
        return False
    if not session.seconds_per_question or not session.question_started_at:
        return False
    return datetime.utcnow() >= session.question_started_at + timedelta(seconds=int(session.seconds_per_question))

def _livequiz_answers_open(session: LiveQuizSessionModel) -> bool:
    if session.state != "live":
        return False
    if _current_question(session) is None:
        return False
    if session.question_closed_at is not None:
        return False
    if _livequiz_timer_expired(session):
        return False
    return True

def _normalise_livequiz_participant_name(value: str) -> str:
    return re.sub(r"\s+", " ", (value or "").strip()).lower()

def _match_livequiz_student_for_attempt(db: Session, class_id: int, display_name: str) -> tuple[Optional[int], Optional[str]]:
    key = _normalise_livequiz_participant_name(display_name)
    if not key:
        return None, None

    students = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id)
        .filter(StudentModel.active == True)  # noqa: E712
        .all()
    )
    matches = [s for s in students if _normalise_livequiz_participant_name(s.first_name) == key]
    if len(matches) != 1:
        return None, None
    return matches[0].id, matches[0].first_name

def _sync_livequiz_attempts(db: Session, session: LiveQuizSessionModel) -> None:
    results = _build_livequiz_results(db, session)
    participants = (
        db.query(LiveQuizParticipantModel)
        .filter(LiveQuizParticipantModel.session_id == session.id)
        .all()
    )
    participant_map = {p.id: p for p in participants}

    latest_answer_times: dict[int, datetime] = {}
    for row in (
        db.query(LiveQuizAnswerModel.participant_id, func.max(LiveQuizAnswerModel.answered_at))
        .filter(LiveQuizAnswerModel.session_id == session.id)
        .group_by(LiveQuizAnswerModel.participant_id)
        .all()
    ):
        if row[0]:
            latest_answer_times[int(row[0])] = row[1]

    db.query(LiveQuizAttemptModel).filter(
        LiveQuizAttemptModel.session_id == session.id
    ).delete(synchronize_session=False)

    total_questions = int(results.get("summary", {}).get("total_questions") or 0)
    scored_mode = bool(results.get("summary", {}).get("scored_mode"))

    for row in results.get("leaderboard", []):
        participant_id = row.get("participant_id")
        participant = participant_map.get(int(participant_id)) if participant_id is not None else None
        participant_name = (
            (participant.nickname or "").strip()
            if participant and participant.nickname
            else str(row.get("name") or "Player").strip()
        ) or "Player"
        student_id, matched_name = _match_livequiz_student_for_attempt(db, session.class_id, participant_name)
        answered = int(row.get("answered") or 0)
        correct = int(row.get("correct") or 0)
        score_percent = int(row.get("percent")) if scored_mode and total_questions > 0 else None
        completed = bool(total_questions > 0 and answered >= total_questions)
        excluded = not completed

        db.add(
            LiveQuizAttemptModel(
                class_id=session.class_id,
                session_id=session.id,
                quiz_id=session.quiz_id,
                participant_id=participant.id if participant else None,
                student_id=student_id,
                participant_identifier=participant.anon_id if participant else None,
                participant_display_name=matched_name or participant_name,
                score=correct,
                score_percent=score_percent,
                total_questions=total_questions,
                completed=completed,
                scored_mode=scored_mode,
                excluded_from_average=excluded,
                submitted_at=latest_answer_times.get(participant.id) if participant else None,
                finished_at=session.ended_at or datetime.utcnow(),
            )
        )

def _finalize_livequiz_session(db: Session, session: LiveQuizSessionModel) -> None:
    if session.state != "ended":
        session.state = "ended"
    if session.ended_at is None:
        session.ended_at = datetime.utcnow()
    if session.question_closed_at is None:
        session.question_closed_at = session.ended_at
    db.flush()
    _sync_livequiz_attempts(db, session)

def _maybe_progress_livequiz_session(db: Session, session: LiveQuizSessionModel) -> LiveQuizSessionModel:
    if session.state != "live":
        return session
    if not _livequiz_timer_expired(session):
        return session

    expired_at = (
        session.question_started_at + timedelta(seconds=int(session.seconds_per_question))
        if session.question_started_at and session.seconds_per_question
        else datetime.utcnow()
    )
    if session.question_closed_at is None:
        session.question_closed_at = expired_at

    if session.auto_play:
        qs = _load_questions(session)
        next_index = int(session.current_index) + 1
        if next_index >= len(qs):
            _finalize_livequiz_session(db, session)
        else:
            session.current_index = next_index
            session.question_started_at = datetime.utcnow()
            session.question_closed_at = None

    db.commit()
    db.refresh(session)
    return session

def ensure_columns():
    # Safely add missing columns to existing SQLite tables
    with engine.connect() as conn:
        # -------------------------
        # classes table
        # -------------------------
        cols = conn.execute(text("PRAGMA table_info(classes)")).fetchall()
        col_names = {c[1] for c in cols}

        if "owner_user_id" not in col_names:
            conn.execute(text("ALTER TABLE classes ADD COLUMN owner_user_id INTEGER"))
            conn.commit()
        if "class_code" not in col_names:
            conn.execute(text("ALTER TABLE classes ADD COLUMN class_code TEXT"))
            conn.commit()
        if "class_pin" not in col_names:
            conn.execute(text("ALTER TABLE classes ADD COLUMN class_pin TEXT"))
            conn.commit()
        if "color" not in col_names:
            conn.execute(text("ALTER TABLE classes ADD COLUMN color TEXT"))
            conn.commit()
        if "preferred_exam_subject" not in col_names:
            conn.execute(text("ALTER TABLE classes ADD COLUMN preferred_exam_subject TEXT"))
            conn.commit()
        if "stream" not in col_names:
            conn.execute(text("ALTER TABLE classes ADD COLUMN stream TEXT"))
            conn.commit()
        if "is_archived" not in col_names:
            conn.execute(text("ALTER TABLE classes ADD COLUMN is_archived BOOLEAN DEFAULT 0 NOT NULL"))
            conn.commit()
        if "archived_at" not in col_names:
            conn.execute(text("ALTER TABLE classes ADD COLUMN archived_at DATETIME"))
            conn.commit()

        quiz_cols = conn.execute(text("PRAGMA table_info(saved_quizzes)")).fetchall()
        quiz_col_names = {c[1] for c in quiz_cols}
        if quiz_cols and "is_starred" not in quiz_col_names:
            conn.execute(text("ALTER TABLE saved_quizzes ADD COLUMN is_starred BOOLEAN DEFAULT 0 NOT NULL"))
            conn.commit()

        conn.execute(
            text("CREATE UNIQUE INDEX IF NOT EXISTS ix_classes_class_code ON classes (class_code)")
        )
        conn.commit()

        # -------------------------
        # users table
        # -------------------------
        user_cols = conn.execute(text("PRAGMA table_info(users)")).fetchall()
        user_col_names = {c[1] for c in user_cols}

        if "subscription_status" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN subscription_status TEXT DEFAULT 'inactive' NOT NULL"))
            conn.commit()

        if "billing_interval" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN billing_interval TEXT"))
            conn.commit()

        if "stripe_customer_id" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN stripe_customer_id TEXT"))
            conn.commit()

        if "stripe_subscription_id" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN stripe_subscription_id TEXT"))
            conn.commit()

        if "stripe_checkout_session_id" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN stripe_checkout_session_id TEXT"))
            conn.commit()

        if "subscription_started_at" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN subscription_started_at DATETIME"))
            conn.commit()

        if "current_period_end" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN current_period_end DATETIME"))
            conn.commit()

        if "launch_offer_applied" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN launch_offer_applied BOOLEAN DEFAULT 0 NOT NULL"))
            conn.commit()
        if "billing_onboarding_required" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN billing_onboarding_required BOOLEAN DEFAULT 0 NOT NULL"))
            conn.commit()
        if "first_name" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN first_name TEXT"))
            conn.commit()
        if "last_name" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN last_name TEXT"))
            conn.commit()
        if "school_name" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN school_name TEXT"))
            conn.commit()
        if "email_verified" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN email_verified BOOLEAN DEFAULT 1 NOT NULL"))
            conn.commit()
        if "trial_started_at" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN trial_started_at DATETIME"))
            conn.commit()
        if "trial_ends_at" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN trial_ends_at DATETIME"))
            conn.commit()
        if "ai_daily_limit" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN ai_daily_limit INTEGER DEFAULT 0 NOT NULL"))
            conn.commit()
        if "ai_prompt_count" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN ai_prompt_count INTEGER DEFAULT 0 NOT NULL"))
            conn.commit()
        if "ai_prompt_count_date" not in user_col_names:
            conn.execute(text("ALTER TABLE users ADD COLUMN ai_prompt_count_date DATETIME"))
            conn.commit()

        # -------------------------
        # notes table
        # -------------------------
        note_cols = conn.execute(text("PRAGMA table_info(notes)")).fetchall()
        note_col_names = {c[1] for c in note_cols}

        if "whiteboard_state_id" not in note_col_names:
            conn.execute(text("ALTER TABLE notes ADD COLUMN whiteboard_state_id INTEGER"))
            conn.commit()

        conn.execute(
            text("CREATE INDEX IF NOT EXISTS ix_notes_whiteboard_state_id ON notes (whiteboard_state_id)")
        )
        conn.commit()

        conn.execute(
            text(
                """
                CREATE TABLE IF NOT EXISTS email_verification_tokens (
                    id INTEGER PRIMARY KEY,
                    user_id INTEGER NOT NULL,
                    token_hash TEXT NOT NULL UNIQUE,
                    expires_at DATETIME NOT NULL,
                    used_at DATETIME NULL,
                    created_at DATETIME NOT NULL DEFAULT CURRENT_TIMESTAMP
                )
                """
            )
        )
        conn.execute(
            text("CREATE INDEX IF NOT EXISTS ix_email_verification_tokens_user_id ON email_verification_tokens (user_id)")
        )
        conn.execute(
            text("CREATE UNIQUE INDEX IF NOT EXISTS ix_email_verification_tokens_token_hash ON email_verification_tokens (token_hash)")
        )
        conn.commit()

        # -------------------------
        # cat4_baseline_sets table
        # -------------------------
        cat4_baseline_cols = conn.execute(text("PRAGMA table_info(cat4_baseline_sets)")).fetchall()
        cat4_baseline_col_names = {c[1] for c in cat4_baseline_cols}
        if "is_locked" not in cat4_baseline_col_names and cat4_baseline_cols:
            conn.execute(text("ALTER TABLE cat4_baseline_sets ADD COLUMN is_locked BOOLEAN DEFAULT 0 NOT NULL"))
            conn.commit()
        if "locked_at" not in cat4_baseline_col_names and cat4_baseline_cols:
            conn.execute(text("ALTER TABLE cat4_baseline_sets ADD COLUMN locked_at DATETIME"))
            conn.commit()

        # -------------------------
        # cat4_term_result_sets table
        # -------------------------
        cat4_term_row_cols = conn.execute(text("PRAGMA table_info(cat4_student_term_results)")).fetchall()
        cat4_term_row_col_names = {c[1] for c in cat4_term_row_cols}
        for column_name in [
            "verbal_domain_score",
            "quantitative_domain_score",
            "non_verbal_domain_score",
            "spatial_domain_score",
        ]:
            if column_name not in cat4_term_row_col_names and cat4_term_row_cols:
                conn.execute(text(f"ALTER TABLE cat4_student_term_results ADD COLUMN {column_name} INTEGER"))
                conn.commit()

        # -------------------------
        # livequiz_sessions table
        # -------------------------
        livequiz_cols = conn.execute(text("PRAGMA table_info(livequiz_sessions)")).fetchall()
        livequiz_col_names = {c[1] for c in livequiz_cols}
        if "quiz_id" not in livequiz_col_names and livequiz_cols:
            conn.execute(text("ALTER TABLE livequiz_sessions ADD COLUMN quiz_id TEXT"))
            conn.commit()
        if "auto_play" not in livequiz_col_names and livequiz_cols:
            conn.execute(text("ALTER TABLE livequiz_sessions ADD COLUMN auto_play BOOLEAN DEFAULT 0 NOT NULL"))
            conn.commit()
        if "question_closed_at" not in livequiz_col_names and livequiz_cols:
            conn.execute(text("ALTER TABLE livequiz_sessions ADD COLUMN question_closed_at DATETIME"))
            conn.commit()


def _ensure_class_access_details(cls: ClassModel, db: Session) -> ClassModel:
    changed = False

    if not (cls.class_code or "").strip():
        cls.class_code = _rand_class_code(db)
        changed = True

    if not (cls.class_pin or "").strip():
        cls.class_pin = _rand_class_pin()
        changed = True

    if changed:
        db.add(cls)

    return cls


DEMO_CLASS_NAME = "Demo Class"
DEMO_CLASS_SUBJECT = "Science"
DEMO_NOTES_PREFIX = "NOTES: "

DEMO_STUDENT_NAMES = [
    "Isaac Newton",
    "Marie Curie",
    "Albert Einstein",
    "Rosalind Franklin",
    "Galileo Galilei",
    "Ada Lovelace",
    "Charles Darwin",
    "Nikola Tesla",
    "Katherine Johnson",
    "Alan Turing",
    "Michael Faraday",
    "Lise Meitner",
    "Niels Bohr",
    "Emmy Noether",
    "Louis Pasteur",
    "Grace Hopper",
    "Johannes Kepler",
    "Jane Goodall",
    "James Clerk Maxwell",
    "Chien-Shiung Wu",
    "Carl Linnaeus",
    "Rachel Carson",
    "Srinivasa Ramanujan",
    "Dorothy Hodgkin",
    "Stephen Hawking",
]

DEMO_ASSESSMENT_TITLES = [
    "Baseline Quiz",
    "Topic Test 1",
    "Homework Check",
    "Midterm Assessment",
    "Practical Task",
    "Final Quiz",
]

DEMO_RESULTS_BY_STUDENT = {
    "Isaac Newton": [68, 71, 74, 72, 76, 78],
    "Marie Curie": [88, 90, 91, 92, 93, 94],
    "Albert Einstein": [91, 93, 92, 95, 94, 96],
    "Rosalind Franklin": [60, 64, 68, 72, 76, 81],
    "Galileo Galilei": [48, 52, 50, 55, 58, 61],
    "Ada Lovelace": [63, 67, 61, 72, 70, 74],
    "Charles Darwin": [69, 71, 73, 74, 72, 75],
    "Nikola Tesla": [82, 65, 88, 70, 91, 77],
    "Katherine Johnson": [58, 63, 69, 74, 79, 84],
    "Alan Turing": [84, 78, 69, 87, 73, 81],
    "Michael Faraday": [66, 68, 70, 72, 73, 75],
    "Lise Meitner": [50, 54, 56, 57, 60, 63],
    "Niels Bohr": [71, 73, 75, 76, 78, 79],
    "Emmy Noether": [89, 90, 92, 93, 94, 95],
    "Louis Pasteur": [67, 69, 72, 71, 74, 76],
    "Grace Hopper": [59, 65, 70, 74, 80, 85],
    "Johannes Kepler": [64, 66, 69, 71, 73, 74],
    "Jane Goodall": [72, 74, 73, 76, 78, 80],
    "James Clerk Maxwell": [79, 70, 82, 73, 86, 78],
    "Chien-Shiung Wu": [61, 66, 71, 77, 82, 86],
    "Carl Linnaeus": [52, 55, 57, 60, 62, 64],
    "Rachel Carson": [70, 72, 74, 75, 77, 79],
    "Srinivasa Ramanujan": [90, 92, 93, 94, 95, 97],
    "Dorothy Hodgkin": [54, 56, 58, 61, 63, 66],
    "Stephen Hawking": [62, 59, 65, 63, 69, 72],
}

DEMO_NOTE_FILENAMES = [
    "Events of the 20th Century.pdf",
    "Science in the 20th Century.pdf",
]
LEGACY_DEMO_NOTE_FILENAMES = {
    "Welcome to Demo Class.txt",
    "Seating Plan.txt",
    "How to Explore Demo Data.txt",
}


def _find_existing_demo_class_for_user(db: Session, user: models.UserModel) -> Optional[ClassModel]:
    return (
        db.query(ClassModel)
        .filter(
            ClassModel.owner_user_id == user.id,
            ClassModel.name == DEMO_CLASS_NAME,
        )
        .order_by(ClassModel.id.asc())
        .first()
    )


def _sync_demo_note_files(class_id: int, topic_id: int, db: Session) -> None:
    source_dir = UPLOADS_DIR / "demo"
    dest_dir = UPLOADS_DIR / "notes" / str(class_id)
    dest_dir.mkdir(parents=True, exist_ok=True)

    existing_notes = (
        db.query(models.Note)
        .filter(
            models.Note.class_id == class_id,
            models.Note.topic_id == topic_id,
        )
        .all()
    )
    existing_by_name = {note.filename: note for note in existing_notes}

    for note in existing_notes:
        if note.filename not in LEGACY_DEMO_NOTE_FILENAMES:
            continue
        try:
            Path(note.stored_path).unlink(missing_ok=True)
        except Exception:
            pass
        db.delete(note)

    for filename in DEMO_NOTE_FILENAMES:
        if filename in existing_by_name:
            continue

        source_path = source_dir / filename
        if not source_path.exists():
            logger.warning("Missing demo PDF: %s", source_path)
            raise HTTPException(status_code=500, detail=f"Missing demo resource: {filename}")

        disk_name = f"demo_{uuid.uuid4().hex}_{Path(filename).name}"
        stored_path = dest_dir / disk_name
        shutil.copyfile(source_path, stored_path)

        db.add(
            models.Note(
                class_id=class_id,
                topic_id=topic_id,
                filename=filename,
                stored_path=str(stored_path),
            )
        )


def _get_or_create_demo_topic(class_id: int, db: Session) -> models.Topic:
    topic = (
        db.query(models.Topic)
        .filter(
            models.Topic.class_id == class_id,
            models.Topic.name == f"{DEMO_NOTES_PREFIX}Getting Started",
        )
        .first()
    )
    if topic:
        return topic

    topic = models.Topic(class_id=class_id, name=f"{DEMO_NOTES_PREFIX}Getting Started")
    db.add(topic)
    db.flush()
    return topic


def _ensure_demo_students(class_id: int, db: Session) -> list[StudentModel]:
    students = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id)
        .filter(StudentModel.active == True)  # noqa: E712
        .all()
    )
    students_by_name = {student.first_name: student for student in students}

    for full_name in DEMO_STUDENT_NAMES:
        student = students_by_name.get(full_name)
        if student:
            if not student.active:
                student.active = True
                db.add(student)
            continue

        student = StudentModel(
            class_id=class_id,
            first_name=full_name,
            active=True,
        )
        db.add(student)
        students.append(student)

    db.flush()
    return students


def _ensure_demo_assessments(class_id: int, now: datetime, db: Session) -> list[ClassAssessmentModel]:
    assessments = db.query(ClassAssessmentModel).filter(ClassAssessmentModel.class_id == class_id).all()
    assessments_by_title = {assessment.title: assessment for assessment in assessments}

    for idx, title in enumerate(DEMO_ASSESSMENT_TITLES):
        assessment = assessments_by_title.get(title)
        if assessment:
            continue

        assessment = ClassAssessmentModel(
            class_id=class_id,
            title=title,
            assessment_date=now - timedelta(days=(len(DEMO_ASSESSMENT_TITLES) - idx) * 7),
        )
        db.add(assessment)
        assessments.append(assessment)

    db.flush()
    return assessments


def _ensure_demo_results(
    class_id: int,
    students: list[StudentModel],
    assessments: list[ClassAssessmentModel],
    db: Session,
) -> None:
    student_by_name = {student.first_name: student for student in students}
    assessment_by_title = {assessment.title: assessment for assessment in assessments}

    assessment_ids = [assessment.id for assessment in assessments if assessment.id is not None]
    existing_results = (
        db.query(AssessmentResultModel)
        .join(ClassAssessmentModel, AssessmentResultModel.assessment_id == ClassAssessmentModel.id)
        .filter(
            ClassAssessmentModel.class_id == class_id,
            AssessmentResultModel.assessment_id.in_(assessment_ids) if assessment_ids else False,
        )
        .all()
    )
    existing_pairs = {(result.assessment_id, result.student_id) for result in existing_results}

    for assessment_index, title in enumerate(DEMO_ASSESSMENT_TITLES):
        assessment = assessment_by_title.get(title)
        if not assessment:
            continue

        for student_name, scores in DEMO_RESULTS_BY_STUDENT.items():
            student = student_by_name.get(student_name)
            if not student:
                continue

            pair = (assessment.id, student.id)
            if pair in existing_pairs:
                continue

            db.add(
                AssessmentResultModel(
                    assessment_id=assessment.id,
                    student_id=student.id,
                    score_percent=int(scores[assessment_index]),
                    absent=False,
                )
            )


def _seed_demo_class(db: Session, user: models.UserModel) -> ClassModel:
    existing = _find_existing_demo_class_for_user(db, user)
    if existing:
        _ensure_class_access_details(existing, db)
        students = _ensure_demo_students(existing.id, db)
        assessments = _ensure_demo_assessments(existing.id, datetime.utcnow(), db)
        _ensure_demo_results(existing.id, students, assessments, db)
        demo_topic = _get_or_create_demo_topic(existing.id, db)
        _sync_demo_note_files(existing.id, demo_topic.id, db)
        db.commit()
        db.refresh(existing)
        _get_or_create_active_student_access_link(existing.id, db)
        return existing

    now = datetime.utcnow()
    demo_class = ClassModel(
        owner_user_id=user.id,
        name=DEMO_CLASS_NAME,
        subject=DEMO_CLASS_SUBJECT,
        class_code=_rand_class_code(db),
        class_pin=_rand_class_pin(),
    )
    db.add(demo_class)
    db.flush()

    students = _ensure_demo_students(demo_class.id, db)
    assessments = _ensure_demo_assessments(demo_class.id, now, db)
    _ensure_demo_results(demo_class.id, students, assessments, db)

    demo_topic = _get_or_create_demo_topic(demo_class.id, db)
    _sync_demo_note_files(demo_class.id, demo_topic.id, db)

    db.commit()
    db.refresh(demo_class)
    _get_or_create_active_student_access_link(demo_class.id, db)
    return demo_class


def _backfill_class_access_details(db: Session) -> None:
    classes = db.query(ClassModel).all()
    changed = False
    for cls in classes:
        before_code = cls.class_code
        before_pin = cls.class_pin
        _ensure_class_access_details(cls, db)
        if cls.class_code != before_code or cls.class_pin != before_pin:
            changed = True

    if changed:
        db.commit()

def _collab_time_left_seconds(s: CollabSessionModel) -> Optional[int]:
    if not s.timer_minutes or not s.breakout_started_at or s.state != "live":
        return None

    elapsed = int((datetime.utcnow() - s.breakout_started_at).total_seconds())
    total = int(s.timer_minutes) * 60
    left = total - elapsed
    return max(0, left)


def _rand_collab_code(db: Session) -> str:
    for _ in range(20):
        code = _rand_code(6)
        exists = db.query(CollabSessionModel).filter(CollabSessionModel.session_code == code).first()
        if not exists:
            return code
    raise HTTPException(status_code=500, detail="Could not generate collaboration code")

# key = (session_code, room_key) -> ordered list of board events
collab_room_history: dict[tuple[str, str], list[dict]] = defaultdict(list)


def _collab_history_key(session_code: str, room_key: str) -> tuple[str, str]:
    return (session_code, room_key)


def _append_collab_event(session_code: str, room_key: str, payload: dict) -> None:
    key = _collab_history_key(session_code, room_key)
    collab_room_history[key].append(deepcopy(payload))


def _replace_collab_history(session_code: str, room_key: str, events: list[dict]) -> None:
    key = _collab_history_key(session_code, room_key)
    collab_room_history[key] = [deepcopy(evt) for evt in events]


def _events_from_snapshot(snapshot: dict) -> list[dict]:
    events: list[dict] = []

    for stroke in snapshot.get("strokes", []):
        events.append({
            "type": "stroke",
            "stroke": stroke,
        })

    for obj in snapshot.get("objects", []):
        events.append({
            "type": "object-create",
            "object": obj,
        })

    return events


def _get_collab_history(session_code: str, room_key: str) -> list[dict]:
    key = _collab_history_key(session_code, room_key)
    return collab_room_history.get(key, [])


async def _seed_breakout_rooms_from_teacher(session_code: str, room_count: int) -> None:
    teacher_events = _get_collab_history(session_code, "teacher-main")

    for i in range(1, int(room_count) + 1):
        room_key = f"room-{i}"

        # 1) Replace stored history for the room
        _replace_collab_history(session_code, room_key, teacher_events)

        # 2) Push the copied events into any sockets already connected to that room
        for evt in teacher_events:
            await collab_room_manager.broadcast(session_code, room_key, deepcopy(evt))


def _clear_collab_session_history(session_code: str) -> None:
    dead_keys = [key for key in collab_room_history.keys() if key[0] == session_code]
    for key in dead_keys:
        del collab_room_history[key]

@app.get("/collab/{code}/me/{anon_id}")
def collab_me(code: str, anon_id: str, db: Session = Depends(get_db)):
    s = db.query(CollabSessionModel).filter(CollabSessionModel.session_code == code).first()
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")

    p = (
        db.query(CollabParticipantModel)
        .filter(CollabParticipantModel.session_id == s.id)
        .filter(CollabParticipantModel.anon_id == anon_id)
        .first()
    )
    if not p:
        raise HTTPException(status_code=404, detail="Participant not found")

    return {
        "id": p.id,
        "anon_id": p.anon_id,
        "name": p.name,
        "room_number": p.room_number,
        "is_online": bool(p.is_online),
        "session_state": s.state,
    }

@app.websocket("/ws/collab/{session_code}/{room_key}")
async def collab_ws(websocket: WebSocket, session_code: str, room_key: str):
    print("WS connect attempt:", session_code, room_key)
    await websocket.accept()
    print("WS accepted")

    collab_room_manager.rooms[(session_code, room_key)].append(websocket)

    try:
        for evt in _get_collab_history(session_code, room_key):
            await websocket.send_json(evt)

        await collab_room_manager.broadcast(session_code, room_key, {
            "type": "presence",
            "message": "peer_joined",
        })

        while True:
            raw = await websocket.receive_text()

            try:
                data = json.loads(raw)
            except Exception:
                await websocket.send_json({
                    "type": "error",
                    "message": "Invalid JSON payload",
                })
                continue

            msg_type = data.get("type")

            if msg_type == "ping":
                await websocket.send_json({"type": "pong"})
                continue

            if msg_type == "pong":
                continue

            if msg_type in {
                "stroke",
                "object-create",
                "object-update",
                "object-delete",
            }:
                _append_collab_event(session_code, room_key, data)
                await collab_room_manager.broadcast(session_code, room_key, data)
                continue

            if msg_type == "snapshot-sync":
                snapshot = data.get("snapshot") or {"strokes": [], "objects": []}
                replacement_events = _events_from_snapshot(snapshot)
                _replace_collab_history(session_code, room_key, replacement_events)
                await collab_room_manager.broadcast(session_code, room_key, data)
                continue

            if msg_type in {
                "stroke-progress",
                "cursor",
                "clear-preview",
            }:
                await collab_room_manager.broadcast(session_code, room_key, data)
                continue

            print("WS unknown message type:", msg_type)

    except WebSocketDisconnect:
        print("WS disconnected:", session_code, room_key)
        collab_room_manager.disconnect(session_code, room_key, websocket)

    except Exception as e:
        print("WS error:", str(e))
        collab_room_manager.disconnect(session_code, room_key, websocket)
        try:
            await websocket.close()
        except:
            pass

def _assert_class_access(class_id: int, db: Session, user: models.UserModel):
    cls = (
        db.query(ClassModel)
        .filter(ClassModel.id == class_id)
        .filter(ClassModel.owner_user_id == user.id)
        .first()
    )
    if not cls:
        raise HTTPException(status_code=404, detail="Class not found")
    return cls


def _quiz_out(q: models.SavedQuizModel) -> dict:
    origin_name = None
    try:
        origin_name = q.class_rel.name if getattr(q, "class_rel", None) else None
    except Exception:
        origin_name = None
    return {
        "id": q.id,
        "class_id": q.class_id,
        "title": q.title,
        "category": q.category,
        "description": q.description,
        "is_starred": bool(getattr(q, "is_starred", False)),
        "origin_class_name": origin_name,
        "created_at": q.created_at,
        "updated_at": q.updated_at,
        "questions": [
            {
                "id": qq.id,
                "prompt": qq.prompt,
                "choices": [qq.choice_a, qq.choice_b, qq.choice_c, qq.choice_d],
                "correct_index": qq.correct_index,
                "explanation": qq.explanation,
                "position": qq.position,
            }
            for qq in sorted(q.questions, key=lambda x: x.position)
        ],
    }


def _normalise_choices(choices: list[str]) -> list[str]:
    vals = [str(x).strip() for x in (choices or [])][:4]
    while len(vals) < 4:
        vals.append("")
    if not all(vals[:4]):
        raise HTTPException(status_code=400, detail="Exactly 4 non-empty choices are required")
    return vals[:4]


def _get_accessible_quiz_or_404(
    quiz_id: int,
    db: Session,
    user: models.UserModel,
) -> models.SavedQuizModel:
    quiz = db.query(models.SavedQuizModel).filter(models.SavedQuizModel.id == quiz_id).first()
    if not quiz:
        raise HTTPException(status_code=404, detail="Quiz not found")

    _assert_class_access(quiz.class_id, db, user)
    return quiz

@app.get("/classes/{class_id}/quizzes", response_model=List[schemas.SavedQuizOut])
def list_saved_quizzes(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _assert_class_access(class_id, db, user)

    rows = (
        db.query(models.SavedQuizModel)
        .filter(models.SavedQuizModel.class_id == class_id)
        .order_by(models.SavedQuizModel.created_at.desc())
        .all()
    )
    return [_quiz_out(row) for row in rows]


@app.get("/quizzes/starred", response_model=List[schemas.SavedQuizOut])
def list_starred_quizzes(
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    rows = (
        db.query(models.SavedQuizModel)
        .join(ClassModel, ClassModel.id == models.SavedQuizModel.class_id)
        .filter(ClassModel.owner_user_id == user.id)
        .filter(models.SavedQuizModel.is_starred == True)
        .order_by(models.SavedQuizModel.updated_at.desc(), models.SavedQuizModel.created_at.desc())
        .all()
    )
    return [_quiz_out(row) for row in rows]


@app.post("/classes/{class_id}/quizzes", response_model=schemas.SavedQuizOut)
def create_saved_quiz(
    class_id: int,
    payload: schemas.SavedQuizCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _assert_class_access(class_id, db, user)

    title = (payload.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title is required")

    now = datetime.utcnow()
    quiz = models.SavedQuizModel(
        class_id=class_id,
        owner_user_id=user.id,
        title=title,
        category=(payload.category or "General").strip() or "General",
        description=(payload.description or "").strip() or None,
        created_at=now,
        updated_at=now,
    )
    db.add(quiz)
    db.flush()

    for idx, q in enumerate(payload.questions or []):
        choices = _normalise_choices(q.choices)
        row = models.SavedQuizQuestionModel(
            quiz_id=quiz.id,
            prompt=(q.prompt or "").strip(),
            choice_a=choices[0],
            choice_b=choices[1],
            choice_c=choices[2],
            choice_d=choices[3],
            correct_index=max(0, min(3, int(q.correct_index or 0))),
            explanation=(q.explanation or "").strip() or None,
            position=int(q.position if q.position is not None else idx),
            created_at=now,
            updated_at=now,
        )
        if not row.prompt:
            raise HTTPException(status_code=400, detail="Question prompt is required")
        db.add(row)

    db.commit()
    db.refresh(quiz)
    return _quiz_out(quiz)


@app.post("/quizzes/{quiz_id}/star", response_model=schemas.SavedQuizOut)
def star_saved_quiz(
    quiz_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    quiz = _get_accessible_quiz_or_404(quiz_id, db, user)

    quiz.is_starred = True
    if quiz.owner_user_id is None:
        quiz.owner_user_id = user.id
    quiz.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(quiz)
    return _quiz_out(quiz)


@app.post("/quizzes/{quiz_id}/unstar", response_model=schemas.SavedQuizOut)
def unstar_saved_quiz(
    quiz_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    quiz = _get_accessible_quiz_or_404(quiz_id, db, user)

    quiz.is_starred = False
    if quiz.owner_user_id is None:
        quiz.owner_user_id = user.id
    quiz.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(quiz)
    return _quiz_out(quiz)


@app.put("/quizzes/{quiz_id}", response_model=schemas.SavedQuizOut)
def update_saved_quiz(
    quiz_id: int,
    payload: schemas.SavedQuizUpdate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    quiz = _get_accessible_quiz_or_404(quiz_id, db, user)

    if payload.title is not None:
        title = payload.title.strip()
        if not title:
            raise HTTPException(status_code=400, detail="Title is required")
        quiz.title = title

    if payload.category is not None:
        quiz.category = payload.category.strip() or "General"

    if payload.description is not None:
        quiz.description = payload.description.strip() or None

    quiz.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(quiz)
    return _quiz_out(quiz)


@app.delete("/quizzes/{quiz_id}")
def delete_saved_quiz(
    quiz_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    quiz = _get_accessible_quiz_or_404(quiz_id, db, user)

    db.delete(quiz)
    db.commit()
    return {"message": "deleted"}


@app.post("/quizzes/{quiz_id}/questions", response_model=schemas.SavedQuizOut)
def add_quiz_question(
    quiz_id: int,
    payload: schemas.SavedQuizQuestionCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    quiz = _get_accessible_quiz_or_404(quiz_id, db, user)

    choices = _normalise_choices(payload.choices)
    prompt = (payload.prompt or "").strip()
    if not prompt:
        raise HTTPException(status_code=400, detail="Question prompt is required")

    max_pos = max([q.position for q in quiz.questions], default=-1)
    row = models.SavedQuizQuestionModel(
        quiz_id=quiz.id,
        prompt=prompt,
        choice_a=choices[0],
        choice_b=choices[1],
        choice_c=choices[2],
        choice_d=choices[3],
        correct_index=max(0, min(3, int(payload.correct_index or 0))),
        explanation=(payload.explanation or "").strip() or None,
        position=int(payload.position) if payload.position is not None else max_pos + 1,
        created_at=datetime.utcnow(),
        updated_at=datetime.utcnow(),
    )
    db.add(row)
    quiz.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(quiz)
    return _quiz_out(quiz)


@app.put("/quizzes/{quiz_id}/questions/{question_id}", response_model=schemas.SavedQuizOut)
def update_quiz_question(
    quiz_id: int,
    question_id: int,
    payload: schemas.SavedQuizQuestionUpdate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    quiz = _get_accessible_quiz_or_404(quiz_id, db, user)

    q = (
        db.query(models.SavedQuizQuestionModel)
        .filter(models.SavedQuizQuestionModel.id == question_id)
        .filter(models.SavedQuizQuestionModel.quiz_id == quiz.id)
        .first()
    )
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")

    if payload.prompt is not None:
        prompt = payload.prompt.strip()
        if not prompt:
            raise HTTPException(status_code=400, detail="Question prompt is required")
        q.prompt = prompt

    if payload.choices is not None:
        choices = _normalise_choices(payload.choices)
        q.choice_a, q.choice_b, q.choice_c, q.choice_d = choices

    if payload.correct_index is not None:
        q.correct_index = max(0, min(3, int(payload.correct_index)))

    if payload.explanation is not None:
        q.explanation = payload.explanation.strip() or None

    if payload.position is not None:
        q.position = int(payload.position)

    q.updated_at = datetime.utcnow()
    quiz.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(quiz)
    return _quiz_out(quiz)


@app.delete("/quizzes/{quiz_id}/questions/{question_id}", response_model=schemas.SavedQuizOut)
def delete_quiz_question(
    quiz_id: int,
    question_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    quiz = _get_accessible_quiz_or_404(quiz_id, db, user)

    q = (
        db.query(models.SavedQuizQuestionModel)
        .filter(models.SavedQuizQuestionModel.id == question_id)
        .filter(models.SavedQuizQuestionModel.quiz_id == quiz.id)
        .first()
    )
    if not q:
        raise HTTPException(status_code=404, detail="Question not found")

    db.delete(q)
    quiz.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(quiz)
    return _quiz_out(quiz)

@app.post("/livequiz/create", response_model=LiveQuizCreateResponse)
def livequiz_create(
    payload: LiveQuizCreateRequest,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(payload.class_id, db, user)
    title = (payload.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title is required")

    questions = payload.questions or []
    if len(questions) == 0:
        raise HTTPException(status_code=400, detail="At least one question is required")

    # normalise + validate choices
    cleaned = []
    for q in questions:
        qid = (q.id or "").strip()
        prompt = (q.prompt or "").strip()
        if not qid or not prompt:
            raise HTTPException(status_code=400, detail="Each question needs id and prompt")

        choices = q.choices or {}
        # enforce A-D keys
        obj = {
            "id": qid,
            "prompt": prompt,
            "choices": {
                "A": str(choices.get("A","")).strip(),
                "B": str(choices.get("B","")).strip(),
                "C": str(choices.get("C","")).strip(),
                "D": str(choices.get("D","")).strip(),
            },
            "correct": (q.correct or None),
        }
        non_empty = [v for v in obj["choices"].values() if v]
        if len(non_empty) < 2:
            raise HTTPException(status_code=400, detail=f"Question '{prompt}' needs at least 2 options")
        cleaned.append(obj)

    if payload.shuffle_questions:
        random.shuffle(cleaned)

    # session code unique
    for _ in range(20):
        code = _rand_code(6)
        exists = db.query(LiveQuizSessionModel).filter(LiveQuizSessionModel.session_code == code).first()
        if not exists:
            break
    else:
        raise HTTPException(status_code=500, detail="Could not generate session code")

    s = LiveQuizSessionModel(
        class_id=payload.class_id,
        session_code=code,
        title=title,
        anonymous=bool(payload.anonymous),
        quiz_id=(str(payload.quiz_id).strip() if payload.quiz_id is not None and str(payload.quiz_id).strip() else None),
        questions_json=json.dumps(cleaned, ensure_ascii=False),
        state="lobby",
        current_index=-1,
        seconds_per_question=(int(payload.seconds_per_question) if payload.seconds_per_question else None),
        shuffle_questions=bool(payload.shuffle_questions),
        auto_play=bool(payload.auto_play),
        auto_end_when_all_answered=bool(payload.auto_end_when_all_answered),
    )
    db.add(s)
    db.commit()
    db.refresh(s)

    return {"session_code": s.session_code, "join_url": None}

@app.get("/livequiz/{code}/status")
def livequiz_status(code: str, db: Session = Depends(get_db)):
    s = db.query(LiveQuizSessionModel).filter(LiveQuizSessionModel.session_code == code).first()
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")
    s = _maybe_progress_livequiz_session(db, s)

    qs = _load_questions(s)
    joined = db.query(LiveQuizParticipantModel).filter(LiveQuizParticipantModel.session_id == s.id).count()

    # answered count for current question
    answered = 0
    cq = _current_question(s)
    if cq:
        qid = cq.get("id")
        answered = (
            db.query(LiveQuizAnswerModel)
            .filter(LiveQuizAnswerModel.session_id == s.id)
            .filter(LiveQuizAnswerModel.question_id == qid)
            .count()
        )

    return {
        "session_code": s.session_code,
        "state": s.state,
        "title": s.title,
        "anonymous": s.anonymous,
        "auto_play": s.auto_play,
        "seconds_per_question": s.seconds_per_question,
        "current_index": s.current_index,
        "total_questions": len(qs),
        "time_left_seconds": _time_left_seconds(s),
        "joined_count": joined,
        "answered_count": answered,
        "answers_open": _livequiz_answers_open(s),
    }

@app.post("/livequiz/{code}/start")
def livequiz_start(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_livequiz_session_or_404(code, db, user)
    qs = _load_questions(s)
    if not qs:
        raise HTTPException(status_code=400, detail="No questions")

    s.state = "live"
    s.started_at = datetime.utcnow()
    s.current_index = 0
    s.question_started_at = datetime.utcnow()
    s.question_closed_at = None
    db.commit()
    return {"message": "started"}

@app.post("/livequiz/{code}/next")
def livequiz_next(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_livequiz_session_or_404(code, db, user)
    qs = _load_questions(s)
    if not qs:
        raise HTTPException(status_code=400, detail="No questions")

    if s.state != "live":
        raise HTTPException(status_code=400, detail="Session is not live")

    nxt = int(s.current_index) + 1
    if nxt >= len(qs):
        _finalize_livequiz_session(db, s)
        db.commit()
        return {"message": "ended"}

    s.current_index = nxt
    s.question_started_at = datetime.utcnow()
    s.question_closed_at = None
    db.commit()
    return {"message": "next"}

@app.post("/livequiz/{code}/end-question")
def livequiz_end_question(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_livequiz_session_or_404(code, db, user)
    if s.state != "live":
        raise HTTPException(status_code=400, detail="Session is not live")

    s.question_closed_at = datetime.utcnow()
    db.commit()
    return {"message": "ended_question"}

@app.post("/livequiz/{code}/end-session")
def livequiz_end_session(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_livequiz_session_or_404(code, db, user)
    _finalize_livequiz_session(db, s)
    db.commit()
    return {"message": "ended"}

@app.post("/livequiz/{code}/join", response_model=LiveQuizJoinResponse)
def livequiz_join(code: str, payload: LiveQuizJoinRequest, db: Session = Depends(get_db)):
    s = db.query(LiveQuizSessionModel).filter(LiveQuizSessionModel.session_code == code).first()
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")

    anon_id = (payload.anon_id or "").strip() or uuid.uuid4().hex

    existing = (
        db.query(LiveQuizParticipantModel)
        .filter(LiveQuizParticipantModel.session_id == s.id)
        .filter(LiveQuizParticipantModel.anon_id == anon_id)
        .first()
    )
    if existing:
        return {"anon_id": existing.anon_id, "nickname": existing.nickname}

    provided = (payload.name or "").strip()

    # If anonymous mode, do not store names
    if s.anonymous:
        nickname = None
    else:
        nickname = provided if provided else _kid_name()

    p = LiveQuizParticipantModel(session_id=s.id, anon_id=anon_id, nickname=nickname)
    db.add(p)
    db.commit()
    db.refresh(p)

    return {"anon_id": p.anon_id, "nickname": p.nickname}

@app.get("/livequiz/{code}/current")
def livequiz_current(code: str, db: Session = Depends(get_db)):
    s = db.query(LiveQuizSessionModel).filter(LiveQuizSessionModel.session_code == code).first()
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")
    s = _maybe_progress_livequiz_session(db, s)

    q = _current_question(s)
    return {
        "state": s.state,
        "title": s.title,
        "anonymous": s.anonymous,
        "auto_play": s.auto_play,
        "current_index": s.current_index,
        "total_questions": len(_load_questions(s)),
        "time_left_seconds": _time_left_seconds(s),
        "answers_open": _livequiz_answers_open(s),
        "question": q,
    }

@app.post("/livequiz/{code}/answer")
def livequiz_answer(code: str, payload: LiveQuizAnswerRequest, db: Session = Depends(get_db)):
    s = db.query(LiveQuizSessionModel).filter(LiveQuizSessionModel.session_code == code).first()
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")
    s = _maybe_progress_livequiz_session(db, s)

    anon_id = (payload.anon_id or "").strip()
    if not anon_id:
        raise HTTPException(status_code=400, detail="anon_id is required")

    choice = (payload.choice or "").strip().upper()
    if choice not in ["A", "B", "C", "D"]:
        raise HTTPException(status_code=400, detail="choice must be A/B/C/D")

    qid = (payload.question_id or "").strip()
    if not qid:
        raise HTTPException(status_code=400, detail="question_id is required")

    p = (
        db.query(LiveQuizParticipantModel)
        .filter(LiveQuizParticipantModel.session_id == s.id)
        .filter(LiveQuizParticipantModel.anon_id == anon_id)
        .first()
    )
    if not p:
        raise HTTPException(status_code=404, detail="Participant not joined")
    if not _livequiz_answers_open(s):
        raise HTTPException(status_code=400, detail="This question is closed")
    current_question = _current_question(s)
    if not current_question or str(current_question.get("id") or "").strip() != qid:
        raise HTTPException(status_code=400, detail="This question is no longer active")

    # prevent double-answering same question by same participant
    existing = (
        db.query(LiveQuizAnswerModel)
        .filter(LiveQuizAnswerModel.session_id == s.id)
        .filter(LiveQuizAnswerModel.participant_id == p.id)
        .filter(LiveQuizAnswerModel.question_id == qid)
        .first()
    )
    if existing:
        existing.choice = choice
        existing.answered_at = datetime.utcnow()
        db.commit()
    else:
        a = LiveQuizAnswerModel(session_id=s.id, participant_id=p.id, question_id=qid, choice=choice)
        db.add(a)
        db.commit()

    # auto-end check (teacher can still press End Q early)
    if s.auto_end_when_all_answered and s.state == "live":
        joined = db.query(LiveQuizParticipantModel).filter(LiveQuizParticipantModel.session_id == s.id).count()
        answered = (
            db.query(LiveQuizAnswerModel)
            .filter(LiveQuizAnswerModel.session_id == s.id)
            .filter(LiveQuizAnswerModel.question_id == qid)
            .count()
        )
        if joined > 0 and answered >= joined and s.seconds_per_question:
            s.question_closed_at = datetime.utcnow()
            db.commit()

    return {"message": "ok"}

def _build_livequiz_results(db: Session, s: LiveQuizSessionModel) -> dict:
    qs = _load_questions(s)

    # Map question_id -> correct option (A/B/C/D) if available
    correct_map = {}
    for q in qs:
        qid = str(q.get("id", "")).strip()
        corr = q.get("correct", None)
        corr = (str(corr).strip().upper() if corr is not None else None)
        if qid:
            correct_map[qid] = corr if corr in ["A", "B", "C", "D"] else None

    participants = (
        db.query(LiveQuizParticipantModel)
        .filter(LiveQuizParticipantModel.session_id == s.id)
        .all()
    )

    # Name display (nickname stored; in anonymous mode nickname is None)
    pid_to_name = {}
    for p in participants:
        if s.anonymous:
            pid_to_name[p.id] = "Anonymous"
        else:
            pid_to_name[p.id] = (p.nickname or "Player")

    answers = (
        db.query(LiveQuizAnswerModel)
        .filter(LiveQuizAnswerModel.session_id == s.id)
        .all()
    )

    # Score per participant
    by_pid = {}
    for a in answers:
        pid = int(a.participant_id)
        qid = str(a.question_id)
        choice = str(a.choice).strip().upper()

        if pid not in by_pid:
            by_pid[pid] = {"answered": 0, "correct": 0}

        # count only valid choices
        if choice in ["A", "B", "C", "D"]:
            by_pid[pid]["answered"] += 1

        corr = correct_map.get(qid, None)
        if corr and choice == corr:
            by_pid[pid]["correct"] += 1

    total_qs = len(qs)
    any_correct_keys = any(v in ["A", "B", "C", "D"] for v in correct_map.values())

    leaderboard = []
    for p in participants:
        stats = by_pid.get(p.id, {"answered": 0, "correct": 0})
        answered = int(stats["answered"])
        correct = int(stats["correct"])

        # Percent must always reflect correct answers
        if total_qs:
            percent = int(round((correct / total_qs) * 100))
        else:
            percent = 0

        leaderboard.append({
            "participant_id": p.id,
            "name": pid_to_name.get(p.id, "Player"),
            "correct": correct,
            "answered": answered,
            "total_questions": total_qs,
            "percent": percent,
        })

    # Sort: if correct keys exist, sort by correct then answered; otherwise answered then name
    leaderboard.sort(key=lambda r: (-r["correct"], -r["answered"], r["name"].lower()))

    top3 = leaderboard[:3]

    # Question stats: counts per option + correct rate when available
    q_counts = {}
    for a in answers:
        qid = str(a.question_id)
        choice = str(a.choice).strip().upper()
        if qid not in q_counts:
            q_counts[qid] = {"A": 0, "B": 0, "C": 0, "D": 0, "total": 0}
        if choice in ["A", "B", "C", "D"]:
            q_counts[qid][choice] += 1
            q_counts[qid]["total"] += 1

    question_stats = []
    for q in qs:
        qid = str(q.get("id", "")).strip()
        prompt = str(q.get("prompt", "")).strip()
        counts = q_counts.get(qid, {"A": 0, "B": 0, "C": 0, "D": 0, "total": 0})
        corr = correct_map.get(qid, None)

        total = int(counts["total"])
        correct_ct = int(counts.get(corr, 0)) if corr else 0
        correct_rate = (correct_ct / total) if (total and corr) else None

        # most common wrong choice
        most_wrong = None
        if corr and total:
            wrong_items = [(k, counts[k]) for k in ["A", "B", "C", "D"] if k != corr]
            wrong_items.sort(key=lambda x: -x[1])
            if wrong_items and wrong_items[0][1] > 0:
                most_wrong = wrong_items[0][0]

        question_stats.append({
            "question_id": qid,
            "prompt": prompt,
            "correct": corr,
            "counts": {k: int(counts[k]) for k in ["A", "B", "C", "D"]},
            "total_answers": total,
            "correct_rate": (float(correct_rate) if correct_rate is not None else None),
            "most_common_wrong": most_wrong,
        })

    joined = len(participants)
    attempted_any = sum(1 for r in leaderboard if r["answered"] > 0)
    avg_percent = int(round(sum(r["percent"] for r in leaderboard) / joined)) if joined else 0

    # hardest question = lowest correct_rate
    hardest = None
    rated = [q for q in question_stats if isinstance(q["correct_rate"], float)]
    if rated:
        rated.sort(key=lambda q: q["correct_rate"])
        hardest = {
            "question_id": rated[0]["question_id"],
            "prompt": rated[0]["prompt"],
            "correct_rate": rated[0]["correct_rate"],
        }

    return {
        "session_code": s.session_code,
        "class_id": s.class_id,
        "title": s.title,
        "anonymous": s.anonymous,
        "auto_play": s.auto_play,
        "state": s.state,
        "created_at": s.created_at.isoformat() if s.created_at else None,
        "started_at": s.started_at.isoformat() if s.started_at else None,
        "ended_at": s.ended_at.isoformat() if s.ended_at else None,
        "summary": {
            "joined": joined,
            "attempted_any": attempted_any,
            "total_questions": total_qs,
            "avg_percent": avg_percent,
            "hardest_question": hardest,
            "scored_mode": bool(any_correct_keys),
        },
        "top3": top3,
        "leaderboard": leaderboard,
        "question_stats": question_stats,
    }


@app.get("/livequiz/{code}/results")
def livequiz_results(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_livequiz_session_or_404(code, db, user)
    if s.state == "ended":
        _sync_livequiz_attempts(db, s)
        db.commit()
    return _build_livequiz_results(db, s)

@app.get("/classes/{class_id}/livequiz/insights")
def class_livequiz_insights(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)

    attempts = (
        db.query(LiveQuizAttemptModel)
        .filter(LiveQuizAttemptModel.class_id == class_id)
        .order_by(LiveQuizAttemptModel.finished_at.desc(), LiveQuizAttemptModel.id.desc())
        .all()
    )

    grouped: dict[str, dict[str, Any]] = {}
    for attempt in attempts:
        group_key = f"student:{attempt.student_id}" if attempt.student_id else f"name:{_normalise_livequiz_participant_name(attempt.participant_display_name)}"
        bucket = grouped.setdefault(
            group_key,
            {
                "group_key": group_key,
                "student_id": attempt.student_id,
                "display_name": attempt.participant_display_name,
                "counted_attempts": 0,
                "average_percent": None,
                "attempts": [],
            },
        )

        is_counted = (
            not attempt.excluded_from_average
            and attempt.score_percent is not None
        )
        attempt_payload = {
            "id": attempt.id,
            "session_id": attempt.session_id,
            "quiz_id": attempt.quiz_id,
            "participant_display_name": attempt.participant_display_name,
            "score": attempt.score,
            "score_percent": attempt.score_percent,
            "total_questions": attempt.total_questions,
            "completed": attempt.completed,
            "scored_mode": attempt.scored_mode,
            "excluded_from_average": attempt.excluded_from_average,
            "submitted_at": attempt.submitted_at.isoformat() if attempt.submitted_at else None,
            "finished_at": attempt.finished_at.isoformat() if attempt.finished_at else None,
            "counted": is_counted,
        }
        bucket["attempts"].append(attempt_payload)

    for bucket in grouped.values():
        counted = [a["score_percent"] for a in bucket["attempts"] if a["counted"] and a["score_percent"] is not None]
        bucket["counted_attempts"] = len(counted)
        bucket["average_percent"] = int(round(sum(counted) / len(counted))) if counted else None

    return {
        "summary": {
            "total_attempts": len(attempts),
            "counted_attempts": sum(1 for attempt in attempts if not attempt.excluded_from_average and attempt.score_percent is not None),
            "excluded_attempts": sum(1 for attempt in attempts if attempt.excluded_from_average),
            "incomplete_attempts": sum(1 for attempt in attempts if not attempt.completed),
        },
        "students": list(grouped.values()),
    }

@app.post("/livequiz/attempts/{attempt_id}/exclude")
def livequiz_attempt_exclude(
    attempt_id: int,
    payload: LiveQuizAttemptExcludeRequest,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    attempt = db.query(LiveQuizAttemptModel).filter(LiveQuizAttemptModel.id == attempt_id).first()
    if not attempt:
        raise HTTPException(status_code=404, detail="Attempt not found")

    get_owned_class_or_404(attempt.class_id, db, user)
    attempt.excluded_from_average = bool(payload.excluded)
    db.commit()
    db.refresh(attempt)

    return {
        "success": True,
        "attempt_id": attempt.id,
        "excluded_from_average": attempt.excluded_from_average,
    }

def normalize_teacher_planner_payload(payload: Any) -> dict:
    if not isinstance(payload, dict):
        payload = {}

    notes = payload.get("notes") if isinstance(payload.get("notes"), list) else []
    tasks = payload.get("tasks") if isinstance(payload.get("tasks"), list) else []

    settings = payload.get("settings") if isinstance(payload.get("settings"), dict) else {}
    raw_slots = settings.get("slotsPerDay", 6)

    try:
        slots_per_day = int(raw_slots)
    except Exception:
        slots_per_day = 6

    slots_per_day = max(6, min(10, slots_per_day))

    return {
        "notes": notes,
        "tasks": tasks,
        "settings": {
            "slotsPerDay": slots_per_day,
        },
    }


@app.get("/teacher-planner")
def get_teacher_planner(
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    planner = (
        db.query(TeacherPlannerStateModel)
        .filter(TeacherPlannerStateModel.teacher_id == current_user.id)
        .first()
    )

    if not planner or not planner.state_json:
        return normalize_teacher_planner_payload({})

    try:
        raw_state = json.loads(planner.state_json)
    except Exception:
        raw_state = {}

    return normalize_teacher_planner_payload(raw_state)


@app.put("/teacher-planner")
def update_teacher_planner(
    body: dict,
    db: Session = Depends(get_db),
    current_user = Depends(get_current_user),
):
    payload = normalize_teacher_planner_payload(body)

    planner = (
        db.query(TeacherPlannerStateModel)
        .filter(TeacherPlannerStateModel.teacher_id == current_user.id)
        .first()
    )

    now_iso = datetime.utcnow().isoformat()

    if not planner:
        planner = TeacherPlannerStateModel(
            teacher_id=current_user.id,
            state_json=json.dumps(payload, ensure_ascii=False),
            updated_at=now_iso,
        )
        db.add(planner)
    else:
        planner.state_json = json.dumps(payload, ensure_ascii=False)
        planner.updated_at = now_iso

    db.commit()
    db.refresh(planner)
    return payload

# =========================================================
# DB
# =========================================================

Base.metadata.create_all(bind=engine)

# =========================================================
# FILES / UPLOADS (ABSOLUTE + STABLE)
# =========================================================
LEGACY_PUBLIC_UPLOAD_DIRS = {"notes", "tests", "posts", "whiteboards"}
EXAM_LIBRARY_DIR = Path(os.getenv("ELUME_EXAM_LIBRARY_DIR") or "/var/lib/elume/exam-library")
EXAM_LIBRARY_MANIFEST = EXAM_LIBRARY_DIR / "manifest.json"


def _read_exam_library_manifest() -> list[schemas.ExamLibraryItemOut]:
    if not EXAM_LIBRARY_MANIFEST.exists():
        return []

    try:
        raw = json.loads(EXAM_LIBRARY_MANIFEST.read_text(encoding="utf-8"))
    except Exception:
        logger.exception("Failed to read exam library manifest: %s", EXAM_LIBRARY_MANIFEST)
        return []

    items = raw.get("items") if isinstance(raw, dict) else raw
    if not isinstance(items, list):
        return []

    out: list[schemas.ExamLibraryItemOut] = []
    seen_ids: set[str] = set()

    for entry in items:
        if not isinstance(entry, dict):
            continue

        item_id = str(entry.get("id") or "").strip()
        cycle = str(entry.get("cycle") or "").strip()
        subject = str(entry.get("subject") or "").strip()
        level = str(entry.get("level") or "").strip()
        year = str(entry.get("year") or "").strip()
        title = str(entry.get("title") or "").strip()
        rel_path = str(entry.get("path") or "").strip().replace("\\", "/")

        if not item_id or item_id in seen_ids:
            continue
        if not cycle or not subject or not level or not year or not title or not rel_path:
            continue

        full_path = (EXAM_LIBRARY_DIR / rel_path).resolve()
        try:
            full_path.relative_to(EXAM_LIBRARY_DIR.resolve())
        except Exception:
            continue

        if not full_path.exists() or not full_path.is_file():
            continue

        seen_ids.add(item_id)
        out.append(
            schemas.ExamLibraryItemOut(
                id=item_id,
                cycle=cycle,
                subject=subject,
                level=level,
                year=year,
                title=title,
                path=rel_path,
                file_url=f"/exam-library/items/{item_id}/download",
            )
        )

    out.sort(key=lambda item: (item.subject.lower(), item.cycle.lower(), item.level.lower(), item.year.lower(), item.title.lower()))
    return out


def _find_exam_library_item_or_404(item_id: str) -> schemas.ExamLibraryItemOut:
    for item in _read_exam_library_manifest():
        if item.id == item_id:
            return item
    raise HTTPException(status_code=404, detail="Exam library item not found")


def _rel_upload_url(stored_path: str) -> str:
    """
    Convert a stored_path on disk into a URL like /uploads/notes/1/123_file.pdf
    Works whether stored_path is absolute or relative.
    """
    p = Path(stored_path)

    # If absolute path inside UPLOADS_DIR, make it relative to uploads root
    try:
        rel = p.resolve().relative_to(UPLOADS_DIR.resolve()).as_posix()
        return f"/uploads/{rel}"
    except Exception:
        pass

    # Fallback: split on "uploads/" if someone stored a path containing it
    s = p.as_posix()
    if "uploads/" in s:
        rel = s.split("uploads/")[-1]
        return f"/uploads/{rel}"

    # Last resort: just expose basename (not ideal, but avoids crashing)
    return f"/uploads/{p.name}"


@app.get("/uploads/{file_path:path}")
def legacy_upload_access(file_path: str):
    rel = Path(file_path)
    if not rel.parts:
        raise HTTPException(status_code=404, detail="File not found")

    if rel.parts[0] not in LEGACY_PUBLIC_UPLOAD_DIRS:
        raise HTTPException(status_code=404, detail="File not found")

    full_path = _resolve_upload_relpath_or_none(str(rel))
    if not full_path:
        raise HTTPException(status_code=404, detail="File not found")

    return FileResponse(
        path=full_path,
        filename=full_path.name,
        headers={"X-Elume-Legacy-Uploads": "true"},
    )


@app.post("/collab/create", response_model=CollabCreateResponse)
def collab_create(
    payload: CollabCreatePayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(payload.class_id, db, user)
    title = (payload.title or "").strip() or "Collaboration Whiteboard"
    room_count = max(1, min(12, int(payload.room_count or 4)))
    timer_minutes = payload.timer_minutes
    if timer_minutes is not None:
        timer_minutes = max(1, min(60, int(timer_minutes)))

    code = _rand_collab_code(db)

    s = CollabSessionModel(
        class_id=payload.class_id,
        session_code=code,
        title=title,
        state="lobby",
        room_count=room_count,
        timer_minutes=timer_minutes,
    )
    db.add(s)
    db.commit()
    db.refresh(s)

    return {"session_code": s.session_code, "join_url": None}


@app.get("/collab/{code}/status", response_model=CollabStatusResponse)
def collab_status(code: str, db: Session = Depends(get_db)):
    s = db.query(CollabSessionModel).filter(CollabSessionModel.session_code == code).first()
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")

    participants = (
        db.query(CollabParticipantModel)
        .filter(CollabParticipantModel.session_id == s.id)
        .all()
    )

    joined_count = len(participants)
    assigned_count = len([p for p in participants if p.room_number is not None])

    return {
        "session_code": s.session_code,
        "title": s.title,
        "state": s.state,
        "room_count": s.room_count,
        "timer_minutes": s.timer_minutes,
        "time_left_seconds": _collab_time_left_seconds(s),
        "joined_count": joined_count,
        "assigned_count": assigned_count,
    }


@app.post("/collab/{code}/join", response_model=CollabJoinResponse)
def collab_join(code: str, payload: CollabJoinPayload, db: Session = Depends(get_db)):
    s = db.query(CollabSessionModel).filter(CollabSessionModel.session_code == code).first()
    if not s:
        raise HTTPException(status_code=404, detail="Session not found")

    anon_id = (payload.anon_id or "").strip()
    name = (payload.name or "").strip()

    if not name and not anon_id:
        raise HTTPException(status_code=400, detail="Name required")

    existing = None
    if anon_id:
        existing = (
            db.query(CollabParticipantModel)
            .filter(CollabParticipantModel.session_id == s.id)
            .filter(CollabParticipantModel.anon_id == anon_id)
            .first()
        )

    if existing:
        existing.last_seen_at = datetime.utcnow()
        existing.is_online = True
        if name:
            existing.name = name
        db.commit()
        db.refresh(existing)
        return {
            "anon_id": existing.anon_id,
            "name": existing.name,
            "room_number": existing.room_number,
        }

    if not name:
        raise HTTPException(status_code=400, detail="Name required")

    anon_id = anon_id or f"cp_{uuid4().hex[:12]}"

    p = CollabParticipantModel(
        session_id=s.id,
        anon_id=anon_id,
        name=name,
        room_number=None,
        joined_at=datetime.utcnow(),
        last_seen_at=datetime.utcnow(),
        is_online=True,
    )
    db.add(p)
    db.commit()
    db.refresh(p)

    return {
        "anon_id": p.anon_id,
        "name": p.name,
        "room_number": p.room_number,
    }


@app.get("/collab/{code}/participants")
def collab_participants(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_collab_session_or_404(code, db, user)

    participants = (
        db.query(CollabParticipantModel)
        .filter(CollabParticipantModel.session_id == s.id)
        .order_by(CollabParticipantModel.joined_at.asc())
        .all()
    )

    return {
        "participants": [
            {
                "id": p.id,
                "anon_id": p.anon_id,
                "name": p.name,
                "room_number": p.room_number,
                "is_online": bool(p.is_online),
            }
            for p in participants
        ]
    }


@app.post("/collab/{code}/assignments")
def collab_assignments(
    code: str,
    payload: CollabAssignmentsPayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_collab_session_or_404(code, db, user)

    for item in payload.assignments:
        p = (
            db.query(CollabParticipantModel)
            .filter(CollabParticipantModel.session_id == s.id)
            .filter(CollabParticipantModel.id == item.participant_id)
            .first()
        )
        if not p:
            continue

        room_number = item.room_number
        if room_number is not None:
            room_number = max(1, min(12, int(room_number)))
            if room_number > int(s.room_count):
                room_number = int(s.room_count)

        p.room_number = room_number

    db.commit()
    return {"message": "ok"}


@app.post("/collab/{code}/start")
async def collab_start(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_collab_session_or_404(code, db, user)

    # Copy the current teacher board into every breakout room
    await _seed_breakout_rooms_from_teacher(code, int(s.room_count or 0))

    s.state = "live"
    s.started_at = s.started_at or datetime.utcnow()
    s.breakout_started_at = datetime.utcnow()
    db.commit()
    return {"message": "started"}


@app.post("/collab/{code}/end")
def collab_end(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_collab_session_or_404(code, db, user)

    s.state = "review"
    db.commit()
    return {"message": "review"}


@app.post("/collab/{code}/end-session")
def collab_end_session(
    code: str,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = get_owned_collab_session_or_404(code, db, user)

    s.state = "ended"
    s.ended_at = datetime.utcnow()
    db.commit()

    _clear_collab_session_history(code)

    return {"message": "ended"}

# =========================================================
# SEED CLASSES (optional)
# =========================================================
def seed_classes(db: Session):
    # ONLY seed if database is empty
    if db.query(ClassModel).count() > 0:
        return

    defaults = [
        ("6th Year Physics", "Physics"),
        ("6th Year Maths", "Maths"),
        ("5th Year Physics", "Physics"),
        ("3rd Year Maths", "Maths"),
    ]

    for name, subject in defaults:
        db.add(ClassModel(name=name, subject=subject))

    db.commit()



@app.on_event("startup")

def on_startup():

    Base.metadata.create_all(bind=engine)

    db = SessionLocal()

    try:

        seed_classes(db)

        _backfill_class_access_details(db)

    finally:

        db.close()

# =========================================================
# CLASSES
# =========================================================
@app.get("/classes")
def get_classes(
    archived: bool = False,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    return (
        db.query(ClassModel)
        .filter(
            ClassModel.owner_user_id == user.id,
            ClassModel.is_archived == archived,
        )
        .all()
    )


@app.get("/classes/archived", response_model=List[schemas.ClassOut])
def get_archived_classes(
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    return (
        db.query(ClassModel)
        .filter(
            ClassModel.owner_user_id == user.id,
            ClassModel.is_archived == True,
        )
        .order_by(ClassModel.archived_at.desc().nullslast(), ClassModel.id.desc())
        .all()
    )


@app.post("/classes", response_model=schemas.ClassOut)
def create_class(
    new_class: schemas.ClassCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    c = ClassModel(
        owner_user_id=user.id,
        name=new_class.name,
        subject=new_class.subject,
        stream=(new_class.stream or "").strip() or None,
        color=(new_class.color or "").strip() or None,
        preferred_exam_subject=(new_class.preferred_exam_subject or "").strip() or None,
        class_code=_rand_class_code(db),
        class_pin=_rand_class_pin(),
    )
    db.add(c)
    db.commit()
    db.refresh(c)

    # auto-create student access token
    _get_or_create_active_student_access_link(c.id, db)

    return c


@app.post("/classes/demo", response_model=schemas.ClassOut)
def create_demo_class(
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    return _seed_demo_class(db, user)


@app.post("/whiteboard/save")
async def save_whiteboard(
    class_id: int = Form(...),
    title: str = Form(...),
    image: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)

    os.makedirs("uploads/whiteboards", exist_ok=True)

    ext = os.path.splitext(image.filename or "")[1].lower()
    if ext not in [".png", ".jpg", ".jpeg", ".webp", ""]:
        ext = ".png"

    filename = f"whiteboard_{class_id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex}{ext}"
    disk_path = os.path.join("uploads", "whiteboards", filename)

    # Save file
    contents = await image.read()
    with open(disk_path, "wb") as f:
        f.write(contents)

    # Create a post on the class feed with a link to the saved image
    url_path = f"/uploads/whiteboards/{filename}"
    post = PostModel(
        class_id=class_id,
        author="Whiteboard",
        content=f"Whiteboard saved: {title}",
        links=json.dumps([url_path]),
    )
    db.add(post)
    db.commit()
    db.refresh(post)

    return {
        "id": post.id,
        "class_id": post.class_id,
        "author": post.author,
        "content": post.content,
        "links": [f"/posts/{post.id}/attachments/0"],
        "createdAt": getattr(post, "created_at", None),
    }


@app.post("/whiteboards/{whiteboard_id}/link-note", response_model=schemas.NoteOut)
def link_whiteboard_note(
    whiteboard_id: int,
    payload: schemas.WhiteboardNoteLinkPayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(payload.class_id, db, user)

    whiteboard = (
        db.query(models.WhiteboardStateModel)
        .filter(
            models.WhiteboardStateModel.id == whiteboard_id,
            models.WhiteboardStateModel.owner_user_id == user.id,
            models.WhiteboardStateModel.class_id == payload.class_id,
        )
        .first()
    )
    if not whiteboard:
        raise HTTPException(status_code=404, detail="Whiteboard not found")

    post = (
        db.query(PostModel)
        .filter(
            PostModel.id == payload.post_id,
            PostModel.class_id == payload.class_id,
        )
        .first()
    )
    if not post:
        raise HTTPException(status_code=404, detail="Saved whiteboard post not found")

    attachment_path, attachment_name = _post_attachment_path_or_404(post, 0)
    whiteboard.preview_image_path = str(attachment_path)

    topic = _get_or_create_saved_whiteboards_topic(payload.class_id, db)

    note = (
        db.query(models.Note)
        .filter(
            models.Note.class_id == payload.class_id,
            models.Note.whiteboard_state_id == whiteboard_id,
        )
        .first()
    )
    if not note:
        note = models.Note(
            class_id=payload.class_id,
            topic_id=topic.id,
            filename=_safe_whiteboard_note_filename(whiteboard.title, Path(attachment_name).suffix),
            stored_path=str(attachment_path),
            whiteboard_state_id=whiteboard_id,
        )
        db.add(note)
    else:
        note.topic_id = topic.id
        note.filename = _safe_whiteboard_note_filename(whiteboard.title, Path(attachment_name).suffix)
        note.stored_path = str(attachment_path)
        note.whiteboard_state_id = whiteboard_id

    db.commit()
    db.refresh(note)

    return schemas.NoteOut(
        id=note.id,
        class_id=note.class_id,
        topic_id=note.topic_id,
        filename=note.filename,
        file_url=f"/notes/{note.id}/download",
        whiteboard_state_id=note.whiteboard_state_id,
        uploaded_at=note.uploaded_at,
        topic_name=strip_prefix(topic.name),
    )


@app.post("/whiteboards", response_model=schemas.WhiteboardStateOut)
def save_whiteboard_state(
    payload: schemas.WhiteboardStateSave,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(payload.class_id, db, user)

    title = (payload.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Whiteboard title is required")

    try:
        state_json = json.dumps(payload.state)
    except (TypeError, ValueError):
        raise HTTPException(status_code=400, detail="Whiteboard state must be JSON serialisable")

    now = datetime.utcnow()

    if payload.whiteboard_id is not None:
        row = db.query(models.WhiteboardStateModel).filter(
            models.WhiteboardStateModel.id == payload.whiteboard_id
        ).first()
        if not row or row.owner_user_id != user.id:
            raise HTTPException(status_code=404, detail="Whiteboard not found")
        if row.class_id != payload.class_id:
            raise HTTPException(status_code=400, detail="Whiteboard does not belong to this class")
        row.title = title
        row.state_json = state_json
        row.updated_at = now
    else:
        row = models.WhiteboardStateModel(
            class_id=payload.class_id,
            owner_user_id=user.id,
            title=title,
            state_json=state_json,
            created_at=now,
            updated_at=now,
        )
        db.add(row)

    db.commit()
    db.refresh(row)

    return {
        "id": row.id,
        "class_id": row.class_id,
        "title": row.title,
        "created_at": row.created_at,
        "updated_at": row.updated_at,
        "state": payload.state,
    }


@app.get("/classes/{class_id}/whiteboards", response_model=schemas.WhiteboardStateListResponse)
def list_whiteboard_states(
    class_id: int,
    limit: int = 5,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    safe_limit = max(1, min(20, int(limit)))
    rows = (
        db.query(models.WhiteboardStateModel)
        .filter(
            models.WhiteboardStateModel.class_id == class_id,
            models.WhiteboardStateModel.owner_user_id == user.id,
        )
        .order_by(models.WhiteboardStateModel.updated_at.desc())
        .limit(safe_limit)
        .all()
    )
    return {"items": rows}


@app.get("/whiteboards/{whiteboard_id}", response_model=schemas.WhiteboardStateOut)
def get_whiteboard_state(
    whiteboard_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    row = db.query(models.WhiteboardStateModel).filter(
        models.WhiteboardStateModel.id == whiteboard_id,
        models.WhiteboardStateModel.owner_user_id == user.id,
    ).first()
    if not row:
        raise HTTPException(status_code=404, detail="Whiteboard not found")

    try:
        state = json.loads(row.state_json)
    except (TypeError, ValueError):
        raise HTTPException(status_code=500, detail="Stored whiteboard state is invalid")

    return {
        "id": row.id,
        "class_id": row.class_id,
        "title": row.title,
        "created_at": row.created_at,
        "updated_at": row.updated_at,
        "state": state,
    }

@app.get("/classes/{class_id}")
def get_class(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    return get_owned_class_or_404(class_id, db, user)


@app.get("/exam-library/items", response_model=List[schemas.ExamLibraryItemOut])
def list_exam_library_items(
    subject: Optional[str] = None,
    cycle: Optional[str] = None,
    level: Optional[str] = None,
    user: models.UserModel = Depends(get_current_user),
):
    del user  # authenticated teacher access only

    def _matches(value: str, expected: Optional[str]) -> bool:
        if not expected:
            return True
        return value.strip().lower() == expected.strip().lower()

    items = [
        item
        for item in _read_exam_library_manifest()
        if _matches(item.subject, subject) and _matches(item.cycle, cycle) and _matches(item.level, level)
    ]
    return items


@app.get("/exam-library/items/{item_id}", response_model=schemas.ExamLibraryItemOut)
def get_exam_library_item(
    item_id: str,
    user: models.UserModel = Depends(get_current_user),
):
    del user
    return _find_exam_library_item_or_404(item_id)


@app.get("/exam-library/items/{item_id}/download")
def download_exam_library_item(
    item_id: str,
    user: models.UserModel = Depends(get_current_user),
):
    del user
    item = _find_exam_library_item_or_404(item_id)
    full_path = (EXAM_LIBRARY_DIR / item.path).resolve()
    filename = Path(item.path).name
    return FileResponse(path=full_path, filename=filename, media_type="application/pdf")

@app.put("/classes/{class_id}")
def update_class(
    class_id: int,
    payload: dict,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    cls = get_owned_class_or_404(class_id, db, user)

    if "name" in payload and isinstance(payload["name"], str):
        cls.name = payload["name"].strip() or cls.name

    if "subject" in payload and isinstance(payload["subject"], str):
        cls.subject = payload["subject"].strip() or cls.subject

    if "stream" in payload and isinstance(payload["stream"], str):
        cls.stream = payload["stream"].strip() or None

    if "color" in payload and isinstance(payload["color"], str):
        cls.color = payload["color"].strip() or None

    if "preferred_exam_subject" in payload and isinstance(payload["preferred_exam_subject"], str):
        cls.preferred_exam_subject = payload["preferred_exam_subject"].strip() or None

    db.commit()
    db.refresh(cls)

    return cls


@app.post("/classes/{class_id}/archive")
def archive_class(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    cls = db.query(ClassModel).filter(
        ClassModel.id == class_id,
        ClassModel.owner_user_id == user.id,
    ).first()

    if not cls:
        raise HTTPException(status_code=404, detail="Class not found")

    if cls.is_archived:
        return cls

    archived_count = (
        db.query(ClassModel)
        .filter(
            ClassModel.owner_user_id == user.id,
            ClassModel.is_archived == True,
        )
        .count()
    )
    if archived_count >= 20:
        raise HTTPException(
            status_code=400,
            detail="You have reached the limit of 20 archived classes. Delete one permanently before archiving another.",
        )

    cls.is_archived = True
    cls.archived_at = datetime.utcnow()
    db.commit()
    db.refresh(cls)
    return cls

# =========================================================
# POSTS (links stored as JSON string)
# =========================================================
def _links_to_list(v):
    if v is None:
        return []
    if isinstance(v, list):
        return [str(x) for x in v if str(x).strip()]
    if isinstance(v, str):
        s = v.strip()
        if not s:
            return []
        try:
            parsed = json.loads(s)
            if isinstance(parsed, list):
                return [str(x) for x in parsed if str(x).strip()]
        except Exception:
            pass
        parts = re.split(r"[\n,]+", s)
        return [p.strip() for p in parts if p.strip()]
    return [str(v)]


def _links_to_storage(v):
    return json.dumps(_links_to_list(v), ensure_ascii=False)


class PostUpdatePayload(BaseModel):
    content: str = ""
    links: list[str] | str | None = None


@app.get("/classes/{class_id}/posts")
def get_posts(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    posts = (
        db.query(PostModel)
        .filter(PostModel.class_id == class_id)
        .order_by(PostModel.id.desc())
        .all()
    )
    out = []
    for p in posts:
        links = _links_to_list(getattr(p, "links", None))
        rewritten_links = []
        for idx, link in enumerate(links):
            if _internal_post_upload_relpath_or_none(link):
                rewritten_links.append(f"/posts/{p.id}/attachments/{idx}")
            else:
                rewritten_links.append(link)
        out.append(
            {
                "id": p.id,
                "class_id": p.class_id,
                "author": getattr(p, "author", ""),
                "content": getattr(p, "content", ""),
                "links": rewritten_links,
                "createdAt": getattr(p, "created_at", None).isoformat() if getattr(p, "created_at", None) else None,
            }
        )
    return out


@app.post("/classes/{class_id}/posts")
async def add_post(
    class_id: int,
    author: str = Form("Teacher"),
    content: str = Form(""),
    links: str = Form("[]"),
    files: list[UploadFile] = File([]),
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)

    content = (content or "").strip()
    author = (author or "Teacher").strip() or "Teacher"

    raw_links = _links_to_list(links)

    upload_dir = Path("uploads") / "posts"
    upload_dir.mkdir(parents=True, exist_ok=True)

    file_links: list[str] = []

    for f in files or []:
        if not f or not f.filename:
            continue

        original_name = Path(f.filename).name
        ext = Path(original_name).suffix
        stem = Path(original_name).stem

        safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._") or "file"
        unique_name = f"{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex}_{safe_stem}{ext}"
        dest_path = upload_dir / unique_name

        contents = await f.read()
        with open(dest_path, "wb") as out:
            out.write(contents)

        file_links.append(f"/uploads/posts/{unique_name}")

    all_links = raw_links + file_links

    if not content and not all_links:
        raise HTTPException(status_code=400, detail="Post must contain text, a link, or an attachment")

    p = PostModel(
        class_id=class_id,
        author=author,
        content=content,
        links=_links_to_storage(all_links),
    )
    db.add(p)
    db.commit()
    db.refresh(p)

    links = _links_to_list(getattr(p, "links", None))
    rewritten_links = []
    for idx, link in enumerate(links):
        if _internal_post_upload_relpath_or_none(link):
            rewritten_links.append(f"/posts/{p.id}/attachments/{idx}")
        else:
            rewritten_links.append(link)

    return {
        "id": p.id,
        "class_id": p.class_id,
        "author": p.author,
        "content": p.content,
        "links": rewritten_links,
        "createdAt": getattr(p, "created_at", None).isoformat() if getattr(p, "created_at", None) else None,
    }


@app.put("/posts/{post_id}")
def update_post(
    post_id: int,
    payload: PostUpdatePayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    post = db.query(PostModel).filter(PostModel.id == post_id).first()
    if not post:
        raise HTTPException(status_code=404, detail="Post not found")
    get_owned_class_or_404(post.class_id, db, user)

    content = (payload.content or "").strip()
    links = _links_to_list(payload.links)

    existing_links = _links_to_list(getattr(post, "links", None))
    kept_links: list[str] = []
    for link in links:
        if link.startswith("/posts/"):
            match = re.match(r"^/posts/\d+/attachments/(\d+)$", link.strip())
            if match:
                idx = int(match.group(1))
                if 0 <= idx < len(existing_links):
                    kept_links.append(existing_links[idx])
                    continue
        kept_links.append(link)

    if not content and not kept_links:
        raise HTTPException(status_code=400, detail="Post must contain text or a link")

    post.content = content
    post.links = _links_to_storage(kept_links)
    db.commit()
    db.refresh(post)

    final_links = _links_to_list(getattr(post, "links", None))
    rewritten_links = []
    for idx, link in enumerate(final_links):
        if _internal_post_upload_relpath_or_none(link):
            rewritten_links.append(f"/posts/{post.id}/attachments/{idx}")
        else:
            rewritten_links.append(link)

    return {
        "id": post.id,
        "class_id": post.class_id,
        "author": getattr(post, "author", ""),
        "content": getattr(post, "content", ""),
        "links": rewritten_links,
        "createdAt": getattr(post, "created_at", None).isoformat() if getattr(post, "created_at", None) else None,
    }


@app.delete("/posts/{post_id}")
def delete_post(
    post_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    post = db.query(PostModel).filter(PostModel.id == post_id).first()
    if not post:
        raise HTTPException(status_code=404, detail="Post not found")
    get_owned_class_or_404(post.class_id, db, user)
    db.delete(post)
    db.commit()
    return {"message": "Post deleted"}


@app.get("/posts/{post_id}/attachments/{attachment_index}")
def download_post_attachment(
    post_id: int,
    attachment_index: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    post = db.query(PostModel).filter(PostModel.id == post_id).first()
    if not post:
        raise HTTPException(status_code=404, detail="Post not found")
    get_owned_class_or_404(post.class_id, db, user)
    path, filename = _post_attachment_path_or_404(post, attachment_index)
    return FileResponse(path=path, filename=filename)


@app.get("/student/{token}/posts/{post_id}/attachments/{attachment_index}")
def student_download_post_attachment(
    token: str,
    post_id: int,
    attachment_index: int,
    db: Session = Depends(get_db),
):
    link = get_active_student_access_link_or_404(token, db)
    post = db.query(PostModel).filter(
        PostModel.id == post_id,
        PostModel.class_id == link.class_id,
    ).first()
    if not post:
        raise HTTPException(status_code=404, detail="Post not found")
    path, filename = _post_attachment_path_or_404(post, attachment_index)
    return FileResponse(path=path, filename=filename)


# =========================================================
# NOTES + TOPICS (Notes vs Exam Papers via prefix)
# =========================================================
EXAM_PREFIX = "EXAM: "
NOTES_PREFIX = "NOTES: "


def strip_prefix(name: str) -> str:
    if name.startswith(EXAM_PREFIX):
        return name[len(EXAM_PREFIX) :]
    if name.startswith(NOTES_PREFIX):
        return name[len(NOTES_PREFIX) :]
    return name


@app.get("/topics/{class_id}", response_model=List[schemas.TopicOut])
def list_topics(
    class_id: int,
    kind: str = "notes",
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    q = db.query(models.Topic).filter(models.Topic.class_id == class_id)
    if kind == "exam":
        q = q.filter(models.Topic.name.startswith(EXAM_PREFIX))
    else:
        q = q.filter(~models.Topic.name.startswith(EXAM_PREFIX))

    topics = q.order_by(models.Topic.name).all()
    return [schemas.TopicOut(id=t.id, class_id=t.class_id, name=strip_prefix(t.name)) for t in topics]


@app.post("/topics", response_model=schemas.TopicOut)
def create_topic(
    payload: schemas.TopicCreate,
    kind: str = "notes",
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(payload.class_id, db, user)
    prefix = EXAM_PREFIX if kind == "exam" else NOTES_PREFIX
    name = (payload.name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="Topic name cannot be empty")

    topic = models.Topic(class_id=payload.class_id, name=f"{prefix}{name}")
    db.add(topic)
    db.commit()
    db.refresh(topic)
    return schemas.TopicOut(id=topic.id, class_id=topic.class_id, name=name)


@app.delete("/topics/{topic_id}")
def delete_topic(
    topic_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    topic = db.query(models.Topic).filter(models.Topic.id == topic_id).first()
    if not topic:
        raise HTTPException(status_code=404, detail="Topic not found")
    get_owned_class_or_404(topic.class_id, db, user)

    # delete notes belonging to this topic
    notes = db.query(models.Note).filter(models.Note.topic_id == topic_id).all()
    for n in notes:
        if n.stored_path and os.path.exists(n.stored_path):
            try:
                os.remove(n.stored_path)
            except Exception:
                pass

    db.query(models.Note).filter(models.Note.topic_id == topic_id).delete()
    db.delete(topic)
    db.commit()
    return {"message": "Topic deleted"}


@app.get("/notes/{class_id}", response_model=List[schemas.NoteOut])
def list_notes(
    class_id: int,
    kind: str = "notes",
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    q = (
        db.query(models.Note, models.Topic)
        .join(models.Topic, models.Note.topic_id == models.Topic.id)
        .filter(models.Note.class_id == class_id)
    )

    if kind == "exam":
        q = q.filter(models.Topic.name.startswith(EXAM_PREFIX))
    else:
        q = q.filter(~models.Topic.name.startswith(EXAM_PREFIX))

    rows = q.order_by(models.Note.id.desc()).all()

    out: List[schemas.NoteOut] = []
    for note, topic in rows:
        out.append(
            schemas.NoteOut(
                id=note.id,
                class_id=note.class_id,
                topic_id=note.topic_id,
                filename=note.filename,
                file_url=f"/notes/{note.id}/download",
                whiteboard_state_id=note.whiteboard_state_id,
                uploaded_at=note.uploaded_at,
                topic_name=strip_prefix(topic.name) if topic else "Unsorted",
            )
        )
    return out


@app.post("/notes/upload", response_model=schemas.NoteOut)
def upload_note(
    class_id: int = Form(...),
    topic_id: int = Form(...),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    topic = db.query(models.Topic).filter(models.Topic.id == topic_id).first()
    if not topic:
        raise HTTPException(status_code=404, detail="Topic not found")
    if topic.class_id != class_id:
        raise HTTPException(status_code=400, detail="Topic does not belong to this class")

    dest_dir = UPLOADS_DIR / "notes" / str(class_id)
    dest_dir.mkdir(parents=True, exist_ok=True)

    safe_name = (file.filename or "file.pdf").replace("\\", "/").split("/")[-1]
    dest_path = dest_dir / f"{int(datetime.utcnow().timestamp())}_{safe_name}"

    with dest_path.open("wb") as f:
        shutil.copyfileobj(file.file, f)

    n = models.Note(
        class_id=class_id,
        topic_id=topic_id,
        filename=safe_name,
        stored_path=str(dest_path),
    )
    db.add(n)
    db.commit()
    db.refresh(n)

    return schemas.NoteOut(
        id=n.id,
        class_id=n.class_id,
        topic_id=n.topic_id,
        filename=n.filename,
        file_url=f"/notes/{n.id}/download",
        whiteboard_state_id=n.whiteboard_state_id,
        uploaded_at=n.uploaded_at,
        topic_name=strip_prefix(topic.name),
    )


@app.delete("/notes/{note_id}")
def delete_note(
    note_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    n = db.query(models.Note).filter(models.Note.id == note_id).first()
    if not n:
        raise HTTPException(status_code=404, detail="Note not found")
    get_owned_class_or_404(n.class_id, db, user)

    if n.stored_path and os.path.exists(n.stored_path):
        try:
            os.remove(n.stored_path)
        except Exception:
            pass

    db.delete(n)
    db.commit()
    return {"message": "Note deleted"}


@app.get("/notes/{note_id}/download")
def download_note(
    note_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    note = db.query(models.Note).filter(models.Note.id == note_id).first()
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    get_owned_class_or_404(note.class_id, db, user)

    resolved_path = _resolve_stored_upload_path_or_none(note.stored_path)
    if not resolved_path:
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=str(resolved_path), filename=note.filename)


@app.get("/student/{token}/notes/{note_id}/download")
def student_download_note(
    token: str,
    note_id: int,
    db: Session = Depends(get_db),
):
    link = get_active_student_access_link_or_404(token, db)
    note = db.query(models.Note).filter(
        models.Note.id == note_id,
        models.Note.class_id == link.class_id,
    ).first()
    resolved_path = _resolve_stored_upload_path_or_none(note.stored_path)
    if not resolved_path:
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=str(resolved_path), filename=note.filename)


# =========================================================
# TESTS (Categories + PDF uploads)  — matches schemas.py
# =========================================================
@app.get("/classes/{class_id}/test-categories", response_model=List[schemas.TestCategoryOut])
def list_test_categories(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    cats = (
        db.query(TestCategory)
        .filter(TestCategory.class_id == class_id)
        .order_by(TestCategory.id.desc())
        .all()
    )
    return [
        schemas.TestCategoryOut(
            id=c.id,
            class_id=c.class_id,
            title=c.title,
            description=c.description,
        )
        for c in cats
    ]


@app.post("/classes/{class_id}/test-categories", response_model=schemas.TestCategoryOut)
def create_test_category(
    class_id: int,
    payload: schemas.TestCategoryCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    title = (payload.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Category title cannot be empty")

    c = TestCategory(
        class_id=class_id,
        title=title,
        description=(payload.description or "").strip() or None,
    )
    db.add(c)
    db.commit()
    db.refresh(c)

    return schemas.TestCategoryOut(
        id=c.id,
        class_id=c.class_id,
        title=c.title,
        description=c.description,
    )


@app.put("/test-categories/{cat_id}", response_model=schemas.TestCategoryOut)
def update_test_category(
    cat_id: int,
    payload: schemas.TestCategoryPatch,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    c = db.query(TestCategory).filter(TestCategory.id == cat_id).first()
    if not c:
        raise HTTPException(status_code=404, detail="Category not found")
    get_owned_class_or_404(c.class_id, db, user)

    if payload.title is not None:
        t = payload.title.strip()
        if not t:
            raise HTTPException(status_code=400, detail="Category title cannot be empty")
        c.title = t

    if payload.description is not None:
        c.description = payload.description.strip() or None

    db.commit()
    db.refresh(c)

    return schemas.TestCategoryOut(
        id=c.id,
        class_id=c.class_id,
        title=c.title,
        description=c.description,
    )


@app.delete("/test-categories/{cat_id}")
def delete_test_category(
    cat_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    c = db.query(TestCategory).filter(TestCategory.id == cat_id).first()
    if not c:
        raise HTTPException(status_code=404, detail="Category not found")
    get_owned_class_or_404(c.class_id, db, user)

    # detach tests from category
    db.query(TestItem).filter(TestItem.category_id == cat_id).update({"category_id": None})
    db.delete(c)
    db.commit()
    return {"message": "Category deleted"}


@app.get("/classes/{class_id}/tests", response_model=List[schemas.TestOut])
def list_tests(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    tests = (
        db.query(TestItem)
        .filter(TestItem.class_id == class_id)
        .order_by(TestItem.id.desc())
        .all()
    )

    out: List[schemas.TestOut] = []
    for t in tests:
        out.append(
            schemas.TestOut(
                id=t.id,
                class_id=t.class_id,
                category_id=t.category_id,
                title=t.title,
                description=t.description,
                filename=t.filename,
                file_url=f"/tests/{t.id}/download",
                uploaded_at=t.uploaded_at,
            )
        )
    return out

@app.delete("/classes/{class_id}")
def delete_class(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    c = db.query(ClassModel).filter(
        ClassModel.id == class_id,
        ClassModel.owner_user_id == user.id
    ).first()

    if not c:
        raise HTTPException(status_code=404, detail="Class not found")

    try:
        _delete_class_dependencies(db, [c.id], [c])
        db.commit()
        return {"ok": True}
    except HTTPException:
        db.rollback()
        raise
    except Exception:
        db.rollback()
        logger.exception("Failed to permanently delete class %s", class_id)
        raise HTTPException(status_code=500, detail="Failed to permanently delete class")

@app.post("/tests", response_model=schemas.TestOut)
def upload_test(
    class_id: int = Form(...),
    title: str = Form(...),
    description: str = Form(""),
    category_id: Optional[int] = Form(None),
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    title = (title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title cannot be empty")
    if category_id is not None:
        category = db.query(TestCategory).filter(TestCategory.id == category_id).first()
        if not category:
            raise HTTPException(status_code=404, detail="Category not found")
        if category.class_id != class_id:
            raise HTTPException(status_code=400, detail="Category does not belong to this class")

    dest_dir = UPLOADS_DIR / "tests" / str(class_id)
    dest_dir.mkdir(parents=True, exist_ok=True)

    safe_name = (file.filename or "test.pdf").replace("\\", "/").split("/")[-1]
    dest_path = dest_dir / f"{int(datetime.utcnow().timestamp())}_{safe_name}"

    with dest_path.open("wb") as f:
        shutil.copyfileobj(file.file, f)

    t = TestItem(
        class_id=class_id,
        category_id=category_id,
        title=title,
        description=(description or "").strip() or None,
        filename=safe_name,
        stored_path=str(dest_path),
    )
    db.add(t)
    db.commit()
    db.refresh(t)

    return schemas.TestOut(
        id=t.id,
        class_id=t.class_id,
        category_id=t.category_id,
        title=t.title,
        description=t.description,
        filename=t.filename,
        file_url=f"/tests/{t.id}/download",
        uploaded_at=t.uploaded_at,
    )


@app.put("/tests/{test_id}", response_model=schemas.TestOut)
def update_test(
    test_id: int,
    payload: schemas.TestPatch,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    t = db.query(TestItem).filter(TestItem.id == test_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Test not found")
    get_owned_class_or_404(t.class_id, db, user)

    if payload.title is not None:
        new_title = payload.title.strip()
        if not new_title:
            raise HTTPException(status_code=400, detail="Title cannot be empty")
        t.title = new_title

    if payload.description is not None:
        t.description = payload.description.strip() or None

    if payload.category_id is not None or payload.category_id is None:
        # allow explicit null to clear category
        if payload.category_id is not None:
            category = db.query(TestCategory).filter(TestCategory.id == payload.category_id).first()
            if not category:
                raise HTTPException(status_code=404, detail="Category not found")
            if category.class_id != t.class_id:
                raise HTTPException(status_code=400, detail="Category does not belong to this class")
        t.category_id = payload.category_id

    db.commit()
    db.refresh(t)

    return schemas.TestOut(
        id=t.id,
        class_id=t.class_id,
        category_id=t.category_id,
        title=t.title,
        description=t.description,
        filename=t.filename,
        file_url=f"/tests/{t.id}/download",
        uploaded_at=t.uploaded_at,
    )


@app.delete("/tests/{test_id}")
def delete_test(
    test_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    t = db.query(TestItem).filter(TestItem.id == test_id).first()
    if not t:
        raise HTTPException(status_code=404, detail="Test not found")
    get_owned_class_or_404(t.class_id, db, user)

    if t.stored_path and os.path.exists(t.stored_path):
        try:
            os.remove(t.stored_path)
        except Exception:
            pass

    db.delete(t)
    db.commit()
    return {"message": "Test deleted"}


@app.get("/tests/{test_id}/download")
def download_test(
    test_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    test = db.query(TestItem).filter(TestItem.id == test_id).first()
    if not test:
        raise HTTPException(status_code=404, detail="Test not found")
    get_owned_class_or_404(test.class_id, db, user)
    if not test.stored_path or not os.path.exists(test.stored_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=test.stored_path, filename=test.filename)


@app.get("/student/{token}/tests/{test_id}/download")
def student_download_test(
    token: str,
    test_id: int,
    db: Session = Depends(get_db),
):
    link = get_active_student_access_link_or_404(token, db)
    test = db.query(TestItem).filter(
        TestItem.id == test_id,
        TestItem.class_id == link.class_id,
    ).first()
    if not test:
        raise HTTPException(status_code=404, detail="Test not found")
    if not test.stored_path or not os.path.exists(test.stored_path):
        raise HTTPException(status_code=404, detail="File not found")
    return FileResponse(path=test.stored_path, filename=test.filename)

# -------------------------
# Calendar Routes (single canonical source of truth)
# - class_id = NULL => global event
# - class_id = <int> => class event
# -------------------------

@app.get("/calendar-events", response_model=list[schemas.CalendarEventOut])
def list_calendar_events(
    class_id: Optional[int] = None,
    global_only: bool = False,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    # Base query: ONLY this user's events
    q = db.query(models.CalendarEvent).filter(
        models.CalendarEvent.owner_user_id == user.id
    )

    # Global only (personal/global events for this user)
    if global_only:
        return q.filter(models.CalendarEvent.class_id.is_(None)).all()

    # If class_id provided, return (global + this class) for this user
    if class_id is not None:
        # ✅ Optional hardening: ensure the class belongs to this user
        cls = db.query(ClassModel).filter(
            ClassModel.id == class_id,
            ClassModel.owner_user_id == user.id
        ).first()
        if not cls:
            raise HTTPException(status_code=404, detail="Class not found")

        return q.filter(
            (models.CalendarEvent.class_id.is_(None))
            | (models.CalendarEvent.class_id == class_id)
        ).all()

    # Otherwise: all events for this user
    return q.all()


# Backwards-compatible endpoint (used by older pages)
@app.get("/classes/{class_id}/calendar-events", response_model=list[schemas.CalendarEventOut])
def get_calendar_events_for_class(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    # ✅ Ensure class belongs to this user
    cls = db.query(ClassModel).filter(
        ClassModel.id == class_id,
        ClassModel.owner_user_id == user.id
    ).first()
    if not cls:
        raise HTTPException(status_code=404, detail="Class not found")

    # Return only THIS user's events for that class
    return db.query(models.CalendarEvent).filter(
        models.CalendarEvent.owner_user_id == user.id,
        models.CalendarEvent.class_id == class_id
    ).all()


@app.post("/calendar-events", response_model=schemas.CalendarEventOut)
def create_calendar_event(
    event: schemas.CalendarEventCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    # ✅ Optional hardening: if class_id set, ensure it belongs to this user
    if event.class_id is not None:
        cls = db.query(ClassModel).filter(
            ClassModel.id == event.class_id,
            ClassModel.owner_user_id == user.id
        ).first()
        if not cls:
            raise HTTPException(status_code=404, detail="Class not found")

    new_event = models.CalendarEvent(**event.dict(), owner_user_id=user.id)
    db.add(new_event)
    db.commit()
    db.refresh(new_event)
    return new_event


@app.put("/calendar-events/{event_id}", response_model=schemas.CalendarEventOut)
def update_calendar_event(
    event_id: int,
    event: schemas.CalendarEventCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    # ✅ Only allow editing your own events
    db_event = db.query(models.CalendarEvent).filter(
        models.CalendarEvent.id == event_id,
        models.CalendarEvent.owner_user_id == user.id
    ).first()

    if not db_event:
        raise HTTPException(status_code=404, detail="Event not found")

    # ✅ Optional hardening: if moving to a class, ensure class belongs to this user
    if event.class_id is not None:
        cls = db.query(ClassModel).filter(
            ClassModel.id == event.class_id,
            ClassModel.owner_user_id == user.id
        ).first()
        if not cls:
            raise HTTPException(status_code=404, detail="Class not found")

    for key, value in event.dict().items():
        setattr(db_event, key, value)

    db.commit()
    db.refresh(db_event)
    return db_event


@app.delete("/calendar-events/{event_id}")
def delete_calendar_event(
    event_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    # ✅ Only allow deleting your own events
    db_event = db.query(models.CalendarEvent).filter(
        models.CalendarEvent.id == event_id,
        models.CalendarEvent.owner_user_id == user.id
    ).first()

    if not db_event:
        raise HTTPException(status_code=404, detail="Event not found")

    db.delete(db_event)
    db.commit()
    return {"message": "Deleted"}

@app.post("/student-access/{class_id}")
def create_student_access(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    cls = db.query(ClassModel).filter(
        ClassModel.id == class_id,
        ClassModel.owner_user_id == user.id
    ).first()

    if not cls:
        raise HTTPException(status_code=404, detail="Class not found")

    # deactivate existing links
    db.query(models.StudentAccessLink).filter(
        models.StudentAccessLink.class_id == class_id,
        models.StudentAccessLink.is_active == True
    ).update({"is_active": False})

    token = secrets.token_urlsafe(24)

    link = models.StudentAccessLink(
        class_id=class_id,
        token=token,
    )

    db.add(link)
    db.commit()
    db.refresh(link)

    return {"token": token}

@app.get("/student-access/{class_id}")
def get_student_access(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    link = db.query(models.StudentAccessLink).filter(
        models.StudentAccessLink.class_id == class_id,
        models.StudentAccessLink.is_active == True
    ).first()

    return {"token": link.token if link else None}


@app.get("/classes/{class_id}/student-access-code", response_model=schemas.ClassAccessOut)
def get_class_access_details(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    cls = get_owned_class_or_404(class_id, db, user)
    _ensure_class_access_details(cls, db)
    db.commit()
    db.refresh(cls)
    return _class_access_out(cls)


@app.post("/classes/{class_id}/regenerate-student-access-pin", response_model=schemas.ClassAccessOut)
def regenerate_class_access_pin(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    cls = get_owned_class_or_404(class_id, db, user)
    cls.class_pin = _rand_class_pin()
    db.add(cls)
    db.commit()
    db.refresh(cls)
    return _class_access_out(cls)


@app.post("/classes/{class_id}/regenerate-student-access-code", response_model=schemas.ClassAccessOut)
def regenerate_class_access_code(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    cls = get_owned_class_or_404(class_id, db, user)
    cls.class_code = _rand_class_code(db)
    db.add(cls)
    db.commit()
    db.refresh(cls)
    return _class_access_out(cls)


@app.post("/student/join/class", response_model=schemas.StudentJoinRedirectOut)
def join_class_by_code(
    payload: schemas.ClassJoinPayload,
    db: Session = Depends(get_db),
):
    code = _normalise_class_code(payload.code)
    pin = re.sub(r"\D", "", payload.pin or "")[:4]

    if len(code) != 6:
        raise HTTPException(status_code=400, detail="Please enter a valid 6-character class code.")
    if len(pin) != 4:
        raise HTTPException(status_code=400, detail="Please enter a valid 4-digit class PIN.")

    cls = db.query(ClassModel).filter(ClassModel.class_code == code).first()
    if not cls or (cls.class_pin or "").strip() != pin:
        raise HTTPException(status_code=400, detail="That class code or PIN is not valid.")

    link = _get_or_create_active_student_access_link(cls.id, db)
    return schemas.StudentJoinRedirectOut(
        ok=True,
        redirect_url=f"/student/{link.token}",
        message="Joined successfully.",
    )

@app.get("/student/{token}")
def student_view(token: str, db: Session = Depends(get_db)):
    link = get_active_student_access_link_or_404(token, db)

    cls = db.query(ClassModel).filter(ClassModel.id == link.class_id).first()
    if not cls:
        raise HTTPException(status_code=404, detail="Class not found")

    posts = (
        db.query(PostModel)
        .filter(PostModel.class_id == link.class_id)
        .order_by(PostModel.created_at.desc())
        .all()
    )

    notes = db.query(models.Note).filter(models.Note.class_id == link.class_id).all()
    tests = db.query(TestItem).filter(TestItem.class_id == link.class_id).all()

    return {
        "class_name": cls.name,
        "subject": cls.subject,
        "posts": [
    {
        "id": p.id,
        "author": p.author,
        "content": p.content,
        "links": [
            (
                f"/student/{token}/posts/{p.id}/attachments/{idx}"
                if _internal_post_upload_relpath_or_none(post_link)
                else post_link
            )
            for idx, post_link in enumerate(_links_to_list(getattr(p, "links", None)))
        ],
        "created_at": p.created_at.isoformat() if p.created_at else None,
    }
    for p in posts
],
        "notes": [
            {
                "id": n.id,
                "filename": n.filename,
                "file_url": f"/student/{token}/notes/{n.id}/download",
            }
            for n in notes
        ],
        "tests": [
            {
                "id": t.id,
                "title": t.title,
                "file_url": f"/student/{token}/tests/{t.id}/download",
            }
            for t in tests
        ],
    }


# =========================================================
# STUDENTS (first names only)
# =========================================================

@app.get("/classes/{class_id}/students")
def list_students(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    # include inactive too (teacher admin needs to see them)
    rows = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id)
        .order_by(StudentModel.id.desc())
        .all()
    )
    return rows


@app.post("/classes/{class_id}/students")
def create_student(
    class_id: int,
    payload: StudentCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    name = (payload.first_name or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="First name cannot be empty")

    s = StudentModel(
        class_id=class_id,
        first_name=name,
        notes=(payload.notes or "").strip() or None,
        active=True,
    )
    db.add(s)
    db.commit()
    db.refresh(s)
    return s


@app.put("/students/{student_id}")
def update_student(
    student_id: int,
    payload: StudentUpdate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = db.query(StudentModel).filter(StudentModel.id == student_id).first()
    if not s:
        raise HTTPException(status_code=404, detail="Student not found")
    get_owned_class_or_404(s.class_id, db, user)

    if payload.first_name is not None:
        name = payload.first_name.strip()
        if not name:
            raise HTTPException(status_code=400, detail="First name cannot be empty")
        s.first_name = name

    if payload.notes is not None:
        s.notes = payload.notes.strip() or None

    if payload.active is not None:
        s.active = bool(payload.active)

    db.commit()
    db.refresh(s)
    return s


@app.delete("/students/{student_id}")
def delete_student(
    student_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    s = db.query(StudentModel).filter(StudentModel.id == student_id).first()
    if not s:
        raise HTTPException(status_code=404, detail="Student not found")
    get_owned_class_or_404(s.class_id, db, user)

    db.query(AssessmentResultModel).filter(
        AssessmentResultModel.student_id == student_id
    ).delete(synchronize_session=False)

    db.delete(s)
    db.commit()
    return {"ok": True, "student_id": student_id}

@app.post("/classes/{class_id}/students/bulk")
def create_students_bulk(
    class_id: int,
    payload: StudentBulkCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    cleaned = []
    seen = set()

    for n in (payload.names or []):
        name = (n or "").strip()
        if not name:
            continue
        key = name.lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned.append(name)

    if not cleaned:
        raise HTTPException(status_code=400, detail="No valid names provided")

    existing = db.query(StudentModel.first_name).filter(
        StudentModel.class_id == class_id
    ).all()

    existing_set = {(x[0] or "").strip().lower() for x in existing}

    to_create = []
    for name in cleaned:
        if name.lower() in existing_set:
            continue
        to_create.append(
            StudentModel(
                class_id=class_id,
                first_name=name,
                active=True,
            )
        )

    if not to_create:
        return []

    db.add_all(to_create)
    db.commit()

    for s in to_create:
        db.refresh(s)

    return to_create


# =========================================================
# ASSESSMENTS (class tests/results tracker) — separate from PDF "tests"
# =========================================================

@app.get("/classes/{class_id}/assessments")
def list_assessments(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    rows = (
        db.query(ClassAssessmentModel)
        .filter(ClassAssessmentModel.class_id == class_id)
        .order_by(ClassAssessmentModel.id.desc())
        .all()
    )
    return [
        {
            "id": a.id,
            "class_id": a.class_id,
            "title": a.title,
            "assessment_date": a.assessment_date.date().isoformat() if a.assessment_date else None,
        }
        for a in rows
    ]


@app.post("/classes/{class_id}/assessments")
def create_assessment(
    class_id: int,
    payload: AssessmentCreate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    title = (payload.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title cannot be empty")

    # parse optional YYYY-MM-DD
    if payload.assessment_date:
        try:
            dt = datetime.strptime(payload.assessment_date, "%Y-%m-%d")
        except Exception:
            raise HTTPException(status_code=400, detail="assessment_date must be YYYY-MM-DD")
    else:
        dt = datetime.utcnow()

    a = ClassAssessmentModel(class_id=class_id, title=title, assessment_date=dt)
    db.add(a)
    db.commit()
    db.refresh(a)

    return {
        "id": a.id,
        "class_id": a.class_id,
        "title": a.title,
        "assessment_date": a.assessment_date.date().isoformat() if a.assessment_date else None,
    }

@app.put("/assessments/{assessment_id}")
def update_assessment(
    assessment_id: int,
    payload: AssessmentUpdate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    a = db.query(ClassAssessmentModel).filter(ClassAssessmentModel.id == assessment_id).first()
    if not a:
        raise HTTPException(status_code=404, detail="Assessment not found")
    get_owned_class_or_404(a.class_id, db, user)

    title = (payload.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title cannot be empty")

    if payload.assessment_date:
        try:
            dt = datetime.strptime(payload.assessment_date, "%Y-%m-%d")
        except Exception:
            raise HTTPException(status_code=400, detail="assessment_date must be YYYY-MM-DD")
    else:
        dt = datetime.utcnow()

    a.title = title
    a.assessment_date = dt

    db.commit()
    db.refresh(a)

    return {
        "id": a.id,
        "class_id": a.class_id,
        "title": a.title,
        "assessment_date": a.assessment_date.date().isoformat() if a.assessment_date else None,
    }


@app.delete("/assessments/{assessment_id}")
def delete_assessment(
    assessment_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    a = db.query(ClassAssessmentModel).filter(ClassAssessmentModel.id == assessment_id).first()
    if not a:
        raise HTTPException(status_code=404, detail="Assessment not found")
    get_owned_class_or_404(a.class_id, db, user)

    # delete linked results first
    db.query(AssessmentResultModel).filter(
        AssessmentResultModel.assessment_id == assessment_id
    ).delete()

    # then delete the assessment
    db.delete(a)
    db.commit()

    return {"ok": True, "deleted_assessment_id": assessment_id}

@app.get("/assessments/{assessment_id}/results")
def get_assessment_results(
    assessment_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    a = db.query(ClassAssessmentModel).filter(ClassAssessmentModel.id == assessment_id).first()
    if not a:
        raise HTTPException(status_code=404, detail="Assessment not found")
    get_owned_class_or_404(a.class_id, db, user)

    # Use ACTIVE students by default for entry grid
    students = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == a.class_id)
        .filter(StudentModel.active == True)  # noqa: E712
        .order_by(StudentModel.id.desc())
        .all()
    )

    # existing results keyed by student_id
    existing = (
        db.query(AssessmentResultModel)
        .filter(AssessmentResultModel.assessment_id == assessment_id)
        .all()
    )
    by_student = {r.student_id: r for r in existing}

    rows = []
    for s in students:
        r = by_student.get(s.id)
        rows.append(
            {
                "student_id": s.id,
                "first_name": s.first_name,
                "score_percent": r.score_percent if r else None,
                "absent": bool(r.absent) if r else False,
            }
        )

    return {
        "assessment": {
            "id": a.id,
            "class_id": a.class_id,
            "title": a.title,
            "assessment_date": a.assessment_date.date().isoformat() if a.assessment_date else None,
        },
        "results": rows,
    }


@app.put("/assessments/{assessment_id}/results")
def upsert_assessment_results(
    assessment_id: int,
    payload: BulkResultsUpdate,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    a = db.query(ClassAssessmentModel).filter(ClassAssessmentModel.id == assessment_id).first()
    if not a:
        raise HTTPException(status_code=404, detail="Assessment not found")
    get_owned_class_or_404(a.class_id, db, user)

    # Build map of existing rows for this assessment
    existing = (
        db.query(AssessmentResultModel)
        .filter(AssessmentResultModel.assessment_id == assessment_id)
        .all()
    )
    by_student = {r.student_id: r for r in existing}

    for item in payload.results:
        # Ensure student belongs to the same class
        s = db.query(StudentModel).filter(StudentModel.id == item.student_id).first()
        if not s or s.class_id != a.class_id:
            raise HTTPException(status_code=400, detail=f"Student {item.student_id} not in this class")

        absent = bool(item.absent)

        # normalise score
        score = item.score_percent
        if absent:
            score = None
        else:
            if score is None:
                score = 0
            score = int(score)
            if score < 0 or score > 100:
                raise HTTPException(status_code=400, detail="score_percent must be 0 - 100")

        r = by_student.get(item.student_id)
        if r:
            r.absent = absent
            r.score_percent = score
        else:
            r = AssessmentResultModel(
                assessment_id=assessment_id,
                student_id=item.student_id,
                absent=absent,
                score_percent=score,
            )
            db.add(r)

    db.commit()

    return {"message": "Saved"}


@app.get("/classes/{class_id}/cat4/meta")
def cat4_meta(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)
    return _build_cat4_meta_payload(class_id, db)


@app.post("/classes/{class_id}/cat4/baselines")
def create_cat4_baseline_set(
    class_id: int,
    payload: Cat4BaselineSetCreatePayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)

    existing_locked = (
        db.query(Cat4BaselineSetModel)
        .filter(Cat4BaselineSetModel.class_id == class_id, Cat4BaselineSetModel.is_locked == True)  # noqa: E712
        .first()
    )
    if existing_locked:
        raise HTTPException(status_code=400, detail="CAT4 baseline is already locked for this class")

    title = (payload.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title required")

    baseline = Cat4BaselineSetModel(
        class_id=class_id,
        title=title,
        test_date=_parse_optional_date(payload.test_date),
        is_locked=False,
    )
    db.add(baseline)
    db.commit()
    db.refresh(baseline)

    return {
        "id": baseline.id,
        "class_id": baseline.class_id,
        "title": baseline.title,
        "test_date": baseline.test_date.date().isoformat() if baseline.test_date else None,
        "is_locked": bool(baseline.is_locked),
        "locked_at": baseline.locked_at.isoformat() if baseline.locked_at else None,
        "created_at": baseline.created_at.isoformat() if baseline.created_at else None,
    }


@app.post("/classes/{class_id}/cat4/baselines/{baseline_id}/rows")
def upsert_cat4_baseline_rows(
    class_id: int,
    baseline_id: int,
    payload: Cat4BaselineRowsPayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)

    baseline = (
        db.query(Cat4BaselineSetModel)
        .filter(Cat4BaselineSetModel.id == baseline_id, Cat4BaselineSetModel.class_id == class_id)
        .first()
    )
    if not baseline:
        raise HTTPException(status_code=404, detail="CAT4 baseline set not found")
    if baseline.is_locked:
        raise HTTPException(status_code=400, detail="CAT4 baseline is locked for this class")

    students = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id)
        .filter(StudentModel.active == True)  # noqa: E712
        .all()
    )
    name_index = _cat4_name_index(students)

    db.query(Cat4StudentBaselineModel).filter(
        Cat4StudentBaselineModel.baseline_set_id == baseline_id
    ).delete()

    matched_count = 0
    unmatched_count = 0
    created_rows = 0

    for item in payload.rows:
        raw_name = (item.raw_name or "").strip()
        if not raw_name:
            continue
        missing = [
            field
            for field in ["verbal_sas", "quantitative_sas", "non_verbal_sas", "spatial_sas", "overall_sas"]
            if getattr(item, field) is None
        ]
        if missing:
            raise HTTPException(status_code=400, detail=f"{raw_name}: missing required CAT4 fields {', '.join(missing)}")

        student_id, matched_name, match_note = _match_cat4_student_name(raw_name, name_index)
        row = Cat4StudentBaselineModel(
            baseline_set_id=baseline_id,
            class_id=class_id,
            student_id=student_id,
            raw_name=raw_name,
            matched_name=matched_name,
            verbal_sas=item.verbal_sas,
            quantitative_sas=item.quantitative_sas,
            non_verbal_sas=item.non_verbal_sas,
            spatial_sas=item.spatial_sas,
            overall_sas=item.overall_sas,
            profile_label=(item.profile_label or "").strip() or None,
            confidence_note=(item.confidence_note or "").strip() or match_note,
        )
        db.add(row)
        created_rows += 1
        if student_id:
            matched_count += 1
        else:
            unmatched_count += 1

    baseline.is_locked = True
    baseline.locked_at = datetime.utcnow()
    db.commit()

    return {
        "ok": True,
        "baseline_id": baseline_id,
        "created_rows": created_rows,
        "matched_count": matched_count,
        "unmatched_count": unmatched_count,
        "baseline_locked": True,
    }


@app.post("/classes/{class_id}/cat4/term-sets")
def create_cat4_term_set(
    class_id: int,
    payload: Cat4TermResultSetCreatePayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)

    title = (payload.title or "").strip()
    if not title:
        raise HTTPException(status_code=400, detail="Title required")

    term_set = Cat4TermResultSetModel(
        class_id=class_id,
        title=title,
        academic_year=(payload.academic_year or "").strip() or None,
        term_key=(payload.term_key or "").strip() or None,
    )
    db.add(term_set)
    db.commit()
    db.refresh(term_set)

    return {
        "id": term_set.id,
        "class_id": term_set.class_id,
        "title": term_set.title,
        "academic_year": term_set.academic_year,
        "term_key": term_set.term_key,
        "created_at": term_set.created_at.isoformat() if term_set.created_at else None,
    }


@app.post("/classes/{class_id}/cat4/term-sets/{term_set_id}/rows")
def upsert_cat4_term_rows(
    class_id: int,
    term_set_id: int,
    payload: Cat4TermResultRowsPayload,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)

    term_set = (
        db.query(Cat4TermResultSetModel)
        .filter(Cat4TermResultSetModel.id == term_set_id, Cat4TermResultSetModel.class_id == class_id)
        .first()
    )
    if not term_set:
        raise HTTPException(status_code=404, detail="CAT4 term result set not found")

    students = db.query(StudentModel).filter(StudentModel.class_id == class_id).all()
    name_index = _cat4_name_index(students)

    db.query(Cat4StudentTermResultModel).filter(
        Cat4StudentTermResultModel.result_set_id == term_set_id
    ).delete()

    matched_count = 0
    unmatched_count = 0
    created_rows = 0

    for item in payload.rows:
        raw_name = (item.raw_name or "").strip()
        if not raw_name:
            continue

        student_id, matched_name, _ = _match_cat4_student_name(raw_name, name_index)
        average_percent, subject_count, domain_scores, raw_json = _calculate_cat4_term_metrics(
            _json_text_or_none(item.raw_subjects_json),
            item.average_percent,
            item.subject_count,
        )
        row = Cat4StudentTermResultModel(
            result_set_id=term_set_id,
            class_id=class_id,
            student_id=student_id,
            raw_name=raw_name,
            matched_name=matched_name,
            average_percent=average_percent,
            subject_count=subject_count,
            raw_subjects_json=raw_json,
            verbal_domain_score=domain_scores.get("verbal_domain_score"),
            quantitative_domain_score=domain_scores.get("quantitative_domain_score"),
            non_verbal_domain_score=domain_scores.get("non_verbal_domain_score"),
            spatial_domain_score=domain_scores.get("spatial_domain_score"),
        )
        db.add(row)
        created_rows += 1
        if student_id:
            matched_count += 1
        else:
            unmatched_count += 1

    db.commit()

    return {
        "ok": True,
        "term_set_id": term_set_id,
        "created_rows": created_rows,
        "matched_count": matched_count,
        "unmatched_count": unmatched_count,
    }


@app.post("/classes/{class_id}/cat4/workbook/validate")
async def validate_cat4_workbook(
    class_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)

    filename = (file.filename or "").strip().lower()
    if not filename.endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="Please upload an .xlsx workbook")

    try:
        preview = _build_cat4_workbook_preview(await file.read())
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="The workbook file could not be opened")
    except Exception:
        logger.exception("Failed to validate CAT4 workbook for class %s", class_id)
        raise HTTPException(status_code=400, detail="Could not read workbook structure")

    return {
        "ok": len(preview["errors"]) == 0,
        "workbook_name": file.filename,
        "baseline_locked": bool(
            db.query(Cat4BaselineSetModel)
            .filter(Cat4BaselineSetModel.class_id == class_id, Cat4BaselineSetModel.is_locked == True)  # noqa: E712
            .first()
        ),
        **preview,
    }


@app.post("/classes/{class_id}/cat4/workbooks")
async def upload_cat4_workbook(
    class_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)

    filename = (file.filename or "").strip().lower()
    if not filename.endswith(".xlsx"):
        raise HTTPException(status_code=400, detail="Please upload an .xlsx workbook")

    try:
        preview = _build_cat4_workbook_preview(_read_xlsx_workbook(await file.read()))
    except zipfile.BadZipFile:
        raise HTTPException(status_code=400, detail="The workbook file could not be opened")
    except Exception:
        logger.exception("Failed to upload CAT4 workbook for class %s", class_id)
        raise HTTPException(status_code=400, detail="Could not read workbook structure")

    if preview["errors"]:
        raise HTTPException(status_code=400, detail={"message": "Workbook validation failed", "errors": preview["errors"], "warnings": preview["warnings"]})

    current_version = (
        db.query(Cat4WorkbookVersionModel)
        .filter(Cat4WorkbookVersionModel.class_id == class_id)
        .order_by(Cat4WorkbookVersionModel.version_number.desc())
        .first()
    )
    next_version_number = (current_version.version_number + 1) if current_version else 1

    db.query(Cat4WorkbookVersionModel).filter(Cat4WorkbookVersionModel.class_id == class_id).update({"is_active": False})
    _replace_cat4_data_from_workbook_payload(class_id, preview, db)

    version = Cat4WorkbookVersionModel(
        class_id=class_id,
        version_number=next_version_number,
        workbook_name=file.filename or f"cat4-workbook-v{next_version_number}.xlsx",
        uploaded_by_email=user.email,
        is_active=True,
        validation_summary_json=json.dumps(
            {
                "baseline_sheet_name": preview.get("baseline_sheet_name"),
                "term_sheet_names": preview.get("term_sheet_names"),
                "warnings": preview.get("warnings", []),
                "matched_student_count": sum(1 for row in preview.get("baseline_rows", []) if row.get("raw_name")),
                "baseline_locked": bool(preview.get("baseline_rows")),
            },
            ensure_ascii=False,
        ),
        parsed_payload_json=json.dumps(preview, ensure_ascii=False),
    )
    db.add(version)
    db.commit()
    db.refresh(version)

    return {
        "ok": True,
        "version_id": version.id,
        "version_number": version.version_number,
        "uploaded_at": version.uploaded_at.isoformat() if version.uploaded_at else None,
        "warnings": preview.get("warnings", []),
    }


@app.post("/classes/{class_id}/cat4/workbooks/{version_id}/restore")
def restore_cat4_workbook_version(
    class_id: int,
    version_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)

    version = (
        db.query(Cat4WorkbookVersionModel)
        .filter(Cat4WorkbookVersionModel.class_id == class_id, Cat4WorkbookVersionModel.id == version_id)
        .first()
    )
    if not version:
        raise HTTPException(status_code=404, detail="Workbook version not found")

    payload = json.loads(version.parsed_payload_json or "{}")
    db.query(Cat4WorkbookVersionModel).filter(Cat4WorkbookVersionModel.class_id == class_id).update({"is_active": False})
    _replace_cat4_data_from_workbook_payload(class_id, payload, db)
    version.is_active = True
    db.commit()
    db.refresh(version)

    return {"ok": True, "restored_version_id": version.id, "version_number": version.version_number}


@app.post("/classes/{class_id}/cat4/baselines/{baseline_id}/reset")
def reset_cat4_baseline(
    class_id: int,
    baseline_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    require_super_admin(user)

    baseline = (
        db.query(Cat4BaselineSetModel)
        .filter(Cat4BaselineSetModel.id == baseline_id, Cat4BaselineSetModel.class_id == class_id)
        .first()
    )
    if not baseline:
        raise HTTPException(status_code=404, detail="CAT4 baseline set not found")

    db.query(Cat4StudentBaselineModel).filter(Cat4StudentBaselineModel.baseline_set_id == baseline_id).delete()
    baseline.is_locked = False
    baseline.locked_at = None
    db.commit()

    return {"ok": True, "baseline_id": baseline_id, "reset": True}


@app.get("/classes/{class_id}/cat4/report")
def cat4_report(
    class_id: int,
    baseline_id: Optional[int] = None,
    term_set_id: Optional[int] = None,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _get_owned_class_for_cat4_or_403(class_id, db, user)
    return _build_cat4_report_payload(class_id, db, baseline_id=baseline_id, term_set_id=term_set_id)

@app.get("/classes/{class_id}/insights")
def class_insights(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    # Active students only (consistent with your results entry grid)
    students = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id)
        .filter(StudentModel.active == True)  # noqa: E712
        .order_by(StudentModel.id.desc())
        .all()
    )

    active_student_count = len(students)
    if active_student_count == 0:
        return {
            "class_id": class_id,
            "class_average": None,
            "assessment_count": 0,
            "active_student_count": 0,
            "student_rankings": [],
            "at_risk": [],
        }

    student_ids = [s.id for s in students]

    # Count assessments for this class
    assessment_count = (
        db.query(ClassAssessmentModel)
        .filter(ClassAssessmentModel.class_id == class_id)
        .count()
    )

    # Pull all results for this class (join assessments -> results)
    rows = (
        db.query(AssessmentResultModel, ClassAssessmentModel)
        .join(ClassAssessmentModel, AssessmentResultModel.assessment_id == ClassAssessmentModel.id)
        .filter(ClassAssessmentModel.class_id == class_id)
        .filter(AssessmentResultModel.student_id.in_(student_ids))
        .all()
    )

    # Aggregate per student
    scores_by_student: dict[int, list[int]] = {sid: [] for sid in student_ids}
    tests_count_by_student: dict[int, int] = {sid: 0 for sid in student_ids}
    absences_by_student: dict[int, int] = {sid: 0 for sid in student_ids}
    latest_by_student: dict[int, int | None] = {sid: None for sid in student_ids}

    for r, a in rows:
        tests_count_by_student[r.student_id] += 1

        if r.absent:
            absences_by_student[r.student_id] += 1
            continue

        if r.score_percent is None:
            continue

        sc = int(r.score_percent)
        scores_by_student[r.student_id].append(sc)
        latest_by_student[r.student_id] = sc  # "latest" in the order rows arrive (good enough for now)

    # Build output rows
    rankings = []
    class_avgs = []

    for s in students:
        scores = scores_by_student.get(s.id, [])
        avg = round(sum(scores) / len(scores), 1) if scores else None

        if avg is not None:
            class_avgs.append(avg)

        rankings.append(
            {
                "student_id": s.id,
                "first_name": s.first_name,
                "average": avg,
                "taken": tests_count_by_student.get(s.id, 0),
                "missed": absences_by_student.get(s.id, 0),
                "latest": latest_by_student.get(s.id, None),
            }
        )

    # Sort: strongest -> weakest by average; students with no avg at bottom
    rankings.sort(
        key=lambda x: (x["average"] is None, -(x["average"] or -1))
    )

    class_avg = round(sum(class_avgs) / len(class_avgs), 1) if class_avgs else None

    # At-risk rule v1 (simple + effective)
    # - average < 50 OR missed >= 2
    at_risk = []
    for r in rankings:
        reasons = []
        if r["average"] is not None and r["average"] < 50:
            reasons.append("Average below 50")
        if r["missed"] >= 2:
            reasons.append(f"Missed {r['missed']} assessments")
        if reasons:
            at_risk.append(
                {
                    "student_id": r["student_id"],
                    "first_name": r["first_name"],
                    "average": r["average"],
                    "missed": r["missed"],
                    "reasons": reasons,
                }
            )

    return {
        "class_id": class_id,
        "class_average": class_avg,
        "assessment_count": assessment_count,
        "active_student_count": active_student_count,
        "student_rankings": rankings,
        "at_risk": at_risk,
    }

@app.get("/classes/{class_id}/report-data")
def class_report_data(
    class_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    get_owned_class_or_404(class_id, db, user)
    students = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id)
        .filter(StudentModel.active == True)  # noqa: E712
        .order_by(StudentModel.first_name.asc())
        .all()
    )

    assessments = (
        db.query(ClassAssessmentModel)
        .filter(ClassAssessmentModel.class_id == class_id)
        .order_by(ClassAssessmentModel.assessment_date.asc(), ClassAssessmentModel.id.asc())
        .all()
    )

    assessment_ids = [a.id for a in assessments]

    results = []
    if assessment_ids:
        results = (
            db.query(AssessmentResultModel)
            .filter(AssessmentResultModel.assessment_id.in_(assessment_ids))
            .all()
        )

    results_by_student: dict[int, list[AssessmentResultModel]] = {}
    for r in results:
        results_by_student.setdefault(r.student_id, []).append(r)

    report_students = []

    for s in students:
        student_results = results_by_student.get(s.id, [])

        numeric_scores = [
            int(r.score_percent)
            for r in student_results
            if not r.absent and r.score_percent is not None
        ]

        average = round(sum(numeric_scores) / len(numeric_scores), 1) if numeric_scores else None
        taken = len(numeric_scores)
        missed = len([r for r in student_results if r.absent])

        report_students.append(
            {
                "id": s.id,
                "name": s.first_name,
                "average": average,
                "taken": taken,
                "missed": missed,
            }
        )

    class_average_values = [s["average"] for s in report_students if s["average"] is not None]
    class_average = (
        round(sum(class_average_values) / len(class_average_values), 1)
        if class_average_values
        else None
    )

    return {
        "class_id": class_id,
        "assessment_count": len(assessments),
        "class_average": class_average,
        "students": report_students,
    }


@app.get("/classes/{class_id}/students/{student_id}/report-data")
def student_report_data(
    class_id: int,
    student_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    student = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id, StudentModel.id == student_id)
        .first()
    )
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")
    get_owned_class_or_404(student.class_id, db, user)

    assessments = (
        db.query(ClassAssessmentModel)
        .filter(ClassAssessmentModel.class_id == class_id)
        .order_by(ClassAssessmentModel.assessment_date.asc(), ClassAssessmentModel.id.asc())
        .all()
    )

    assessment_ids = [a.id for a in assessments]

    results = []
    if assessment_ids:
        results = (
            db.query(AssessmentResultModel)
            .filter(
                AssessmentResultModel.student_id == student_id,
                AssessmentResultModel.assessment_id.in_(assessment_ids),
            )
            .all()
        )

    results_by_assessment = {r.assessment_id: r for r in results}

    assessment_rows = []
    numeric_scores = []

    for a in assessments:
        r = results_by_assessment.get(a.id)

        if not r:
            result_text = "—"
            absent = False
            score_value = None
        elif r.absent:
            result_text = "Absent"
            absent = True
            score_value = None
        else:
            score_value = int(r.score_percent) if r.score_percent is not None else None
            result_text = f"{score_value}%" if score_value is not None else "—"
            absent = False
            if score_value is not None:
                numeric_scores.append(score_value)

        assessment_rows.append(
            {
                "assessment_id": a.id,
                "title": a.title,
                "date": a.assessment_date.date().isoformat() if a.assessment_date else None,
                "result": result_text,
                "absent": absent,
                "score_percent": score_value,
            }
        )

    average = round(sum(numeric_scores) / len(numeric_scores), 1) if numeric_scores else None
    taken = len(numeric_scores)
    missed = len([r for r in results if r.absent])

    return {
        "class_id": class_id,
        "student": {
            "id": student.id,
            "name": student.first_name,
            "average": average,
            "taken": taken,
            "missed": missed,
        },
        "assessment_count": len(assessments),
        "assessments": assessment_rows,
    }

@app.get("/classes/{class_id}/students/{student_id}/history")
def student_history(
    class_id: int,
    student_id: int,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    # ✅ Security: ensure class belongs to logged-in teacher
    cls = db.query(ClassModel).filter(
        ClassModel.id == class_id,
        ClassModel.owner_user_id == user.id
    ).first()
    if not cls:
        raise HTTPException(status_code=404, detail="Class not found")

    # Ensure student belongs to this class
    student = db.query(StudentModel).filter(
        StudentModel.id == student_id,
        StudentModel.class_id == class_id
    ).first()
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    # Assessments in chronological order
    assessments = (
        db.query(ClassAssessmentModel)
        .filter(ClassAssessmentModel.class_id == class_id)
        .order_by(ClassAssessmentModel.assessment_date.asc(), ClassAssessmentModel.id.asc())
        .all()
    )

    if not assessments:
        return {
            "student": {"id": student.id, "first_name": student.first_name},
            "points": [],
        }

    assessment_ids = [a.id for a in assessments]

    # Student results for those assessments
    student_results = (
        db.query(AssessmentResultModel)
        .filter(AssessmentResultModel.assessment_id.in_(assessment_ids))
        .filter(AssessmentResultModel.student_id == student_id)
        .all()
    )
    by_assessment = {r.assessment_id: r for r in student_results}

    # Class averages per assessment (exclude absent + null scores)
    avg_rows = (
        db.query(
            AssessmentResultModel.assessment_id.label("assessment_id"),
            func.avg(AssessmentResultModel.score_percent).label("avg_score"),
        )
        .filter(AssessmentResultModel.assessment_id.in_(assessment_ids))
        .filter(AssessmentResultModel.absent == False)  # noqa: E712
        .filter(AssessmentResultModel.score_percent.isnot(None))
        .group_by(AssessmentResultModel.assessment_id)
        .all()
    )
    avg_by_assessment = {int(r.assessment_id): float(r.avg_score) for r in avg_rows}

    points = []
    for a in assessments:
        r = by_assessment.get(a.id)

        # student score: null if absent or missing row
        student_score = None
        absent = False

        if r:
            absent = bool(r.absent)
            if not absent and r.score_percent is not None:
                student_score = int(r.score_percent)

        class_avg = avg_by_assessment.get(a.id, None)
        class_avg_rounded = round(class_avg, 1) if class_avg is not None else None

        points.append(
            {
                "assessment_id": a.id,
                "title": a.title,
                "date": a.assessment_date.date().isoformat() if a.assessment_date else None,
                "student": student_score,
                "absent": absent,
                "class_avg": class_avg_rounded,
            }
        )

    return {
        "student": {"id": student.id, "first_name": student.first_name},
        "points": points,
    }

# -------------------------
# AI Calendar Assistant (draft only - NEVER writes to DB)
# -------------------------
@app.post("/ai/parse-event", response_model=schemas.AIParseEventResponse)
def ai_parse_event(
    payload: schemas.AIParseEventRequest,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    text = (payload.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="Text is required")
    _enforce_ai_prompt_limit(db, user)

    # Lazy import so your backend still runs even if openai isn't installed in this env
    try:
        from openai import OpenAI  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI client not available: {e}")

    # ✅ Make key usage explicit (prevents silent 500s)
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail="OPENAI_API_KEY is missing. Check backend .env and that load_dotenv() runs on startup.",
        )

    client = OpenAI(api_key=api_key)

    system = (
       "Today is " + datetime.now().date().isoformat() + " in timezone " + payload.timezone + ". "
        "If the user does not specify a year, choose the next occurrence in the future (not past). "
        "You convert teacher calendar requests into a single JSON object for an event draft. "
        "Return ONLY valid JSON. No markdown. "
        "Timezone: " + payload.timezone + ". "
        "If date is relative (e.g. next Friday), resolve it relative to today's date. "
        "If no end time is provided, set end_date to start + default duration. "
        "If it's clearly an all-day item (e.g. 'midterm break', 'bank holiday'), set all_day=true and set times to 09:00 start, 09:00 end. "
        "Use event_type one of: general, test, homework, trip. "
        "Required keys: title, description, event_date, end_date, all_day, event_type, class_id."
    )

    user_message = (
        f"Text: {text}\n"
        f"class_id: {payload.class_id}\n"
        f"default_duration_minutes: {payload.default_duration_minutes}\n"
        f"timezone: {payload.timezone}"
        f"today: {datetime.now().date().isoformat()}\n"
    )

    resp = client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user_message},
        ],
        temperature=0.2,
    )

    content = (resp.choices[0].message.content or "").strip()

    # Extract the first JSON object defensively (supports extra text)
    m = re.search(r"\{[\s\S]*\}", content)
    if not m:
        raise HTTPException(status_code=500, detail=f"AI did not return JSON. Got: {content[:200]}")

    try:
        data = json.loads(m.group(0))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not parse AI JSON: {e}")

    # Some models return {"draft": {...}}; accept both shapes
    if isinstance(data, dict) and "draft" in data and isinstance(data["draft"], dict):
        data = data["draft"]

    warnings: list[str] = []

    # Force class_id from payload if provided (front-end controls target class)
    if payload.class_id is not None:
        data["class_id"] = payload.class_id

    # If end_date missing, we’ll let schema validation complain nicely (422),
    # but keep your warning behavior too:
    if data.get("end_date") in (None, ""):
        warnings.append(f"End time assumed as {payload.default_duration_minutes} minutes (AI did not supply end_date).")

    # ✅ Catch schema validation errors and return 422 instead of 500
    try:
        draft = schemas.CalendarEventCreate(**data)
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail=f"AI draft failed validation against CalendarEventCreate. Error: {e}. Draft keys: {list(data.keys())}",
        )

    _record_ai_prompt_usage(db, user)
    return schemas.AIParseEventResponse(draft=draft, warnings=warnings)

# -------------------------
# AI CreateResources (draft generator - DOES NOT write to DB)
# -------------------------

from typing import Any, Literal

class AICreateResourcesSource(BaseModel):
    id: str | None = None
    title: str = ""
    pages: Any | None = None   # front-end sends array or null; keep flexible
    text: str = ""


class AICreateResourcesSaveTarget(BaseModel):
    bucket: str | None = None
    folder: str | None = None


class AICreateResourcesManualFile(BaseModel):
    id: str | None = None
    filename: str | None = None
    mime_type: str | None = None
    size_bytes: int | None = None

class AICreateResourcesScope(BaseModel):
    mode: Literal["general", "single", "group"]
    classId: int | None = None
    classIds: list[int] | None = None
    groupName: str | None = None

class AICreateResourcesRequest(BaseModel):
    kind: str                       # e.g. "lesson_plan", "scheme", "dept_plan", "worksheet", "ideas"
    template: str | None = None     # legacy payload
    tone: str | None = None         # legacy payload
    level: str | None = None
    detail: str | None = None
    scope: AICreateResourcesScope
    prompt: str
    sources: list[AICreateResourcesSource] = []
    manual_sources: list[AICreateResourcesSource] = []
    manual_file_sources: list[AICreateResourcesManualFile] = []
    save_target: AICreateResourcesSaveTarget | None = None
    source_context: dict[str, Any] | None = None
    brandingChoice: str | None = None
    include_answer_key: bool | None = None
    worksheet_options: dict[str, Any] | None = None
    branding: dict[str, Any] | None = None
    teacherDisplayNameShort: str | None = None
    schoolName: str | None = None
    outputKind: str | None = None
    output_intent: str | None = None
    audience: str | None = None
    timezone: str = "Europe/Dublin"

    class Config:
        extra = "allow"

class AICreateResourcesResponse(BaseModel):
    title: str
    content: str


def _ai_create_resources_template_for_kind(kind: str) -> str:
    k = (kind or "").strip().lower()
    if k == "ideas":
        return "3 Ideas"
    if k == "lesson_plan":
        return "Lesson Plan"
    if k == "worksheet":
        return "Worksheet"
    if k == "scheme":
        return "Scheme of Work"
    if k == "dept_plan":
        return "Department Plan"
    return "Resource"


def _ai_create_resources_tone(kind: str, detail: str | None, tone: str | None) -> str:
    if tone and tone.strip():
        return tone.strip()
    d = (detail or "").strip().lower()
    k = (kind or "").strip().lower()
    base = "Clear and teacher-friendly"
    if k == "worksheet":
        base = "Clear, student-facing, and printable"
    elif k == "dept_plan":
        base = "Professional and department-facing"
    elif k == "scheme":
        base = "Professional and sequenced"
    elif k == "ideas":
        base = "Quick, practical, and teacher-facing"
    if d == "detailed":
        return f"{base}; more detailed"
    if d == "concise":
        return f"{base}; concise"
    return base


def _ai_create_resources_include_answer_key(payload: AICreateResourcesRequest) -> bool:
    if payload.include_answer_key is not None:
        return bool(payload.include_answer_key)
    if isinstance(payload.worksheet_options, dict):
        if "include_answer_key" in payload.worksheet_options:
            return bool(payload.worksheet_options.get("include_answer_key"))
    return True


def _ai_create_resources_source_bundle(payload: AICreateResourcesRequest) -> tuple[str, str]:
    combined: list[AICreateResourcesSource] = []
    if payload.sources:
        combined.extend(payload.sources)
    if payload.manual_sources:
        combined.extend(payload.manual_sources)

    chunks: list[str] = []
    char_budget = 12000
    used = 0

    for idx, s in enumerate(combined[:24], start=1):
        title = (s.title or "").strip() or f"Source {idx}"
        text_value = (s.text or "").strip()
        if not text_value:
            continue
        text_value = re.sub(r"\s+\n", "\n", text_value)
        text_value = re.sub(r"\n{3,}", "\n\n", text_value)
        text_value = text_value[:1800]
        pages = f" (pages: {s.pages})" if s.pages else ""
        chunk = f"---\nSOURCE: {title}{pages}\n{text_value}"
        if used + len(chunk) > char_budget:
            remaining = max(0, char_budget - used)
            if remaining < 200:
                break
            chunk = chunk[:remaining]
        chunks.append(chunk)
        used += len(chunk)
        if used >= char_budget:
            break

    file_notes = []
    for f in payload.manual_file_sources[:12]:
        filename = (f.filename or "").strip()
        if filename:
            file_notes.append(filename)
    file_hint = ""
    if file_notes:
        file_hint = "Teacher also selected uploaded files for context: " + ", ".join(file_notes)

    sources_txt = "\n".join(chunks)
    return sources_txt, file_hint

@app.post("/ai/create-resources", response_model=AICreateResourcesResponse)
def ai_create_resources(
    payload: AICreateResourcesRequest,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),  # keep it teacher-only
):
    p = (payload.prompt or "").strip()
    if not p:
        raise HTTPException(status_code=400, detail="prompt is required")
    _enforce_ai_prompt_limit(db, user)

    # Lazy import so backend still runs without openai installed
    try:
        from openai import OpenAI  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI client not available: {e}")

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail="OPENAI_API_KEY is missing. Check backend .env and that load_dotenv() runs on startup.",
        )

    client = OpenAI(api_key=api_key)

    template = (payload.template or "").strip() or _ai_create_resources_template_for_kind(payload.kind)
    tone = _ai_create_resources_tone(payload.kind, payload.detail, payload.tone)
    level = (payload.level or "").strip() or "Not specified"
    answer_key = _ai_create_resources_include_answer_key(payload)
    sources_txt, file_hint = _ai_create_resources_source_bundle(payload)

    scope_desc = ""
    if payload.scope.mode == "general":
        scope_desc = "General (not tied to a specific class)"
    elif payload.scope.mode == "single":
        scope_desc = f"Single class_id: {payload.scope.classId}"
    else:
        scope_desc = f"Group class_ids: {payload.scope.classIds} | groupName: {payload.scope.groupName}"

    save_bucket = (payload.save_target.bucket if payload.save_target else None) or ""
    save_folder = (payload.save_target.folder if payload.save_target else None) or ""
    source_context = payload.source_context or {}
    selected_folder = (
        (source_context.get("selected_folder") if isinstance(source_context, dict) else None)
        or save_folder
        or ""
    )
    selected_bucket = (
        (source_context.get("selected_bucket") if isinstance(source_context, dict) else None)
        or save_bucket
        or ""
    )
    output_kind = (payload.outputKind or payload.kind or "").strip()
    teacher_name_short = (
        (payload.teacherDisplayNameShort or "").strip()
        or str((payload.branding or {}).get("teacherDisplayNameShort") or "").strip()
        or "Teacher"
    )
    school_name = (
        (payload.schoolName or "").strip()
        or str((payload.branding or {}).get("schoolName") or "").strip()
    )
    branding_choice = (
        (payload.brandingChoice or "").strip()
        or str((payload.branding or {}).get("brandingChoice") or "").strip()
        or "elume"
    )
    audience_rule = {
        "ideas": "Produce three structured teacher-facing teaching ideas for immediate classroom use, not a lesson plan.",
        "lesson_plan": "Produce a concise teacher-facing lesson plan for an Irish post-primary classroom using the exact requested section structure.",
        "worksheet": "Produce a clean student-facing printable worksheet for an Irish post-primary classroom, not a lesson plan.",
        "scheme": "Produce sequenced teacher planning for a scheme of work.",
        "dept_plan": "Produce department-facing planning for shared use across a subject team.",
    }.get((payload.kind or "").strip().lower(), "Produce a clear classroom resource.")

    system = (
        f"Today is {datetime.now().date().isoformat()} in timezone {payload.timezone}. "
        "You are an assistant for teachers writing classroom resources for the Irish secondary school / post-primary context. "
        "Use British English spelling and terminology throughout. "
        "Prioritise Irish classroom framing, curriculum assumptions, and school language. "
        "Do not drift into US or non-Irish assumptions unless the teacher explicitly asks for that context. "
        "Treat the selected class or group as the default working context. "
        "Treat the selected folder as the priority search context when relevant. "
        "Manual pasted or uploaded sources are teacher-selected guidance and should be used carefully when present. "
        "CRITICAL RULE: You may ONLY use the provided SOURCE EXCERPTS as factual grounding. "
        "If something is not in the sources, write it as a sensible teaching suggestion, not as a quoted fact, and do not invent curriculum citations, exam-board references, or school policies. "
        "Do not clutter the resource with branding or footer metadata; that is handled separately. "
        "Return ONLY valid JSON (no markdown, no backticks). "
        "JSON must have exactly these keys: title, content. "
        "content should be plain text with clear headings and bullet points where useful. "
        "Do not include any extra keys."
    )

    user_msg = (
        f"Teacher: {user.email}\n"
        f"teacherDisplayNameShort: {teacher_name_short}\n"
        f"schoolName: {school_name or '[Not provided]'}\n"
        f"brandingChoice: {branding_choice}\n"
        f"kind: {payload.kind}\n"
        f"outputKind: {output_kind or payload.kind}\n"
        f"template: {template}\n"
        f"tone: {tone}\n"
        f"detail: {(payload.detail or '').strip() or '[Not provided]'}\n"
        f"level: {level}\n"
        f"audience: {(payload.audience or '').strip() or '[Inferred from kind]'}\n"
        f"output_intent: {(payload.output_intent or '').strip() or template}\n"
        f"scope: {scope_desc}\n"
        f"save_target_bucket: {save_bucket or '[Not provided]'}\n"
        f"selected_folder_priority: {selected_folder or '[Top level / none]'}\n"
        f"selected_bucket_context: {selected_bucket or '[Not provided]'}\n"
        f"include_answer_key: {'Yes' if answer_key else 'No'}\n"
        f"audience_rule: {audience_rule}\n"
        f"prompt: {p}\n"
        "\n"
        "Generation rules:\n"
        "- Write for Irish post-primary teaching and learning.\n"
        "- Use British English spelling and terminology.\n"
        "- Keep the response appropriate to the output kind.\n"
        "- When sources are present, use them carefully and do not invent unsupported facts.\n"
        "- When no sources are present, draft sensibly but avoid fabricated curriculum citations.\n"
        "- When teacher-selected files or manual notes are present for lesson_plan or worksheet, treat them as the primary truth and preserve their facts, framing, vocabulary, and sequence where possible.\n"
        "- For lesson_plan or worksheet, do not drift into generic content if usable source content exists.\n"
        "- If you add support content beyond the source, frame it as sensible teaching support rather than as a source fact.\n"
        "- If ideas, produce exactly 3 structured teaching ideas and do not turn them into a lesson plan.\n"
        "- If ideas, follow this exact top-level structure: Idea 1: Thought-provoking question; Idea 2: Thought-provoking activity; Idea 3: Collaborative Board session.\n"
        "- If ideas, Idea 1 must centre on a question that can be answered by an individual or a group.\n"
        "- If ideas, Idea 2 must centre on an activity that can be completed by an individual or a group.\n"
        "- If ideas, Idea 3 must centre on a Collaborative Board task using Elume's built-in collaboration technology.\n"
        "- If ideas, for each idea include practical teacher-facing elements such as a Title, The Hook, The Task, optional Board / Whiteboard / Collaboration integration where relevant, and Why it works or the key discussion angle.\n"
        "- If ideas, keep the ideas concise, practical, varied in wording, and ready for immediate use in an Irish post-primary classroom.\n"
        "- If ideas, do not generate lesson-plan sections such as Learning Intentions, Success Criteria, Lesson Flow, Resources, Assessment, or Homework.\n"
        "- If worksheet, write for Irish post-primary / secondary students using British English.\n"
        "- If worksheet, make it feel like a clean printable classroom handout, not a lesson plan or teacher note.\n"
        "- If worksheet, keep the wording concise, student-facing, and easy to follow in print.\n"
        "- If worksheet, do not write like a raw markdown dump.\n"
        "- If worksheet, follow this structure: Worksheet Title, Student Name line, Class line, Date line, short Instructions, 3 to 5 worksheet task sections, Extension Challenge if useful, and Answer Key only if included.\n"
        "- If worksheet, anchor the content closely to the selected source notes or selected folder context whenever sources are available.\n"
        "- If worksheet, use the selected source notes first and derive named facts, keywords, processes, definitions, and examples directly from the provided material.\n"
        "- If worksheet, avoid broad generic worksheet questions when more specific source-grounded questions are possible.\n"
        "- If worksheet, prefer real handout tasks such as a fill-in-the-blanks paragraph, label or identify, short-answer questions tied to exact concepts from the notes, sequencing or process steps from the notes, and compare or contrast tasks using concepts explicitly present in the notes.\n"
        "- If worksheet, one task should usually be a proper fill-in-the-blanks paragraph when the topic suits it; for science or biology, use topic vocabulary drawn from the notes.\n"
        "- If worksheet, the fill-in-the-blanks paragraph should read like a real cloze exercise, not a loose bullet list.\n"
        "- If worksheet, reduce generic filler such as broad reflection prompts, placeholder task summaries, generic research tasks, or 'write a definition in your own words' when a more specific notes-based question is possible.\n"
        "- If worksheet, do not refer to a diagram, image, chart, or labelled visual unless one is actually provided.\n"
        "- If worksheet, do not say 'Below is a diagram' or ask students to label an image unless that asset is actually included; instead use wording such as 'Draw a simple diagram of ...' or 'Draw and label ...' when a text-only task is needed.\n"
        "- If worksheet, do not include lesson-plan sections such as Learning Intentions, Success Criteria, Lesson Flow, Assessment, or Homework.\n"
        "- If worksheet, do not let markdown markers like ### Task 1 or ### Reflection appear in the final worksheet body.\n"
        "- If worksheet, include an answer key unless told not to.\n"
        "- If lesson_plan, do not use tables.\n"
        "- If lesson_plan, keep it concise, printable, teacher-facing, and suitable for roughly 12 pages or less when exported.\n"
        "- If lesson_plan, use clear headings only, keep spacing clean, and avoid unnecessary explanation, symbols, or excessive line breaks.\n"
        "- If lesson_plan, do not write like a raw markdown dump or a generic scaffold.\n"
        "- If lesson_plan, write for Irish post-primary / secondary classrooms using practical, realistic classroom phrasing.\n"
        "- If lesson_plan, avoid US curriculum assumptions unless the teacher explicitly asks for them.\n"
        "- If lesson_plan, title it as Lesson Plan: {topic/title} and follow with one clean metadata line in the form Subject | Level | Duration.\n"
        "- If lesson_plan, keep Learning Overview short and topic-led rather than padded with a long overview paragraph.\n"
        "- If lesson_plan, write Success Criteria as checklist-style 'I can ...' statements.\n"
        "- If lesson_plan, follow exactly this structure and order:\n"
        "  1. Learning Overview\n"
        "  2. Learning Intentions\n"
        "  3. Success Criteria\n"
        "  4. Lesson Flow\n"
        "  5. Starter\n"
        "  6. Teaching and Development\n"
        "  7. Activity and Application\n"
        "  8. Plenary and Closure\n"
        "  9. Resources\n"
        "  10. Differentiation\n"
        "  11. Assessment\n"
        "  12. Suggested Homework\n"
        "  13. Reflection\n"
        "- If lesson_plan, make each lesson-flow subsection usable and specific rather than filler.\n"
        "- If lesson_plan, make the four timed lesson-flow subsections read as concrete classroom actions, not vague scaffold text.\n"
        "- If lesson_plan, if usable source content exists, preserve its topic framing, vocabulary, named facts, and sequence where possible.\n"
        "- If lesson_plan, keep the tone practical, school-ready, and concise.\n"
        "\n"
        "SOURCE EXCERPTS (allowed):\n"
        f"{sources_txt if sources_txt else '[No sources selected]'}\n"
        f"{file_hint + chr(10) if file_hint else ''}"
        "\n"
        "Write the best possible resource for the teacher."
    )

    resp = client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.3,
    )

    content = (resp.choices[0].message.content or "").strip()

    # Extract JSON defensively (handles occasional extra text)
    m = re.search(r"\{[\s\S]*\}", content)
    if not m:
        raise HTTPException(status_code=500, detail=f"AI did not return JSON. Got: {content[:200]}")

    try:
        data = json.loads(m.group(0))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Could not parse AI JSON: {e}")

    title = (data.get("title") or "").strip()
    body = (data.get("content") or "").strip()

    if not title:
        title = f"{template} - {p}"[:80]
    if not body:
        raise HTTPException(status_code=500, detail="AI returned empty content")

    _record_ai_prompt_usage(db, user)
    return {"title": title, "content": body}
class AIReportCommentRequest(BaseModel):
    length: str = "Medium"            # Short | Medium | Long
    indicators: List[str] = []
    sign_off: Optional[str] = None


class AIReportCommentResponse(BaseModel):
    comment: str


@app.post(
    "/classes/{class_id}/students/{student_id}/generate-report-comment",
    response_model=AIReportCommentResponse,
)
def generate_report_comment(
    class_id: int,
    student_id: int,
    payload: AIReportCommentRequest,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _enforce_ai_prompt_limit(db, user)
    get_owned_class_or_404(class_id, db, user)
    student = (
        db.query(StudentModel)
        .filter(StudentModel.class_id == class_id, StudentModel.id == student_id)
        .first()
    )
    if not student:
        raise HTTPException(status_code=404, detail="Student not found")

    assessments = (
        db.query(ClassAssessmentModel)
        .filter(ClassAssessmentModel.class_id == class_id)
        .order_by(ClassAssessmentModel.assessment_date.asc(), ClassAssessmentModel.id.asc())
        .all()
    )

    assessment_ids = [a.id for a in assessments]

    results = []
    if assessment_ids:
        results = (
            db.query(AssessmentResultModel)
            .filter(
                AssessmentResultModel.student_id == student_id,
                AssessmentResultModel.assessment_id.in_(assessment_ids),
            )
            .all()
        )

    by_assessment = {r.assessment_id: r for r in results}

    numeric_scores = [
        int(r.score_percent)
        for r in results
        if not r.absent and r.score_percent is not None
    ]
    average = round(sum(numeric_scores) / len(numeric_scores), 1) if numeric_scores else None
    taken = len(numeric_scores)
    missed = len([r for r in results if r.absent])

    latest_score = None
    for a in reversed(assessments):
        r = by_assessment.get(a.id)
        if r and (not r.absent) and r.score_percent is not None:
            latest_score = int(r.score_percent)
            break

    try:
        from openai import OpenAI  # type: ignore
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI client not available: {e}")

    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=500,
            detail="OPENAI_API_KEY is missing. Check backend .env and that load_dotenv() runs on startup.",
        )

    client = OpenAI(api_key=api_key)

    desired_length = (payload.length or "Medium").strip()
    if desired_length not in ["Short", "Medium", "Long"]:
        desired_length = "Medium"

    system = (
        f"Today is {datetime.now().date().isoformat()}. "
        "You are an assistant helping a teacher write a school report comment. "
        "Write in a professional, natural, supportive school-report style suitable for Ireland/UK schools. "
        "Start the comment with the student's first name. "
        "ALWAYS include the student's first name. "
        "Do not use bullet points. Do not use markdown. "
        "Write one polished paragraph only. "
        "Do not invent facts beyond the data provided. "
        "If behaviour or effort indicators are provided, weave them in naturally. "
        "If a sign-off is provided, place it naturally at the end."
        "Avoid repeating identical sentence structures across comments. "
        "When mentioning scores, averages, or test results, always format them as whole-number percentages (e.g., 82%) with no decimal places."
    )

    if desired_length == "Short":
        length_rule = "Write about 2 short sentences."
    elif desired_length == "Long":
        length_rule = "Write about 4 to 5 sentences."
    else:
        length_rule = "Write about 3 sentences."

    user_msg = (
        f"Teacher email: {user.email}\n"
        f"Student name: {student.first_name}\n"
        f"Class id: {class_id}\n"
        f"Average score: {average if average is not None else 'N/A'}\n"
        f"Latest score: {latest_score if latest_score is not None else 'N/A'}\n"
        f"Assessments completed: {taken}\n"
        f"Assessments missed: {missed}\n"
        f"Indicators: {', '.join(payload.indicators) if payload.indicators else 'None'}\n"
        f"Sign-off: {(payload.sign_off or '').strip() or 'None'}\n"
        f"Length instruction: {length_rule}\n"
        "\n"
        "Write the final student report comment now."
    )

    resp = client.chat.completions.create(
        model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
        messages=[
            {"role": "system", "content": system},
            {"role": "user", "content": user_msg},
        ],
        temperature=0.4,
    )

    comment = (resp.choices[0].message.content or "").strip()
    if not comment:
        raise HTTPException(status_code=500, detail="AI returned empty comment")

    _record_ai_prompt_usage(db, user)
    return {"comment": comment}


# ================= AI QUIZ GENERATION =================

import requests
from pypdf import PdfReader

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
OPENAI_MODEL = "gpt-4o-mini"
OPENAI_URL = "https://api.openai.com/v1/chat/completions"


class GenerateQuizRequest(BaseModel):
    class_id: int
    note_id: int
    num_questions: int = 10


@app.get("/ai/status")
def ai_status():
    return {
        "has_key": bool(OPENAI_API_KEY),
        "model": OPENAI_MODEL,
    }


def extract_pdf_text(path: str):
    reader = PdfReader(path)
    text = ""

    for page in reader.pages:
        try:
            text += page.extract_text() or ""
        except:
            pass

    return text[:12000]

@app.post("/ai/generate-quiz-from-note")
def generate_quiz(
    payload: GenerateQuizRequest,
    db: Session = Depends(get_db),
    user: models.UserModel = Depends(get_current_user),
):
    _enforce_ai_prompt_limit(db, user)

    note = db.query(models.Note).filter(models.Note.id == payload.note_id).first()

    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    get_owned_class_or_404(note.class_id, db, user)

    resolved_path: Optional[Path] = None
    if note.stored_path:
        direct_path = Path(note.stored_path)
        if direct_path.exists():
            resolved_path = direct_path
        else:
            normalized = str(note.stored_path).replace("\\", "/")
            marker = "uploads/"
            marker_index = normalized.lower().find(marker)
            if marker_index != -1:
                rel_suffix = normalized[marker_index + len(marker):].lstrip("/")
                fallback_path = (UPLOADS_DIR / Path(rel_suffix)).resolve()
                try:
                    fallback_path.relative_to(UPLOADS_DIR.resolve())
                except ValueError:
                    fallback_path = None
                if fallback_path and fallback_path.exists():
                    resolved_path = fallback_path

    if not resolved_path:
        raise HTTPException(status_code=404, detail="PDF missing")

    text = extract_pdf_text(str(resolved_path))

    prompt = f"""
Create {payload.num_questions} multiple choice quiz questions from this content.

Return JSON format:

{{
"title": "Quiz",
"questions":[
{{
"prompt": "",
"choices": ["","","",""],
"correctIndex": 0
}}
]
}}

CONTENT:
{text}
"""

    if not OPENAI_API_KEY:
        raise HTTPException(status_code=500, detail="Missing OpenAI API key")

    try:
        try:
            from openai import OpenAI  # type: ignore
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"OpenAI client not available: {e}")

        client = OpenAI(api_key=OPENAI_API_KEY)
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.4,
        )
        content = (resp.choices[0].message.content or "").strip()
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=500,
            detail=f"OpenAI quiz generation failed: {str(exc)}",
        )

    cleaned = content.strip()

    if cleaned.startswith("```"):
        cleaned = cleaned.replace("```json", "").replace("```", "").strip()

    try:
        quiz = json.loads(cleaned)
    except Exception:
        return {
            "error": "OpenAI did not return valid JSON",
            "raw_response": cleaned
        }
    _record_ai_prompt_usage(db, user)
    return quiz


