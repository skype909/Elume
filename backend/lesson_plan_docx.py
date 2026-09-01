"""Native DOCX renderer for validated Create Resources 2.0 Lesson Plans."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from structured_documents import StructuredLessonPlanDocument


EMERALD = "087F5B"
TEAL = "0F766E"
CYAN = "E6FFFB"
VIOLET = "6D28D9"
SLATE = "334155"
LIGHT_EMERALD = "ECFDF5"
LIGHT_AMBER = "FFFBEB"
LIGHT_VIOLET = "F5F3FF"
LIGHT_SLATE = "F8FAFC"


def _set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def _set_cell_margins(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _repeat_table_header(row) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)


def _keep_row_together(row) -> None:
    from docx.oxml import OxmlElement

    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def _set_keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def _set_cell_text(cell, text: str, *, bold: bool = False, size: float = 9, colour: str = SLATE) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = __import__("docx").shared.Pt(size)
    run.font.color.rgb = __import__("docx").shared.RGBColor.from_string(colour)
    _set_cell_margins(cell)


def _section_heading(doc, title: str):
    paragraph = doc.add_paragraph()
    _set_keep_with_next(paragraph)
    paragraph.paragraph_format.space_before = __import__("docx").shared.Pt(13)
    paragraph.paragraph_format.space_after = __import__("docx").shared.Pt(5)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = __import__("docx").shared.Pt(12)
    run.font.color.rgb = __import__("docx").shared.RGBColor.from_string(EMERALD)
    return paragraph


def _add_bullets(doc, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = __import__("docx").shared.Pt(3)
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = __import__("docx").shared.Pt(10)


def _add_panel(doc, title: str, text: str, *, fill: str, title_colour: str = EMERALD) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _set_cell_shading(cell, fill)
    _set_cell_margins(cell, top=150, start=170, bottom=150, end=170)
    title_paragraph = cell.paragraphs[0]
    title_paragraph.paragraph_format.space_after = __import__("docx").shared.Pt(4)
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = __import__("docx").shared.Pt(10)
    title_run.font.color.rgb = __import__("docx").shared.RGBColor.from_string(title_colour)
    body = cell.add_paragraph()
    body.paragraph_format.space_after = 0
    body_run = body.add_run(text)
    body_run.font.name = "Arial"
    body_run.font.size = __import__("docx").shared.Pt(10)
    _keep_row_together(table.rows[0])


def _add_definitions_table(doc, definitions: list[dict[str, str]]) -> None:
    _section_heading(doc, "Key definitions")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    headers = ("Term", "Definition")
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        _set_cell_shading(cell, EMERALD)
        _set_cell_text(cell, header, bold=True, size=9, colour="FFFFFF")
    _repeat_table_header(table.rows[0])
    for definition in definitions:
        row = table.add_row()
        _set_cell_text(row.cells[0], definition["term"], bold=True)
        _set_cell_text(row.cells[1], definition["definition"])
        _keep_row_together(row)


def _add_lesson_flow(doc, items: list[dict[str, Any]]) -> None:
    _section_heading(doc, "Timed lesson flow")
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    table.autofit = False
    headers = ("Time", "Phase", "Teacher", "Students", "Check")
    widths = (0.62, 0.78, 1.82, 1.82, 1.42)
    for index, (header, width) in enumerate(zip(headers, widths)):
        cell = table.rows[0].cells[index]
        cell.width = __import__("docx").shared.Inches(width)
        _set_cell_shading(cell, TEAL)
        _set_cell_text(cell, header, bold=True, size=8.5, colour="FFFFFF")
    _repeat_table_header(table.rows[0])
    for item in items:
        row = table.add_row()
        values = (item["minutes"], item["phase"], item["teacher_action"], item["student_action"], item.get("check_for_understanding") or "—")
        for index, value in enumerate(values):
            cell = row.cells[index]
            cell.width = __import__("docx").shared.Inches(widths[index])
            cell.vertical_alignment = __import__("docx").enum.table.WD_CELL_VERTICAL_ALIGNMENT.TOP
            _set_cell_text(cell, value, bold=index in {0, 1}, size=8.5)
        _keep_row_together(row)


def render_structured_lesson_plan_docx(document: StructuredLessonPlanDocument, *, teacher: str | None = None, meta: dict | None = None) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Mm, Pt, RGBColor

    meta = meta or {}
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(15)
    section.left_margin = Mm(15)
    section.right_margin = Mm(15)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("ELUME  |  LESSON PLAN")
    header_run.font.name = "Arial"
    header_run.font.size = Pt(8)
    header_run.font.color.rgb = RGBColor.from_string(EMERALD)

    footer_parts = ["Elume"]
    school_name = str(meta.get("schoolName") or "").strip()
    if teacher:
        footer_parts.append(teacher)
    if school_name:
        footer_parts.append(school_name)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("  |  ".join(footer_parts))
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor.from_string(SLATE)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run(document.title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor.from_string(EMERALD)

    metadata = [("Subject", document.subject), ("Level / year", document.level), ("Class / group", document.class_context), ("Duration", document.duration), ("Teacher", teacher), ("School", school_name or None)]
    metadata = [(label, value) for label, value in metadata if value]
    if metadata:
        table = doc.add_table(rows=1, cols=len(metadata))
        table.autofit = False
        width = Inches(6.7 / len(metadata))
        for index, (label, value) in enumerate(metadata):
            cell = table.cell(0, index)
            cell.width = width
            _set_cell_shading(cell, LIGHT_SLATE)
            _set_cell_margins(cell, top=90, start=100, bottom=90, end=100)
            label_paragraph = cell.paragraphs[0]
            label_paragraph.paragraph_format.space_after = Pt(1)
            label_run = label_paragraph.add_run(label.upper())
            label_run.bold = True
            label_run.font.name = "Arial"
            label_run.font.size = Pt(7.5)
            label_run.font.color.rgb = RGBColor.from_string(TEAL)
            value_paragraph = cell.add_paragraph()
            value_paragraph.paragraph_format.space_after = 0
            value_run = value_paragraph.add_run(str(value))
            value_run.font.name = "Arial"
            value_run.font.size = Pt(9)
            value_run.font.color.rgb = RGBColor.from_string(SLATE)
        _keep_row_together(table.rows[0])

    for block in document.blocks:
        block_type = block["type"]
        if block_type == "info_panel" and block.get("label") == "Primary learning outcome":
            _add_panel(doc, "PRIMARY LEARNING OUTCOME", block.get("text") or document.primary_outcome, fill=LIGHT_EMERALD)
        elif block_type == "info_panel" and block.get("definitions"):
            _add_definitions_table(doc, block["definitions"])
        elif block_type == "info_panel" and block.get("text"):
            _add_panel(doc, block["label"], block["text"], fill=LIGHT_EMERALD)
        elif block_type == "bullet_list":
            _section_heading(doc, block["title"])
            _add_bullets(doc, block["items"])
        elif block_type == "timeline":
            _add_lesson_flow(doc, block["items"])
        elif block_type == "student_task":
            _section_heading(doc, block["title"])
            _add_bullets(doc, block["items"])
        elif block_type == "assessment_checkpoint":
            _add_panel(doc, block["title"], "\n".join(f"• {item}" for item in block["items"]), fill=CYAN, title_colour=TEAL)
        elif block_type == "teacher_note":
            _add_panel(doc, block["title"], block["text"], fill=LIGHT_VIOLET, title_colour=VIOLET)
        elif block_type == "homework":
            _add_panel(doc, block["title"], block["text"], fill=LIGHT_SLATE, title_colour=SLATE)
        elif block_type == "callout":
            fill = LIGHT_AMBER if block.get("tone") == "warning" else CYAN
            title_colour = "A16207" if block.get("tone") == "warning" else TEAL
            text = block.get("text") or "\n".join(f"• {item}" for item in block.get("items", []))
            _add_panel(doc, block["title"], text, fill=fill, title_colour=title_colour)
        elif block_type == "heading":
            _section_heading(doc, block["text"])
        elif block_type == "paragraph":
            paragraph = doc.add_paragraph(block["text"])
            paragraph.paragraph_format.space_after = Pt(5)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
